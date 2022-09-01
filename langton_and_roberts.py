from asyncio.windows_events import NULL
from cProfile import label
from cgitb import text
import imghdr
from telnetlib import STATUS
import pyttsx3
import PyPDF2 
import datetime 
import datetime
from datetime import timedelta 
from distutils import text_file
import random
#from curses.textpad import Textbox
#from curses.textpad import Textbox
from distutils.command.config import config
from email.mime import image
from faulthandler import disable
from itertools import product
from msilib.schema import ComboBox
from msvcrt import setmode
from optparse import Values
from random import random
from sre_parse import State
import tkinter as tk
from tkinter import ANCHOR, W, Canvas, ttk
import time 
from tkinter import*
import sqlite3
from tkinter import colorchooser
from tkinter import messagebox
from tkinter import font
from tkinter.tix import IMAGE
from turtle import bgcolor, fillcolor, home, pd, position, title, width
from tkinter import filedialog as fd
from typing import Counter, Literal, Type
from tkinter import ttk
from unicodedata import category
from urllib import response
from click import command
from colorama import Cursor
from numpy import integer, record
from pyparsing import col
#from click import command 
from tkcalendar import*
from PIL import Image
from datetime import date
from tkinter import filedialog
import win32api
import os,sys
import csv
import win32print
import win32api
from datetime import datetime
from PIL import ImageTk, Image
from pydoc import doc
from turtle import width
from docx import Document
from docx.shared import Inches
import win32com.client
from PIL import Image, ImageTk
import csv
import random 
import cv2 
import numpy 

global engine  
global open_status_name
open_status_name=False
global selected
selected=False

LARGE_FONT= ("Verdana", 12)
NORM_FONT = ("Helvetica", 11)
SMALL_FONT = ("Helvetica", 10)
HEIGHT = 700
WIDTH = 500

Langton_and_Roberts="Langton & Roberts"
boading_rate_1=(float(10.00))
boading_rate_1=format(boading_rate_1,'.2f')

#def ring():
    #app.bell()
        
#,,,opens file search window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,    
def callback():
    name= fd.askopenfilename()
    print(name)
    


def popupmsg(msg):
    popup = Tk()
    #popup.overrideredirect(True)
    popup.wm_title("!")
    label = Label(popup, text=msg, font=NORM_FONT)
    label.pack(side="top", fill="x", pady=10,padx=10)
    B1 = Button(popup, text="Close",command = popup.destroy)
    B1.pack()
    app.bell()
    popup.resizable(False,False)
    popup.geometry('300x200+700+400')
    popup.mainloop()


#from PIL import ImageTk,Image
class HoverButton(Button):
    """ Button that changes color to activebackground when mouse is over it. """
    def __init__(self, master, **kw):
        super().__init__(master=master, **kw)
        self.default_Background = self.cget('background')
        self.hover_Background = self.cget('activebackground')
        self.bind('<Enter>', self.on_enter)
        self.bind('<Leave>', self.on_leave)

    def on_enter(self, e):
        self.config(background=self.hover_Background)

    def on_leave(self, e):
        self.config(background=self.default_Background)

#,,,,password window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
'''class Login:
    def __init__(self,root):
        self.root=root
        self.root.title("Login Screen")
        self.root.geometry("1199x600+200+80")
        self.root.resizable(False,False)
        self.root.overrideredirect(True)
        self.root.config(background="#a6a6a6")
        #self.root.after(3000)
        #background image
        #self.bg=PhotoImage(file="Picture1.png")
        #self.bg_image=Label(self.root,image=self.bg).place(x=0,y=0,relheight=1,relwidth=1)
        
        #,,,login frame.
        Frame_login=Frame(self.root,bg="white")
        Frame_login.place(x=360,y=100,width=440,height=380)
        lb1_copyright=Label(self.root,text="Copyright\nAll rights reserved",font=("Goudy old style",15,"bold"),fg="white",bg="#a6a6a6",justify=LEFT).place(x=20,y=540)
        lb2_copyright=Label(self.root,text="\nLangton & Roberts",font=("Goudy old style",15,"bold"),fg="white",bg="#a6a6a6",justify=LEFT).place(x=999,y=540)
        
        def register_user():
            global t1r
            global t2r
            global t3r
            #messagebox.showinfo("Dog Lick","You are unable to register at present.")
            window=tk.Tk()
            window.resizable(0,0)#deactavate maximise button
            #window.configure(bg="red")
            window.title("Register Window")
            window.geometry("470x220")
            l1=Label(window, text="Username:", font=("Arial", 15), bg="systembuttonface")
            l1.place(x=10, y=10)
            t1r=Entry(window, width=30, bd=1)
            t1r.place(x=200, y=10)
            
            l2=Label(window, text="Password:", font=("Arial", 15), bg="systembuttonface")
            l2.place(x=10, y=60)
            t2r=Entry(window, width=30, bd=1)
            t2r.place(x=200, y=60)
            
            l3=Label(window, text="Confirm Password:", font=("Arial", 15), bg="systembuttonface")
            l3.place(x=10, y=110)
            t3r=Entry(window, width=30, bd=1)
            t3r.place(x=200, y=110)
            
            b1=Button(window, text="Sign in", font=("Arial", 15), bg="#ffc22a",command=register_login_password)
            b1.place(x=170, y=150)
            
            
            window.mainloop()
            
        # this function will open a file and add name and password for checking against
        # https://youtu.be/tpGjHRDEjCE
        # look at https://youtu.be/geDjQUaTKAY for data base linking.
        def register_login_password():
            if t1r.get()!="" or t2r.get()!="" or t3r.get()!="":
                messagebox.showinfo("ERROR","Password did not match")
                if t2r.get()==t3r.get():
                    with open("langton_login.txt", "a") as f:
                        f.write(t1r.get()+","+t2r.get()+"\n")
                        messagebox.showinfo("LOGIN SUCSESSFUL", f"Welcome\nregester successull\n\nPress ok to open application")
                        #messagebox.showerror("ERROR","All Feilds are required",parent=self.root)
                
        

        def forgot_password():
            messagebox.showinfo("Dog Lick","Username:= Joe\nPassword:= password\ncase sensertive")
        
        def reg():
            messagebox.showinfo("Dog Lick","Not able to register at this time.")
        
        
        #,,,title and subtitle
        title=Label(Frame_login,text="Login Here",font=("Impact",35,"bold"),fg="orange",bg="white").place(x=65,y=30)
        subtitle=Label(Frame_login,text="Kennel Managment System.",font=("Goudy old style",15,"bold"),fg="#1d1d1d",bg="white").place(x=65,y=100)
        
        #,,username
        lb1_user=Label(Frame_login,text="Username",font=("Goudy old style",15,"bold"),fg="gray",bg="white").place(x=70,y=140)
        self.username=Entry(Frame_login,font=("Goudy old style",15),bg="#E7E6E6")
        self.username.place(x=70,y=170,width=320,height=35)
        
        #,,,password
        lb1_password=Label(Frame_login,text="Password",font=("Goudy old style",15,"bold"),fg="gray",bg="white").place(x=70,y=210)
        self.password=Entry(Frame_login,font=("Goudy old style",15),bg="#E7E6E6",show="*")
        self.password.place(x=70,y=240,width=320,height=35)

        #,,,button
        forget=Button(Frame_login,text="forgot password",bd=0,cursor="hand2",font=("Goudy old style",12,"bold"),fg="#6162FF",bg="white",command=forgot_password).place(x=70,y=280)
        register=Button(Frame_login,text="New User",bd=0,cursor="hand2",font=("Goudy old style",12,"bold"),fg="#6162FF",bg="white",command=reg).place(x=320,y=280)
        submit=Button(Frame_login,command=self.check_function,cursor="hand2",text="Login?",bd=0,font=("Goudy old style",15),bg="orange",fg="white").place(x=70,y=320,width=138,height=40)
        exit_button=Button(Frame_login,text="Exit",cursor="hand2",fg="white",font=("Goudy old style",15),bg="orange",bd=0,command=quit).place(x=250,y=320,width=138,height=40)
    
    
    def verify(self):
        try:
            with open("langton_login.txt","r") as f:
                info=f.readlines()
                i = 0
                for e in info:
                    user, password = e.split(",")
                    if user.strip() == self.username.get() and password.strip() == self.password.get():
                        messagebox.showinfo("LOGIN SUCSESSFUL", f"Welcome {self.username.get()}\nRegistration was successfull")
                        i=1
                        root.destroy()
                    if i==0:
                        messagebox.showerror("ERROR","Invalid Username or Password",parent=self.root)
        except:
            messagebox.showerror("ERROR","Please provide Username and Password",parent=self.root)
                
                
    
    def check_function(self):
        if self.username.get()=="" or self.password.get()=="":
            messagebox.showerror("ERROR","All Feilds are required",parent=self.root)
        elif self.username.get()!="Joe" or self.password.get()!="password":
            messagebox.showerror("ERROR","Invalid Username or Password",parent=self.root)
        else:
            messagebox.showinfo("LOGIN SUCSESSFUL", f"Welcome {self.username.get()}\nYou have successfully loged in\n\nPress ok to open application")
            root.destroy()
            #self.root.after(2000) 
        
                   
root=Tk()
obj=Login(root)
root.mainloop()'''
  
#,,,New Supplier Popup Window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,    
def settings_window():#,Popup window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    #,,,New Supplier Form Popup Window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,   
    #ring()
    settings_form = tk.Toplevel()
    frame_1 = Frame(settings_form, height=HEIGHT, width=WIDTH)
    frame_1.pack() 
    settings_form.title("New Supplier form(Enter Details)")
    settings_form.geometry('500x520+500+52')
    settings_form.resizable(False,False)
    #booking_form.overrideredirect(True)#removes header.
    
    #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    #,,Frames for Suppliers Form,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    master_frame=LabelFrame(frame_1,text="",font=LARGE_FONT)
    master_frame.grid(row=0,column=0,pady=20)
    contact_frame=LabelFrame(master_frame,text="Supplier Details",font=LARGE_FONT)
    contact_frame.pack(padx=10,pady=5,ipadx=10)#grid(row=0,column=0)
    main_contact_frame=LabelFrame(master_frame,text="Main Contact",font=LARGE_FONT)
    main_contact_frame.pack(padx=10,pady=5,ipadx=10)#grid(row=0,column=0)
    #notes_contact_frame=LabelFrame(master_frame,text="Notes",font=LARGE_FONT)
    #notes_contact_frame.pack(padx=10,pady=5,ipadx=5,ipady=5)#grid(row=0,column=0)
    button_frame=LabelFrame(master_frame,text="Comand Buttons")
    button_frame.pack(padx=5,pady=5)#grid(row=0,column=0)
    
    def clear_suppliers_box():
        ent_1.delete(0,END)
        ent_3.delete(0,END)
        ent_4.delete(0,END)
        ent_5.delete(0,END)
        ent_6.delete(0,END)
        ent_7.delete(0,END)
        ent_2.delete(0,END)
        combo_8.delete(0,END)
        ent_9.delete(0,END)
        ent_10.delete(0,END)
        ent_11.delete(0,END)
        ent_12.delete(0,END)
        ent_13.delete(0,END)
        ent_14.delete(0,END)

    #,,,supplier labels and entries,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    lbl_1 = Label(contact_frame, text='Supplier Name:')
    lbl_1.grid(row=0, column=0, padx=4, pady=10,sticky=W)
    ent_1 = Entry(contact_frame)
    ent_1.grid(row=0, column=1)
    lbl_2 = Label(contact_frame, text='ID:')
    lbl_2.grid(row=0, column=2, padx=4, pady=10,sticky=E)
    ent_2 = Entry(contact_frame,width=7,bd=0,bg="SystemButtonFace")
    ent_2.grid(row=0, column=3,sticky=W)
    lbl_3 = Label(contact_frame, text='Address:')
    lbl_3.grid(row=1, column=0, padx=4, pady=10,sticky=W)
    ent_3 = Entry(contact_frame)
    ent_3.grid(row=1, column=1)
    lbl_4 = Label(contact_frame, text='Town:')
    lbl_4.grid(row=1, column=2, padx=4, pady=10,sticky=E)
    ent_4 = Entry(contact_frame)
    ent_4.grid(row=1, column=3)
    lbl_5 = Label(contact_frame, text='City:')
    lbl_5.grid(row=2, column=0, padx=4, pady=10,sticky=W)
    ent_5 = Entry(contact_frame)
    ent_5.grid(row=2, column=1)
    lbl_6 = Label(contact_frame, text='Postcode:')
    lbl_6.grid(row=2, column=2, padx=4, pady=10,sticky=E)
    ent_6 = Entry(contact_frame)
    ent_6.grid(row=2, column=3)
    lbl_7 = Label(contact_frame, text='Website:')
    lbl_7.grid(row=3, column=0, padx=4, pady=10,sticky=W)
    ent_7 = Entry(contact_frame,width=53)
    ent_7.grid(row=3, column=1, columnspan=3,sticky=W)
    #,,,Main Contact,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    lbl_8 = Label(main_contact_frame, text='Title:')
    lbl_8.grid(row=0, column=0, padx=4, pady=10,sticky=W)
    lbl_14 = Label(main_contact_frame, text='Position held:')
    lbl_14.grid(row=0, column=2, padx=4, pady=10,sticky=E)
    ent_14 = Entry(main_contact_frame)
    ent_14.grid(row=0, column=3)
    combo_8=ttk.Combobox(main_contact_frame,width=7)
    combo_8['values']=['Mr', 'Miss', 'Mrs', 'Mis', 'Dr']
    combo_8.current(0)
    combo_8.grid(row=0,column=1,sticky=W)
    lbl_9 = Label(main_contact_frame, text='First Name:')
    lbl_9.grid(row=1, column=0, padx=4, pady=10,sticky=W)
    ent_9 = Entry(main_contact_frame)
    ent_9.grid(row=1, column=1)
    lbl_10 = Label(main_contact_frame, text='Last Name:')
    lbl_10.grid(row=1, column=2, padx=4, pady=10,sticky=E)
    ent_10 = Entry(main_contact_frame)
    ent_10.grid(row=1, column=3)
    lbl_11 = Label(main_contact_frame, text='Work Phone:')
    lbl_11.grid(row=2, column=0, padx=4, pady=10,sticky=W)
    ent_11 = Entry(main_contact_frame)
    ent_11.grid(row=2, column=1)
    lbl_12 = Label(main_contact_frame, text='Mobile:')
    lbl_12.grid(row=2, column=2, padx=4, pady=10,sticky=E)
    ent_12 = Entry(main_contact_frame)
    ent_12.grid(row=2, column=3)
    lbl_13 = Label(main_contact_frame, text='Email:')
    lbl_13.grid(row=3, column=0, padx=4, pady=10,sticky=W)
    ent_13 = Entry(main_contact_frame,width=53)
    ent_13.grid(row=3, column=1,columnspan=3,sticky=W)
    lbl_info = Label(master_frame, text='Supplier details can be found on invoices and reciepts:',fg="brown")
    lbl_info.pack()#.grid(row=4, column=0, padx=4, pady=10,sticky=W)
    
    #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    #,,Supplier database functions,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    #,,,add new supplier to database.
    def add_supplier():
        if ent_1.get()=="":
            popupmsg("Missing Data\n\nLooks like you forgot something\nSupplier name needed.")
        elif ent_3.get()=="":
            popupmsg("Missing Data\n\nLooks like you forgot something\nAddress needed.")
        elif ent_4.get()=="":
            popupmsg("Missing Data\n\nLooks like you forgot something\nTown needed.")
        elif ent_5.get()=="":
            popupmsg("Missing Data\n\nLooks like you forgot something\nCity needed.")
        elif ent_6.get()=="":
            popupmsg("Missing Data\n\nLooks like you forgot something\nPostcode needed.")
        elif ent_7.get()=="":
            popupmsg("Missing Data\n\nLooks like you forgot something\nWebsite needed\nIf not aplicable enter NA")
        else:    
            #,,, Connect to database. to add suppliers,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,add new suppliers.
            c.execute("INSERT INTO suppliers VALUES(:supplier_name,:address,:town,:city,:postcode,:website,:id,:title,:first_name,:last_name,:work_phone,:mobile_phone,:email,:position_held)",
                    {
                    "supplier_name":ent_1.get().capitalize(),
                    "address":ent_3.get().capitalize(),
                    "town":ent_4.get().capitalize(),
                    "city":ent_5.get().capitalize(),
                    "postcode":ent_6.get().capitalize(),
                    "website":ent_7.get(),
                    "id":ent_2.get(),
                    "title":combo_8.get(),
                    "first_name":ent_9.get().capitalize(),
                    "last_name":ent_10.get().capitalize(),
                    "work_phone":ent_11.get(),
                    "mobile_phone":ent_12.get(),
                    "email":ent_13.get(),
                    "position_held":ent_14.get().capitalize(),
                    })
            
            ent_1.delete(0,END)
            ent_3.delete(0,END)
            ent_4.delete(0,END)
            ent_5.delete(0,END)
            ent_6.delete(0,END)
            ent_7.delete(0,END)
            ent_2.delete(0,END)
            combo_8.delete(0,END)
            ent_9.delete(0,END)
            ent_10.delete(0,END)
            ent_11.delete(0,END)
            ent_12.delete(0,END)
            ent_13.delete(0,END)
            ent_14.delete(0,END)
            #,,,commit the changes
            conn.commit()
            #,,, close the connection
            conn.close()
            popupmsg('Your entrey has been saved!.')
            
    
    clear_all_button=HoverButton(button_frame,text="CLEAR ALL",font=LARGE_FONT,cursor="hand2",activebackground='#e2f723',command=clear_suppliers_box)#lambda: popupmsg('Not supported just yet!'))
    clear_all_button.grid(row=0,column=1,padx=2,pady=2)        
    save_button=HoverButton(button_frame,text="SAVE",cursor="hand2",font=LARGE_FONT,activebackground='#e2f723',command=add_supplier)#lambda: popupmsg('Not supported just yet!'))
    save_button.grid(row=0,column=2,padx=2,pady=2)
    exit_button=HoverButton(button_frame,text="EXIT FORM",font=LARGE_FONT,activebackground='orange',cursor="hand2",bg="#e2f723",fg="black",command=settings_form.destroy)
    exit_button.grid(row=0,column=3,padx=2,pady=2) 
             
#,,,New Client Registration Form Popup Window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
def quick_reg_window():#,Popup window
    #ring()
    quick_reg_window = Toplevel()
    quick_reg_window.title("New Client Register Form (For new client's only)")
    quick_reg_window.geometry('800x310+360+240')
    quick_reg_window.resizable(False,False)
    frame_1 = Frame(quick_reg_window, height=HEIGHT, width=WIDTH)
    frame_1.pack()
    #booking_form.overrideredirect(True)#,,removes top title from window.
    
    #,,Master frame Quick Client Register Form,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    master_frame=LabelFrame(frame_1,text='Quick Client Registration Form')
    master_frame.grid(row=0,column=0,padx=10,pady=10)
    contact_info=LabelFrame(master_frame,text="Contact information")
    contact_info.grid(row=0,column=0,padx=10,pady=10)
    button_command=LabelFrame(master_frame,text="What next buttons")
    button_command.grid(row=1,column=0,padx=10,pady=3,ipadx=3,ipady=3,sticky="E")
    info_notes=LabelFrame(master_frame,text="Employee Info:")
    info_notes.grid(row=1,column=0,padx=10,pady=10,ipadx=3,ipady=3,sticky="W")

    #,,,add new record to database.
    def save_client_record():
        if title_combo.get()=="":
            popupmsg("Title Needed")
        elif tb_2.get()=="":
            popupmsg("First name needed.")
        elif tb_3.get()=="":
            popupmsg("Please enter\nLast name.")
        elif tb_4.get()=="":
            popupmsg("Please enter\nAddress.")
        elif tb_5.get()=="":
            popupmsg("Please enter\nTown.")
        elif tb_6.get()=="":
            popupmsg("Please enter\ncity.")
        elif tb_7.get()=="":
            popupmsg("County needed.")
        elif tb_8.get()=="":
            popupmsg("Please enter\nPostcode.")
        elif tb_9.get()=="":
            popupmsg("Phone number needed.")
        elif tb_10.get()=="":
            popupmsg("Email needed.")
        else:
            #,,, Create a database or connect to one that executes.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,add new record.
            c.execute("INSERT INTO clients VALUES(:title,:first,:last,:id,:address,:town,:city,:county,:postcode,:phone,:email)",
                    {
                    "title":title_combo.get().capitalize(),
                    "first":tb_2.get().capitalize(),
                    "last":tb_3.get().capitalize(),
                    "id":tb_11.get(),
                    "address":tb_4.get().title(),
                    "town":tb_5.get().title(),
                    "city":tb_6.get().title(),
                    "county":tb_7.get().title(),
                    "postcode":tb_8.get().upper(),
                    "phone":tb_9.get(),
                    "email":tb_10.get(),
                    })
            #,,,commit the changes
            conn.commit()
            #,,, close the connection
            conn.close()
            #,,,Clear entry boxes.
            title_combo.delete(0,END)
            tb_2.delete(0,END)
            tb_3.delete(0,END)
            tb_4.delete(0,END)
            tb_5.delete(0,END)
            tb_6.delete(0,END)
            tb_7.delete(0,END)
            tb_8.delete(0,END)
            tb_9.delete(0,END)
            tb_10.delete(0,END)
            #pet1_ent.delete(0,END)
            #pet2_ent.delete(0,END)
            quick_reg_window.destroy()
            messagebox.showinfo("UPDATED", "New Client Successfully\nSaved to client database!")
            
            #my_tree.delete(*my_tree.get_children())
            #query_database()
    
    
    
    #,,,clear entry
    def clear_boxes():
        #,,,Clear entry boxes.
        title_combo.delete(0,END)
        tb_2.delete(0,END)
        tb_3.delete(0,END)
        tb_4.delete(0,END)
        tb_5.delete(0,END)
        tb_6.delete(0,END)
        tb_7.delete(0,END)
        tb_8.delete(0,END)
        tb_9.delete(0,END)
        tb_10.delete(0,END)
        
    
    
    lab_1=Label(contact_info,text="Title:")
    lab_1.grid(row=0,column=0,padx=3,pady=3,sticky="w")
    title_combo=ttk.Combobox(contact_info,width=7,state="readonly")
    title_combo['values']=["Dr", "Mr", "Mrs", "Miss", "Mis"]
    title_combo.grid(row=0,column=1,padx=3,pady=3,sticky="w")
    lab_2=Label(contact_info,text="First Name:")
    lab_2.grid(row=1,column=0,padx=3,pady=3,sticky="w")
    tb_2=Entry(contact_info,width=15)
    tb_2.grid(row=1,column=1,padx=3,pady=3,sticky="w")
    lab_3=Label(contact_info,text="Last Name:")
    lab_3.grid(row=1,column=2,padx=3,pady=3,sticky="w")
    tb_3=Entry(contact_info,width=15)
    tb_3.grid(row=1,column=3,padx=3,pady=3,sticky="w")
    lab_4=Label(contact_info,text="Address:")
    lab_4.grid(row=2,column=0,padx=3,pady=3,sticky="w")
    tb_4=Entry(contact_info,width=25)
    tb_4.grid(row=2,column=1,padx=3,pady=3,sticky="w")
    lab_5=Label(contact_info,text="Town:")
    lab_5.grid(row=2,column=2,padx=3,pady=3,sticky="w")
    tb_5=Entry(contact_info,width=25)
    tb_5.grid(row=2,column=3,padx=3,pady=3,sticky="w")
    lab_6=Label(contact_info,text="City:")
    lab_6.grid(row=2,column=4,padx=3,pady=3,sticky="w")
    tb_6=Entry(contact_info,width=28)
    tb_6.grid(row=2,column=5,padx=(3,10),pady=3,sticky="w")
    lab_7=Label(contact_info,text="County:")
    lab_7.grid(row=3,column=0,padx=3,pady=3,sticky="w")
    tb_7=Entry(contact_info,width=25)
    tb_7.grid(row=3,column=1,padx=3,pady=3,sticky="w")
    lab_8=Label(contact_info,text="Postcode:")
    lab_8.grid(row=3,column=2,padx=3,pady=3,sticky="w")
    tb_8=Entry(contact_info,width=10)
    tb_8.grid(row=3,column=3,padx=3,pady=3,sticky="w")
    lab_9=Label(contact_info,text="Mobile Phone:")
    lab_9.grid(row=4,column=0,padx=3,pady=3,sticky="w")
    tb_9=Entry(contact_info,width=15)
    tb_9.grid(row=4,column=1,padx=3,pady=3,sticky="w")
    lab_10=Label(contact_info,text="Email:")
    lab_10.grid(row=4,column=2,padx=3,pady=3,sticky="w")
    tb_10=Entry(contact_info,width=30)
    tb_10.grid(row=4,column=3,padx=3,pady=3,sticky="w")
    lab_11=Label(contact_info,text="")
    lab_11.grid(row=0,column=4,padx=3,pady=3,sticky="w")
    tb_11=Entry(contact_info,width=5,relief=FLAT,bg="SystemButtonFace")
    tb_11.grid(row=0,column=5,padx=3,pady=3,sticky="w")
    lab_12=Label(info_notes,text="Once completed continue to  New Pet Registration in the ribbon menu\nPet's are regestered seperate and given a uniqe ID.",fg="brown",font=SMALL_FONT)
    lab_12.grid(row=0,column=0,padx=3,pady=3,sticky="w")
    
    
    #,,,Buttons,Client regestration form,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    clear_button=HoverButton(button_command,text="CLEAR ALL",cursor="hand2",font=LARGE_FONT,activebackground='#e2f723',bg="systembuttonface",fg="black",command=clear_boxes)#lambda: popupmsg('Not supported just yet!'))
    clear_button.grid(row=0,column=2,padx=2,pady=2)#,sticky="E")
    save_button=HoverButton(button_command,text="SAVE",cursor="hand2",font=LARGE_FONT,activebackground='#e2f723',bg="systembuttonface",fg="black",command=save_client_record)#lambda: popupmsg('Not supported just yet!'))
    save_button.grid(row=0,column=1,padx=(5,2),pady=2)#,sticky="E")
    exit_button=HoverButton(button_command,text="EXIT FORM",cursor="hand2",font=LARGE_FONT,activebackground='orange',bg="#e2f723",fg="black",command=quick_reg_window.destroy)
    exit_button.grid(row=0,column=3,padx=2,pady=2,sticky="E")
    

#,,,New product form
def new_product_window():#,Popup window
    #app.bell()
    product_form = tk.Toplevel()
    product_form.title("Product Form (Stock)")
    product_form.geometry('460x345+410+104')
    product_form.resizable(False,False)
    
    #,,Frames,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    frame_1 = Frame(product_form, height=HEIGHT, width=WIDTH)
    frame_1.pack()#grid(row=0,column=0,padx=10,pady=10)
    prod_lbl_frame=LabelFrame(frame_1,text="Enter product details",font=20)#,bg="#e6f5f3")
    prod_lbl_frame.grid(row=0,column=0,padx=10,pady=20,ipadx=0,ipady=0)#,sticky=W)#fill="x",expand="yes"
    frame=LabelFrame(prod_lbl_frame,text='')#,bg="#e6f5f3")
    frame.grid(row=0,column=0,padx=5,pady=5,ipadx=5,ipady=5)
    
    def add_product():
            if descrip_ent.get()=="":
                popupmsg("It seems you left an entry blank\n\nPlease enter Product Name to continue.")
            elif sup_combo.get()=="":
                popupmsg("It seems you left an entry blank\n\nSupplier name needed, If not in the list\nregister Supplier first using the\nNew supplier Form.")
            elif cat_combo.get()=="":
                popupmsg("Please choose\na category.")
            elif cost_ent.get()=="":
                popupmsg("It seems you left an entry blank\n\nPlease enter\nCost Price.")
            elif retail_ent.get()=="":
                popupmsg("It seems you left an entry blank\n\nPlease enter\nRetail Price.")
            #elif lev_ent.get()=="":
                #popupmsg("It seems you left an entry blank\n\nPlease enter\nCurrent Stock Level\nJust enter 0 if not sure.")
            #elif minimum_ent.get()=="":
                #popupmsg("It seems you left an entry blank\n\nMin Stock Level Needed.")
            #elif maximum_ent.get()=="":
                #popupmsg("It seems you left an entry blank\n\nMax Stock Level Needed.")
            elif in_ent.get()=="":
                popupmsg("It seems you left an entry blank\n\nIncomeing Qt Needed.")
            elif barcode_ent.get()=="":
                popupmsg("It seems you left an entry blank\n\nPlease enter\nBarcode or Product ID.")
            else:
                #,,, Connect to database.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record.
                c.execute("INSERT INTO products_two VALUES(:product_name,:supplier,:category,:product_id,:cost_price,:retail_price,:inbound_qt,:stock_level,:min_stock_level,:max_stock_level,:barcode,:date_purchased)",
                        {
                        "product_name":descrip_ent.get().capitalize(),
                        "supplier":sup_combo.get().capitalize(),
                        "category":cat_combo.get().capitalize(),
                        "product_id":id_ent.get(),
                        "cost_price":cost_ent.get(),
                        "retail_price":retail_ent.get(),
                        "inbound_qt":in_ent.get(),
                        "stock_level":lev_ent.get(),
                        "min_stock_level":minimum_ent.get(),
                        "max_stock_level":maximum_ent.get(),
                        "barcode":barcode_ent.get().upper(),
                        "date_purchased":date_pur_picker.get_date(),
                        #"vat_inclusive":check_vat.getboolean(0),
                        })
            
                #,,,commit the changes
                conn.commit()
                #,,, close the connection
                conn.close()
            
                #,,,Clear entry boxes.
                descrip_ent.delete(0,END)
                cat_combo.set('')
                sup_combo.set('')
                id_ent.delete(0,END)
                cost_ent.delete(0,END)
                retail_ent.delete(0,END)
                lev_ent.delete(0,END)
                in_ent.delete(0,END)
                minimum_ent.delete(0,END)
                maximum_ent.delete(0,END)
                barcode_ent.delete(0,END)
                product_form.destroy()
                query_product_database()
                messagebox.showinfo("INFO","Saved\nProduct added to Inventory")
    
    
    #,,,Clear Add New Product entry boxes,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    def clear_products():
        #prod_save_button["state"]="active"
        #,,,Clear entry boxes.
        descrip_ent.delete(0,END)
        cat_combo.set('')
        sup_combo.set('')
        cost_ent.delete(0,END)
        retail_ent.delete(0,END)
        in_ent.delete(0,END)
        barcode_ent.delete(0,END)
        lev_ent.delete(0,END)
        minimum_ent.delete(0,END)
        maximum_ent.delete(0,END)
        id_ent.delete(0,END)
        #query_database()
    
    
    lbl_2=Label(frame,text='Name:')#,width=50,height=25,background="yellow")
    lbl_2.grid(row=0,column=0,padx=5,pady=5,sticky=W)#side="left")
    descrip_ent=Entry(frame,width=22)
    descrip_ent.grid(row=0,column=1,columnspan=2,padx=5,pady=5,sticky=W)
    lbl_3=Label(frame,text='Department:')#,width=50,height=25,background="yellow")
    lbl_3.grid(row=1,column=0,padx=5,pady=5,sticky=W)#side="left")
    cat_combo=ttk.Combobox(frame,width=15,state="readonly")
    cat_combo['values']=department_list
    cat_combo.current()
    cat_combo.grid(row=1,column=1,padx=5,pady=5,sticky=W)
    lev_lbl=Label(frame,text='')#,width=50,height=25,background="yellow")
    lev_lbl.grid(row=1,column=2,padx=5,pady=5,sticky=E)#side="left")
    lev_ent=Entry(frame,width=6,relief=FLAT,bg="systembuttonface")
    lev_ent.grid(row=1,column=3,padx=5,pady=5,sticky=W)
    lbl_4=Label(frame,text='Cost price: £')#,width=50,height=25,background="yellow")
    lbl_4.grid(row=2,column=0,padx=5,pady=5,sticky=W)#side="left")
    cost_ent=Entry(frame,width=6)
    cost_ent.grid(row=2,column=1,padx=5,pady=5,sticky=W)
    lbl_5=Label(frame,text='Retail price: £')#,width=50,height=25,background="yellow")
    lbl_5.grid(row=2,column=2,padx=5,pady=5,sticky=W)#side="left")
    retail_ent=Entry(frame,width=6)
    retail_ent.grid(row=2,column=3,padx=5,pady=5,sticky=W)
    lbl_6=Label(frame,text='Supplier:')#,width=50,height=25,background="yellow")
    lbl_6.grid(row=3,column=0,padx=5,pady=5,sticky=W)#side="left")
    sup_combo=ttk.Combobox(frame,width=15,state="readonly")
    sup_combo['values']=sup_names
    sup_combo.grid(row=3,column=1,padx=5,pady=5,sticky=W)
    lbl_min=Label(frame,text='')#,width=50,height=25,background="yellow")
    lbl_min.grid(row=3,column=2,padx=5,pady=5,sticky=W)#side="left")
    minimum_ent=Entry(frame,width=6,relief=FLAT,bg="systembuttonface")#,relief=FLAT)
    minimum_ent.grid(row=3,column=3,padx=5,pady=5,sticky=W)
    lbl_7=Label(frame,text='Quantity:')#,width=50,height=25,background="yellow")
    lbl_7.grid(row=4,column=0,padx=5,pady=5,sticky=W)#side="left")
    in_ent=Entry(frame,width=6)#,relief=FLAT)
    in_ent.grid(row=4,column=1,padx=5,pady=5,sticky=W)
    lbl_max=Label(frame,text='')#,width=50,height=25,background="yellow")
    lbl_max.grid(row=4,column=2,padx=5,pady=5,sticky=W)#side="left")
    maximum_ent=Entry(frame,width=6,relief=FLAT,bg="systembuttonface")#,relief=FLAT)
    maximum_ent.grid(row=4,column=3,padx=5,pady=5,sticky=W)
    lbl_8=Label(frame,text='Barcode:')#,width=50,height=25,background="yellow")
    lbl_8.grid(row=5,column=0,padx=5,pady=5,sticky=W)#side="left")
    barcode_ent=Entry(frame,width=13)
    barcode_ent.grid(row=5,column=1,padx=5,pady=5,sticky=W)
    date_pur=Label(frame,text='Date purchased')#,width=50,height=25,background="yellow")
    date_pur.grid(row=5,column=2,padx=5,pady=5,sticky=W)#side="left")
    date_pur_picker=DateEntry(frame,setmode="day",date_pattern="d/m/yyyy",state='readonly')
    date_pur_picker.grid(row=5,column=3)#grid(row=1,coloumn=2)
    id_lbl=Label(frame,text='')#,width=50,height=25,background="yellow")
    id_lbl.grid(row=0,column=2,padx=5,pady=5,sticky=E)#side="left")
    id_ent=Entry(frame,width=5,relief=FLAT,bg="systembuttonface")
    id_ent.grid(row=0,column=3,padx=5,pady=5,sticky=W)
    
    #,,,Buttons,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    # Command Frame in Products Popup,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    button_frame=LabelFrame(prod_lbl_frame,text="Comand Buttons")
    button_frame.grid(row=6,column=0,columnspan=5,padx=5,pady=5,sticky=SW)#fill="x",expand="yes"
    clear_button=HoverButton(button_frame,text='CLEAR',cursor="hand2",font=LARGE_FONT,activebackground= '#e2f723',command=clear_products)#lambda: popupmsg('Not supported just yet!'))
    clear_button.grid(row=0,column=0,padx=5,pady=5,ipadx=5,ipady=5)
    prod_save_button=HoverButton(button_frame,text='SAVE',cursor="hand2",font=LARGE_FONT,activebackground= '#e2f723',command=add_product)#lambda: popupmsg('Not supported just yet!'))
    prod_save_button.grid(row=0,column=1,padx=5,pady=5,ipadx=5,ipady=5)
    button_close=HoverButton(button_frame,text="EXIT FORM",cursor="hand2",font=LARGE_FONT,activebackground= 'orange',bg="#e2f723",fg="black",command=product_form.destroy)
    button_close.grid(row=0,column=2,padx=5,pady=5,ipadx=5,ipady=5)
    
#,,,Pet Booking in form Popup Window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,    
def pet_window():#,Popup window
    #app.bell()
    pet_window = Toplevel()
    pet_window.title("Pet Details Form.(can be filled in on arrival)")
    pet_window.geometry('695x310+400+104')
    pet_window.resizable(False,False)
    dt_stamp=time.strftime("%d/%m/%Y")
    #,,Frames,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    frame_1 = Frame(pet_window, height=HEIGHT, width=WIDTH)
    frame_1.grid(row=0,column=0,padx=5,pady=5)
    
    widget_box_frame = LabelFrame(frame_1)#frame for combo boxes.
    widget_box_frame.pack(padx=10,pady=10,ipadx=10)
    button_frame = Frame(widget_box_frame)
    button_frame.grid(row=7,column=0,columnspan=4,pady=10)
    
    #pet_info_lb_frame = LabelFrame(frame_1,text='Pet Registration & Information',font=LARGE_FONT)
    #pet_info_lb_frame.pack(padx=5,pady=5)
    
    
    #vacination_frame = LabelFrame(widget_box_frame,text='Vacinations')
    #vacination_frame.grid(row=6,column=0,columnspan=4,padx=10,pady=5,ipadx=2,sticky=W)
    
    #button_frame.pack(side="left",padx=5,pady=5,)
    
    click_get=StringVar()
    click_get.set("")
    vaccin_get=StringVar()
    vaccin_get.set("")
    
    
    def query_pet_database():
        global pet_tree
        #,,,clear the treeview
        for record in pet_tree.get_children():
            pet_tree.delete(record)
        #,,, Create a database or connect to one.
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,get row using ID,,,,,,,
        c.execute("SELECT rowid, * FROM customer_pets")
        records=c.fetchall()
        #,,,Add Record Entry Boxes,,,
        global count
        count = 0 #,,,start of ID
        
        for record in records:
            if count % 2==0:
                pet_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14],record[15],record[16],record[17]), tags=("evenrow",))
            else:
                pet_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14],record[15],record[16],record[17]), tags=("oddrow",))
            #,,,increment counter adds 1 to each record entered (ID number),,,,
            count +=1   
        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        #petinfo.set("")
    query_pet_database()
        
    
    
    
    def add_pet():
        if entry_1.get()=="":
            popupmsg("Owner Name needed.")
        elif pet_name_ent.get()=="":
            popupmsg("Pet name needed.")
        elif entry_6.get()=="":
            popupmsg("Ener a status.")
        elif entry_2.get()=="":
            popupmsg("Type of Pet needed.")
        elif entry_10.get()=="":
            popupmsg("Ener the Pet's age\nIn years (aprox).")
        elif entry_7.get()=="":
            popupmsg("Ener the Pet's D.O.B.")
        elif entry_3.get()=="":
            popupmsg("Please enter Breed.")
        elif entry_11.get()=="":
            popupmsg("Size of Pet Needed.")
        elif micro_combo.get()=="":
            popupmsg("Is the Pet Micro Chipped?.")
        elif entry_9.get()=="":
            popupmsg("Do you Have the chip number?\nEnter NO if NO.")
        elif pet_gender_ent.get()=="":
            popupmsg("Please enter Gender.")
        elif entry_12.get()=="":
            popupmsg("Please enter Weight.")
        elif entry_5.get()=="":
            popupmsg("Please enter Colour.")
        elif entry_13.get()=="":
            popupmsg("Error","Please choose temper.")
        #elif pet_gender_ent.get()=="":
            #popupmsg("Date Regirstered.")
        elif entry_health.get()=="":
            popupmsg("What is the Pet\nHealth status?.")
        elif attribute_ent.get()=="":
            popupmsg("Doe's the Pat have any habits?.")
        else:
            #,,, Connect to database.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,add new record.
            c.execute("INSERT INTO customer_pets VALUES(:pet_owner,:pet_name,:pet_age,:status,:pet_type,:pet_dob,:breed,:size,:micro_chip,:chip_number,:gender,:weight,:colour,:temper,:reg_date,:health,:attributes,:photo)",
                    {
                    "pet_owner":entry_1.get(),
                    "pet_name":pet_name_ent.get(),
                    "pet_age":entry_6.get(),
                    "status":entry_2.get(),
                    "pet_type":entry_10.get(),
                    "pet_dob":entry_7.get(),
                    "breed":entry_3.get(),
                    "size":entry_11.get(),
                    "micro_chip":micro_combo.get(),
                    "chip_number":entry_9.get(),
                    "gender":pet_gender_ent.get(),
                    "weight":entry_12.get().upper(),
                    "colour":entry_5.get(),
                    "temper":entry_13.get(),
                    "reg_date":dt_stamp,#date
                    "health":entry_health.get(),
                    "attributes":attribute_ent.get(),
                    "photo":NULL,
                    })
            #,,,commit the changes
            conn.commit()
            #,,, close the connection
            conn.close()
            #,,,Clear entry boxes.
            entry_1.set('')
            pet_name_ent.set('')
            entry_6.delete(0,END)
            entry_2.set('')
            entry_10.set('')
            entry_7.delete(0,END)
            entry_3.set('')
            entry_11.set('')
            micro_combo.set('')
            entry_9.delete(0,END)
            pet_gender_ent.set('')
            entry_12.delete(0,END)
            entry_5.set('')
            entry_13.set('')
            entry_health.current(5)
            attribute_ent.set('')
            pet_window.destroy()
            messagebox.showinfo("COMPLETED", "Pet details successfully saved!")
            #popupmsg("SAVED.")
            
        
    #,,,Clear entry boxes.
    def clear_pet():
        entry_1.set('')
        pet_name_ent.set('')
        entry_6.delete(0,END)
        entry_2.set('')
        entry_10.set('')
        entry_3.set('')
        entry_11.set('')
        micro_combo.set('')
        entry_9.delete(0,END)
        pet_gender_ent.set('')
        entry_12.delete(0,END)
        entry_5.set('')
        entry_13.set('')
        entry_health.current(5)
        attribute_ent.set('')
        entry_9.configure(bd=0,bg="systembuttonface",state="disabled")
        click_get.set("")
        ent9_text.set("")
        vaccin_combo.current(0)
        vac_name_combo.current()
        vaccin_get.set("")
        
        
    
    
    def vaccin_click(e):
        if vaccin_combo.get()== "Yes":
            vaccin_get.set("Vaccine name:")
            #vaccine_date = DateEntry(widget_box_frame,setmode="day",date_pattern="dd/mm/yyyy",state='readonly',width=11)
            #vaccine_date.grid(row=6,column=3,padx=3,pady=5,sticky=W)  
            vac_name_combo=ttk.Combobox(widget_box_frame,width=11,state="normal")#,state='readonly')
            vac_name_combo['values']=['Bordetella', 'DHPP', 'Leptosirosis', 'Parvovirus', 'Hepatitis', 'Distemper', 'Kennel Cough']
            vac_name_combo.current(0)
            vac_name_combo.grid(row=6,column=3,padx=3,pady=5,sticky=W) 
        elif vaccin_combo.get()== "No":
            vac_name_combo=ttk.Combobox(widget_box_frame,width=11,state='readonly')#,state='readonly')
            vac_name_combo['values']=['N/A']
            vac_name_combo.current(0)
            vac_name_combo.grid(row=6,column=3,padx=3,pady=5,sticky=W) 
            vaccin_get.set("")
        else:
            vaccin_get.set("")
            
    
    
    def combo_click(e):
        if micro_combo.get()== "Yes":
            click_get.set("Chip ID:")
            entry_9.delete(0,END)
            entry_9.configure(bg="white",fg="black",state="normal",bd=1)
        else:
            click_get.set("")
            entry_9.delete(0,END)
            entry_9.configure(bg="systembuttonface",fg="gray",bd=1)
            ent9_text.set("N/A")
            
    
    
    
    #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    #,,,create Pet table in database,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    #,,,create a cursor.
    c=conn.cursor()
    #,,,create Pet table.,,,,,,,,,,,,,,,,,,,,,,,
    c.execute("""CREATE TABLE if not exists customer_pets(
        pet_owner text,
        pet_name text,
        pet_age integer,
        status text,
        pet_type text,
        pet_dob integar,
        breed text,
        size text,
        micro_chip text,
        chip_number integer,
        gender text,
        weight real,
        colour text,
        temper text,
        reg_date integer integer,
        health text,
        attributes text,
        photo BLOB)
        """)
    #,,,commit the changes
    conn.commit()
    #,,,close the connection
    conn.close()
    
    #,,takes names from data base and appends to combo box,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    #,,,create a cursor.
    c=conn.cursor()
    owners=[]
    sql5=("SELECT rowid,  first_name, last_name FROM clients")
    #,,,get row using ID,,,,,,,
    c.execute(sql5)
    ids=c.fetchall()
    for i in ids:
        owners.append(str(i[1])+" "+i[2])#+" "+i[2])#,,i[0]=ID/i[1]=first_name/i[2]=last_name

    conn.commit()
    conn.close()
    pet=[]
    #,,takes pet names from Clients data base table and appends to combo box,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    #,,,create a cursor.
    c=conn.cursor()
    pet=[]
    sqlpet=("SELECT rowid,  pet_name FROM customer_pets")
    #,,,get row using ID,,,,,,,
    c.execute(sqlpet)
    ids=c.fetchall()
    for i in ids:
        pet.append(str(i[1]))#+" ")+i[2])
    conn.commit()
    conn.close()
        
    #,,labels and buttons,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    lbl_1 = Label(widget_box_frame,text='Pet Owner:')
    lbl_1.grid(row=0,column=0,padx=3,pady=5,sticky=W)
    entry_1 = ttk.Combobox(widget_box_frame,width=15,state='readonly')
    entry_1['values']=owners
    entry_1.grid(row=0,column=1,sticky=W)
    
    lbl_2 = Label(widget_box_frame,text='Status:')
    lbl_2.grid(row=1,column=0,padx=3,pady=5,sticky=W)
    entry_2=ttk.Combobox(widget_box_frame,width=11,state='readonly')
    entry_2['values']=['active', 'not active']
    entry_2.grid(row=1,column=1,padx=3,pady=5,sticky=W)
    
    lbl_3 = Label(widget_box_frame,text='Breed:')
    lbl_3.grid(row=2,column=0,padx=3,pady=5,sticky=W)
    entry_3=ttk.Combobox(widget_box_frame,width=15,state='readonly')
    entry_3['values']=['Grayhound', 'Mixed breed', 'Pussy']
    entry_3.current()
    entry_3.grid(row=2,column=1,padx=3,pady=5,sticky=W)
    
    lbl_4 = Label(widget_box_frame,text='Gender:')
    lbl_4.grid(row=3,column=0,padx=3,pady=5,sticky=W)
    pet_gender_ent=ttk.Combobox(widget_box_frame,width=7,state='readonly')
    pet_gender_ent['values']=['Dog', 'Bitch', 'Tom', 'Queen', 'Molly', 'Other']
    pet_gender_ent.current()
    pet_gender_ent.grid(row=3,column=1,padx=3,pady=5,sticky=W)
    
    lbl_5 = Label(widget_box_frame,text='Colour:')
    lbl_5.grid(row=4,column=0,padx=3,pady=5,sticky=W)
    entry_5=ttk.Combobox(widget_box_frame,width=15,state='readonly')
    entry_5['values']=['Red', 'Brindel', 'Black', 'White', 'Blue','Gray','Brown','Spotty','Patchy']
    entry_5.current()
    entry_5.grid(row=4,column=1,padx=3,pady=5,sticky=W)
    
    lbl_6 = Label(widget_box_frame,text='Pet age:')
    lbl_6.grid(row=0,column=4,padx=3,pady=5,sticky=W)
    entry_6 = ttk.Combobox(widget_box_frame,width=10,state='readonly')
    entry_6['values']=['0.5','1','2','3','4','5','6','7','8','9','10','11','12','13','14','15']
    entry_6.grid(row=0,column=5,sticky=W)
    
    lbl_7 = Label(widget_box_frame,text='Pet D.O.B:')
    lbl_7.grid(row=1,column=4,padx=3,pady=5,sticky=W)
    entry_7 = DateEntry(widget_box_frame,setmode="day",date_pattern="dd/mm/yyyy",state='readonly')
    entry_7.grid(row=1,column=5,sticky=W)
    
    lbl_8 = Label(widget_box_frame,text='Micro chip:')
    lbl_8.grid(row=2,column=4,padx=3,pady=5,sticky=W)
    micro_combo = ttk.Combobox(widget_box_frame,width=10,state='readonly')
    micro_combo['values']=['Yes','No']
    micro_combo.grid(row=2,column=5,sticky=W)
    micro_combo.bind("<<ComboboxSelected>>",combo_click)
    
    lbl_9 = Label(widget_box_frame,textvariable=click_get)
    lbl_9.grid(row=3,column=4,padx=3,pady=5,sticky=W)
    ent9_text=StringVar(value="")
    entry_9 = Entry(widget_box_frame,width=17,textvariable=ent9_text,bd=0,bg="systembuttonface",state="disabled")
    entry_9.grid(row=3,column=5,sticky=W)
    
    lbl_pet_name = Label(widget_box_frame,text='Pet name:')
    lbl_pet_name.grid(row=0,column=2,padx=3,pady=5,sticky=E)
    pet_name_ent = ttk.Combobox(widget_box_frame,width=15)
    pet_name_ent['values']=pet
    pet_name_ent.current()
    pet_name_ent.grid(row=0,column=3,sticky=W)
    
    lbl_10 = Label(widget_box_frame,text='Pet type:')
    lbl_10.grid(row=1,column=2,padx=(50,3),pady=5,sticky=E)
    entry_10 = ttk.Combobox(widget_box_frame,width=9,state='readonly')
    entry_10['values']=['Dog', 'Cat', 'Other']
    entry_10.grid(row=1,column=3,sticky=W)
    
    lbl_11 = Label(widget_box_frame,text='Size:')
    lbl_11.grid(row=2,column=2,padx=3,pady=5,sticky=E)
    entry_11=ttk.Combobox(widget_box_frame,width=9,state='readonly')
    entry_11['values']=['Giant', 'Large', 'Medium', 'Small', 'Toy']
    entry_11.grid(row=2,column=3,padx=3,pady=5,sticky=W)
    
    lbl_12 = Label(widget_box_frame,text='Wieght (kg):')
    lbl_12.grid(row=3,column=2,padx=3,pady=5,sticky=E)
    entry_12 = Entry(widget_box_frame,width=7)
    entry_12.grid(row=3,column=3,sticky=W)
    
    lbl_13 = Label(widget_box_frame,text='Temper:')
    lbl_13.grid(row=4,column=2,padx=3,pady=5,sticky=E)
    entry_13=ttk.Combobox(widget_box_frame,width=11,state='readonly')
    entry_13['values']=['Difficult', 'Exellent', 'Good', 'Nervous','Agresive']
    entry_13.current()
    entry_13.grid(row=4,column=3,padx=3,pady=5,sticky=W)
    
    lbl_health = Label(widget_box_frame,text='Health:')
    lbl_health.grid(row=5,column=0,padx=3,pady=5,sticky=W)
    entry_health=ttk.Combobox(widget_box_frame,width=11,state='readonly')
    entry_health['values']=['Un-well', 'Blind', 'Def', 'Old', 'Arthritus', 'Fit & well']
    entry_health.current(5)
    entry_health.grid(row=5,column=1,padx=3,pady=5,sticky=W)
    
    lbl_atributes = Label(widget_box_frame,text='Atributes:')
    lbl_atributes.grid(row=5,column=2,padx=3,pady=5,sticky=E)
    attribute_ent=ttk.Combobox(widget_box_frame,width=15,state='readonly')
    attribute_ent['values']=['N/A', 'Barks', 'Bites', 'Escapes', 'Not social','Fights', 'Needs company', 'Muzel']
    attribute_ent.current()
    attribute_ent.grid(row=5,column=3,padx=3,pady=5,sticky=W)
    
    lbl_v1 = Label(widget_box_frame,text='Vaccinated:')
    lbl_v1.grid(row=6,column=0,padx=3,pady=5,sticky=W)
    vaccine_date_lbl = Label(widget_box_frame,textvariable=vaccin_get)
    vaccine_date_lbl.grid(row=6,column=2,padx=2,pady=5,sticky=E)
    vaccin_combo=ttk.Combobox(widget_box_frame,width=11,state='readonly')
    vaccin_combo['values']=['No', 'Yes']
    vaccin_combo.current(0)
    vaccin_combo.grid(row=6,column=1,padx=3,pady=5,sticky=W)
    vaccin_combo.bind("<<ComboboxSelected>>",vaccin_click)
    
    vac_name_combo=ttk.Combobox(widget_box_frame,width=11,state='readonly')#,state='readonly')
    vac_name_combo['values']=['N/A']
    vac_name_combo.current(0)
    vac_name_combo.grid(row=6,column=3,padx=3,pady=5,sticky=W) 
    
    lbl_info = Label(widget_box_frame,text="Notice:\nPet's are linked to pet owners (client).\nplease register the client details first.\nthis can be found in the ribbon menu\nNew Client Registration Form.",fg="brown")
    lbl_info.grid(row=4,column=4,columnspan=2,rowspan=3,padx=10,pady=2)
    
    clear_button=HoverButton(button_frame,text="CLEAR",font=LARGE_FONT,cursor="hand2",activebackground='#e2f723',bg="systembuttonface",fg="black",command=clear_pet)#lambda: popupmsg('Not supported just yet!'))
    clear_button.grid(row=0,column=0,padx=3,pady=6)#,sticky="E")
    save_button=HoverButton(button_frame,text="SAVE",font=LARGE_FONT,cursor="hand2",activebackground='#e2f723',bg="systembuttonface",fg="black",command=add_pet)#add_pet)popupmsg('Not supported just yet!'))
    save_button.grid(row=0,column=1,padx=3,pady=6)#,sticky="E")
    
    exit_button=HoverButton(button_frame,text="EXIT FORM",font=LARGE_FONT,cursor="hand2",activebackground='orange',bg="#e2f723",fg="black",command=pet_window.destroy)
    exit_button.grid(row=0,column=2,padx=3,pady=6)#,sticky="E")
        
#,,,Quick Booking in Form Popup window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
def book_in_window():#,Popup window
    #app.bell()
    booking_window = Toplevel()
    booking_window.title("Quick Booking Form (Book Accomodation)")
    booking_window.geometry('805x470+323+104')
    booking_window.resizable(False,False)
    #booking_form.config(bg="lightblue")
    
    #,,Your Booking Frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    frame_1 = Frame(booking_window)#, height=HEIGHT, width=WIDTH)
    frame_1.grid(row=0,column=0,padx=0,pady=0)
    your_booking_frame=LabelFrame(frame_1,text="Your Booking",font=LARGE_FONT)
    your_booking_frame.grid(row=0,rowspan=4,column=0,padx=(20,5),pady=10,ipadx=0,ipady=2)#,fill="x",expand="yes")  
    #startup_data=[('test client', 'test pet', '1','Admin staff', 'Dog', 'No', 'No', 'Standard', '12/12/2021', '14/12/2021', '3', '10.00', '36.00', '12/12/2021')]
    dt_stamp=time.strftime("%d%m%Y-%H%M%S")
    #,,,connect to database,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    #,,,create a cursor.
    c=conn.cursor()
    #,,,create a bookings table.
    c.execute("""CREATE TABLE if not exists kennel_bookings(
        client_name text,
        pet_name text,
        enclosure integer,
        staff text,
        pet_type text,
        pick_up text,
        drop_off text,
        care_option text,
        in_date int,
        out_date int,
        units integer,
        unit_price real,
        total real,
        date_stamp integer
        )
        """)
    #,enters test startup data in database.
    #c.executemany("INSERT or IGNORE INTO kennel_bookings VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)",startup_data)
    #,,,commit the changes
    conn.commit()
    #,,,close the connection
    conn.close()
    
    #,, Today Time and date booking in form,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    def start():
        text=time.strftime("%d/%m/%Y")
        lb_date_today.config(text=text)
        lb_date_today.after(200,start)
    lb_date_today=Label(your_booking_frame,width=12)
    lb_date_today.grid(row=0,column=1,padx=5,pady=4)
    start()
    
    
    #,,takes client names from data base and appends to combo box,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    #,,,create a cursor.
    c=conn.cursor()
    options=[]#,,this is the name of the combobox values,,,,,,,,,,,,,,,
    sql_client=("SELECT rowid, first_name, last_name FROM clients ORDER BY last_name")
    #,,,get row using ID,,,,,,,
    c.execute(sql_client)
    ids=c.fetchall()
    for i in ids:
        options.append(str(i[1])+" "+i[2])#,,i[0]=ID/i[1]=first_name/i[2]=last_name
    conn.commit()
    conn.close()    
    
    #,,takes pet names from Clients database table and appends to combo box,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    c=conn.cursor()
    pet_names=[]
    sql_pet_name=("SELECT rowid,  pet_name FROM customer_pets")
    #,,,get row using ID,,,,,,,
    c.execute(sql_pet_name)
    ids=c.fetchall()
    for i in ids:
        pet_names.append(str(i[1]))#+i[2])
    conn.commit()
    conn.close()

    
    #,,takes names from Employee data base and appends to combo box,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    c=conn.cursor()
    staff=[]
    sql_employee=("SELECT rowid,  first_name, last_name FROM employees")
    #,,,get row using ID,,,,,,,
    c.execute(sql_employee)
    ids=c.fetchall()
    for i in ids:
        staff.append(str(i[1])+" "+i[2])
    conn.commit()
    conn.close()
           
    """#,,NOT WORKING?????,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    #,,function to put pet name in combobox linked from client name,,,,,,,,,,,,,,,,
    def get_pet_name(*args):
        pet_names=[]
        conn=sqlite3.connect("tree_crm.db")
        c=conn.cursor()
        pet_name_combo.set('')#clears combobox for next click.
        query="SELECT pet_name1 FROM clients WHERE first_name= '" +sel.get()+"'"
        my_data=c.execute(query)
        pet_names=[r for r, in my_data]
        pet_name_combo['values']=pet_names
    sel=tk.StringVar() 
    sel.trace("w",get_pet_name)"""
      
    #,,,Function to count days between dates,,,,,,,,,,,,,,,,,,,,,,,,,,
    def calculate(grab_out_date,grab_in_date):
        global comfirm_details
        global total_price2
        global plan
        global stay
        global arrival
        global depart
        dt_ref=time.strftime("%d%m%Y%H%M%S")
        difference=grab_in_date.get_date() - grab_out_date.get_date()
        lb_3_calc.config(text=f"{difference.days + 1}")
        time_stamp1=time.time()
        #print(time_stamp1)
        #print(time.ctime(time_stamp1))
        
        #,,, connect to database,,,,,,,,,,.
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,get company name from Database (should have only 1 row),,,,,,,
        c.execute("SELECT taxi_rate, standard_rate, luxury_rate, premium_rate FROM general_settings")
        record=c.fetchall()
        for i in record:
            taxi_rate=(i[0])
            standard_rate=(i[1])
            luxury_rate=(i[2])
            premium_rate=(i[3])
            
        #,,,commit the changes
        conn.commit()
        #,,, close the connection
        conn.close()
        
        pick_up=pet_collect_combo.get()
        drop_off=pet_doff_combo.get()
        
        #rates from database.
        both=(taxi_rate+taxi_rate)
        rate_1=standard_rate
        rate_2=luxury_rate
        rate_3=premium_rate
        
        #if statements to decide rate charge.
        if care_combo.get() == 'Standard':
            care=(rate_1)
        elif care_combo.get() == 'Luxury':
            care=(rate_2)
        else: 
            care=(rate_3)    
        #print(care)
        
        plan=(care)
        plan=float(plan)
        total_price=(int(difference.days + 1)*plan)
        total_price2=total_price/100*20+total_price
        format(total_price2,'.2f')
        stay=str(difference.days +1)
        arrival=date_in_picker.get()
        depart=date_out_picker.get()
        
        #if statements for taxi charge.
        if pick_up == 'Yes' and drop_off =='Yes':
            a=(f"Transport service (pick-up & drop-off) £{both:.2f}")
        elif pick_up == 'Yes' and drop_off == 'No':
            a=(f"Transport collection service £{taxi_rate:.2f}")
        elif pick_up == 'No' and drop_off == 'Yes':
            a=(f'Transport drop off service £{taxi_rate:.2f}')
        else:
            a=('Transport service not needed.')
        
        
        if client_combo.get()=="":
                popupmsg("Missing Entries?\nPlease Select Client Name?")
        elif employee_combo.get()=="":
                popupmsg("Missing Entries?\nPlease Enter Booked in by?\n\nSelect employee's name")
        else:
            greetings1= f"""{employee_combo.get()} is confirming that {client_combo.get()} 
has booked accommodation for {pet_name_combo.get()} 
to stay from {arrival} untill {depart} 
to be alocated to {name} {kennal_combo.get()} on arrival.
 
Charge will be {stay} days of {care_combo.get()} care plan at £{str(format(care,'.2f'))}+vat 
Total accomodation balance to pay at this time is £{str(format(total_price2,'.2f'))}

{a}
"""

            #shows message comfirming booking details            
            comfirm_details=Label(frame_1,text=greetings1,font=LARGE_FONT,fg="brown",justify=LEFT)
            comfirm_details.grid(row=1,column=1,columnspan=4,padx=20,pady=(5),sticky=W)
            confirm["state"]="normal"
            save["state"]="disabled"
            
    
    
            #collect records from database and place in tree view.
            global fetch_bookings
        def fetch_bookings():
            #,,,clear the treeview
            for record in calendar_tree.get_children():
                calendar_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute("SELECT rowid, * FROM kennel_bookings ORDER by in_date DESC")
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            
            for record in records:
                if count % 2==0:
                    calendar_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[4],record[3],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14]), tags=("evenrow",))
                else:
                    calendar_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[4],record[3],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14]), tags=("oddrow",))
                #,,,increment counter adds 1 to each record entered (ID number),,,,
                count +=1   
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            
    
    
    
    def go_ahead():
        if client_combo.get()=="":
            popupmsg("Client Name needed.")
        elif pet_name_combo.get()=="":
            popupmsg("Pet name needed.")
        elif employee_combo.get()=="":
            popupmsg("Booked in by needed.")
        else:
            #,,, Connect to database.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,add new record.
            c.execute("INSERT INTO kennel_bookings VALUES(:client_name,:pet_name,:enclosure,:staff,:pet_type,:pick_up,:drop_off,:care_option,:in_date,:out_date,:units,:unit_price,:total,:date_stamp)",
                    {
                    #"booking_ref":lb_6_ent.get(),
                    "client_name":client_combo.get(),
                    "pet_name":pet_name_combo.get(),
                    "enclosure":kennal_combo.get(),
                    "staff":employee_combo.get(),
                    "pet_type":pet_combo.get(),
                    "pick_up":pet_collect_combo.get(),
                    "drop_off":pet_doff_combo.get(),
                    "care_option":care_combo.get(),
                    "in_date":arrival,
                    "out_date":depart,
                    "units":stay,
                    "unit_price":plan,
                    "total":total_price2,
                    "date_stamp":dt_stamp#time.strftime("%d/%m/%Y-%H:%M:%S")
                    })
            conn.commit()
            conn.close()
            #,,,Clear/reset entry/combo boxes.
            comfirm_details.config(text="")
            confirm["state"]="disabled"
            client_combo.set('')
            employee_combo.delete(0,END)
            lb_3_calc.config(text="")
            kennal_combo.current()
            pet_name_combo.set('')
            pet_combo.current(0)
            pet_collect_combo.current(0)
            pet_doff_combo.current(0)
            care_combo.current(0)
            lb_6_ent.config(text="")
            fetch_bookings()
            save["state"]="normal"
            booking_window.destroy()
            fetch_bookings()
            messagebox.showinfo("Booking Saved",f"Booking ref: ACCOM-{dt_stamp}\n\nPlease continue to BOOKING RECORDS\nand confirm payment.")
            
        
    def cancel_comfirm():
        comfirm_details.config(text="")
        confirm["state"]="disabled"
        client_combo.set('')
        employee_combo.delete(0,END)
        lb_3_calc.config(text="")
        kennal_combo.current()
        pet_name_combo.set('')
        pet_combo.current(0)
        pet_collect_combo.current(0)
        pet_doff_combo.current(0)
        care_combo.current(0)
        lb_6_ent.config(text="")
        save["state"]="normal"  
        
     
    def cancel_boxes():
        client_combo.set('')
        pet_name_combo.set('')
        employee_combo.delete(0,END)
        lb_3_calc.config(text="")
        kennal_combo.set('')
        pet_combo.current(0)
        pet_collect_combo.current(0)
        pet_doff_combo.current(0)
        care_combo.current(0)
        lb_6_ent.config(text="")
        save["state"]="normal"  
        
          
    def get_pet_name(h):
        pass
    
    
    global enclosure
    enclosure=[]#,this list is dynamic and extends when number of enclosures entered.
    
    #,,takes max enclosure from general settings database table and appends to combo box in booking form,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    c=conn.cursor()
    runs=''
    sql_capacity=("SELECT max_capacity FROM general_settings")
    c.execute(sql_capacity)
    max_capacity=c.fetchone()
    runs=max_capacity
    
    for num in range(runs[0]):
        enclosure.append(num + 1)
    
    conn.commit()
    conn.close()
    
    #,,takes enclosure name from general settings database table and changes label text in booking form,,,,,,,,,,,,
    conn=sqlite3.connect("tree_crm.db")
    c=conn.cursor()
    sql_enc_name=("SELECT enclosure_name FROM general_settings")
    c.execute(sql_enc_name)
    enclosure_name=c.fetchone()
    global name
    name=enclosure_name[0]
    taxi=["No","Yes"]
    
    lb_date=Label(your_booking_frame,text='Todays Date:')#,width=50,height=25,background="yellow")
    lb_date.grid(row=0,column=0,padx=5,pady=4,sticky=E)#side="left")
    lbl_cust=Label(your_booking_frame,text='Client Name:')#,width=50,height=25,background="yellow")
    lbl_cust.grid(row=1,column=0,padx=5,pady=4,sticky=E)#side="left")
    client_combo=ttk.Combobox(your_booking_frame,width=17)#,textvariable=sel)
    client_combo['values']=options
    client_combo.current()
    client_combo.bind("<<ComboboxSelected>>",get_pet_name)
    client_combo.grid(row=1,column=1,padx=5,pady=4,sticky=E)
    lbl_pet=Label(your_booking_frame,text='Pet Name:')#,width=50,height=25,background="yellow")
    lbl_pet.grid(row=2,column=0,padx=5,pady=4,sticky=E)#side="left")
    pet_name_combo=ttk.Combobox(your_booking_frame,width=12)#,textvariable=sel)
    pet_name_combo['values']=pet_names
    
    #pet_name_combo.current()
    #pet_name_combo.bind("<<ComboboxSelected>>",get_pet_name)
    pet_name_combo.grid(row=2,column=1,padx=5,pady=4,sticky=E)
    lb_1=Label(your_booking_frame,text="Arival Date:")
    lb_1.grid(row=3,column=0,padx=5,pady=4,sticky=E)
    #date entry info #https://youtu.be/zn7I0xIsRTc
    date_in_picker=DateEntry(your_booking_frame,selectmode="day",date_pattern="dd/mm/yyyy",state='readonly')
    date_in_picker.grid(row=3,column=1,padx=5,pady=4,sticky=E)
    lb_2=Label(your_booking_frame,text="Checkout Date:")
    lb_2.grid(row=4,column=0,padx=5,pady=4,sticky=E)
    date_out_picker=DateEntry(your_booking_frame,setmode="day",date_pattern="dd/mm/yyyy",state='readonly')
    date_out_picker.grid(row=4,column=1,padx=5,pady=4,sticky=E)#grid(row=0,column=0,padx=5,pady=5) 
    lb_3=Label(your_booking_frame,text="Total Days:")
    lb_3.grid(row=5,column=0,padx=5,pady=4,sticky=E) 
    lb_3_calc=Label(your_booking_frame,text='',width=12)
    lb_3_calc.grid(row=5,column=1,padx=5,pady=4)
    global enc_name
    enc_name=(str(name)+':')
    lbl_5=Label(your_booking_frame,text=enc_name)#enclosure_ent.get())#,width=50,height=25,background="yellow")
    lbl_5.grid(row=6,column=0,padx=5,pady=4,sticky=E)#side="left")
    kennal_combo=ttk.Combobox(your_booking_frame,state='readonly',width=12)
    kennal_combo['values']=enclosure#['On Arrival', 'Holding', 'Quarantine']#Dynamic list.
    kennal_combo.current()
    kennal_combo.grid(row=6,column=1,padx=5,pady=4,sticky=E)
    lbl_4=Label(your_booking_frame,text='Booked in By:')#,width=50,height=25,background="yellow")
    lbl_4.grid(row=7,column=0,padx=5,pady=4,sticky=E)#side="left")
    global employee_combo
    employee_combo=ttk.Combobox(your_booking_frame,width=17)
    employee_combo['values']=staff
    employee_combo.current()
    employee_combo.grid(row=7,column=1,padx=5,pady=4,sticky=E)
    lbl_pet=Label(your_booking_frame,text='Pet Type:')#,width=50,height=25,background="yellow")
    lbl_pet.grid(row=8,column=0,padx=5,pady=4,sticky=E)#side="left")
    pet_combo=ttk.Combobox(your_booking_frame,state='readonly',width=12)
    pet_combo['values']=['Dog', 'Cat']
    pet_combo.current(0)
    pet_combo.grid(row=8,column=1,padx=5,pady=4,sticky=E)
    lbl_collect=Label(your_booking_frame,text='Pick Up:')#,width=50,height=25,background="yellow")
    lbl_collect.grid(row=9,column=0,padx=5,pady=4,sticky=E)#side="left")
    pet_collect_combo=ttk.Combobox(your_booking_frame,state='readonly',width=12)
    pet_collect_combo['values']=taxi
    pet_collect_combo.current(0)
    pet_collect_combo.grid(row=9,column=1,padx=5,pady=4,sticky=E)
    lbl_doff=Label(your_booking_frame,text='Drop Off:')#,width=50,height=25,background="yellow")
    lbl_doff.grid(row=10,column=0,padx=5,pady=4,sticky=E)#side="left")
    pet_doff_combo=ttk.Combobox(your_booking_frame,state='readonly',width=12)
    pet_doff_combo['values']=taxi
    pet_doff_combo.current(0)
    pet_doff_combo.grid(row=10,column=1,padx=5,pady=4,sticky=E)
    lb_care=Label(your_booking_frame,text="Care Option:")
    lb_care.grid(row=11,column=0,padx=5,pady=4,sticky=E) 
    care_combo=ttk.Combobox(your_booking_frame,state='readonly',width=17)
    care_combo['values']=['Standard', 'Premium', 'Luxury']
    care_combo.current(0)
    care_combo.grid(row=11,column=1,padx=5,pady=4,sticky=E)
    
    lb_6=Label(your_booking_frame,text="Booking Ref:")
    lb_6.grid(row=12,column=0,padx=5,pady=4,sticky=E) 
    lb_6_ent=Label(your_booking_frame,text="",width=14)
    lb_6_ent.grid(row=12,column=1,padx=5,pady=4,sticky=W)
    
    #,,Command Button Frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    command_button_frame=LabelFrame(your_booking_frame,text="")
    command_button_frame.grid(row=13,column=0,columnspan=2,padx=5,pady=0,ipadx=0,ipady=0,sticky=W)#,fill="x",expand="yes")
    save=HoverButton(command_button_frame,text="CONTINUE",cursor="hand2",activebackground="#e2f723",command=lambda:calculate(date_in_picker,date_out_picker))#lambda: popupmsg('Not supported just yet!'))
    save.grid(row=0,column=1,padx=3,pady=2)
    clear=HoverButton(command_button_frame,text="CLEAR ALL",cursor="hand2",activebackground="#e2f723",command=cancel_boxes)#lambda: popupmsg('Not supported just yet!'))
    clear.grid(row=0,column=2,padx=3,pady=2)
    exit=HoverButton(command_button_frame,text="EXIT FORM",cursor="hand2",activebackground='orange',bg="#e2f723",fg="black",command=booking_window.destroy)
    exit.grid(row=0,column=3,padx=3,pady=2)
    
    #,,,Comfirm your booking frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    comfirm_frame=LabelFrame(frame_1,text="Comfirming Your Booking",font=LARGE_FONT)#,width=515,height=100)
    comfirm_frame.grid(row=0,column=1,columnspan=2,padx=20,pady=(20,40),ipadx=2,ipady=2,sticky=N)#,fill="x",expand="yes")  
    comfirm_1=Label(comfirm_frame,text="Bookings can only be made for registered clients.\nunless in 'EMERGENCY' or 'GUEST'.\nPlease register your clients details first.",justify=LEFT)
    comfirm_1.grid(row=0,column=0,padx=5,pady=4,sticky=E) 
    
    cancel=HoverButton(frame_1,text="CANCEL",cursor="hand2",activebackground="#e2f723",command=cancel_comfirm)#lambda:popupmsg("Not yet completed"))
    cancel.grid(row=2,column=2,padx=3,pady=2,sticky=SE)
    confirm=HoverButton(frame_1,text="CONFIRM THIS BOOKING",cursor="hand2",bg="orange",fg="black",state="disabled",command=go_ahead)
    confirm.grid(row=2,column=3,padx=3,pady=2,sticky=SW)
         
#,,,,Tutorial step through popup windows,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
def tutorial():

    def page2():
        tut.destroy()
        tut2 = Tk()

        def page3():
            tut2.destroy()
            tut3 = Tk()

            tut3.wm_title("Part 3!")

            label = Label(tut3, text="Part 3", font=NORM_FONT)
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(tut3, text="Done!", command= tut3.destroy)
            B1.pack()
            tut3.mainloop()

        tut2.wm_title("Part 2!")
        label = Label(tut2, text="Part 2", font=NORM_FONT)
        label.pack(side="top", fill="x", pady=10)
        B1 = Button(tut2, text="Next", command= page3)
        B1.pack()
        tut2.mainloop()

    tut = Tk()
    tut.wm_title("Tutorial")
    
    label = Label(tut, text="What do you need help with?", font=NORM_FONT,width=50,height=2)
    label.pack(side="top")

    B1 = Button(tut, text = "Overview of the application", width=25,command=page2)
    B1.pack(pady=3)

    B2 = Button(tut, text = "How do I enter client?", width=25,command=lambda:popupmsg("Not yet completed"))
    B2.pack(pady=3)

    B3 = Button(tut, text = "How do I enter expenses?", width=25,command=lambda:popupmsg("Not yet completed"))
    B3.pack(pady=3)

    tut.mainloop()

#,,,,,Start of App Screen,,,(when App Starts),,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
class LangtonAndRoberts(Tk):#,,,startup config,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,

    def __init__(self, *args, **kwargs):
        Tk.__init__(self, *args, **kwargs)
        Tk.iconbitmap(self,default='home.ico')
        Tk.wm_title(self, Langton_and_Roberts)
        
        container = Frame(self)
        container.pack(side="top", fill="both", expand = True)
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)
        #,,,status bar at bottom,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        global status_bar
        status_bar=Label(self,text=Langton_and_Roberts,height=2,relief=FLAT,anchor=CENTER,bg="#e2f723",fg="brown")
        status_bar.pack(fill=X,side=BOTTOM,ipady=2,padx=0)#,expand=True)
        status_bar_text=StringVar()
        status_bar_text.set("This is a hidded Label")
        label2=Label(status_bar,textvariable=status_bar_text,bg="#e2f723",fg="black")
        label2.place(y=8,x=160)
        
        #,,Time ,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def start():
            text=time.strftime("%H:%M:%S %p")#("%A  %d/%m/%Y  %H:%M:%S %p")
            label.config(text=text)
            label.after(200,start)
        label=Label(status_bar,font=("ds-digital",9),bg="#e2f723",fg="black")
        label.place(y=8,x=40)
        start()
        
        
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create a table.
        c.execute("""CREATE TABLE if not exists suppliers(
            supplier_name text,
            address text,
            town text,
            city text,
            postcode text,
            website text,
            id integer,
            title text,
            first_name text,
            last_name text,
            work_phone integar,
            mobile_phone integar,
            email text,
            position_held text)
            """)

        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        
        
        def primary_colour():
            primary_colour=colorchooser.askcolor()[1]
        def secondary_colour():
            primary_colour=colorchooser.askcolor()[1]
        def highlight_colour():
            primary_colour=colorchooser.askcolor()[1]
        
        """def save_csv(records):
            with open("general_settings.csv",'a',newline='') as f:
                w=csv.writer(f,dialect="excel")
                for record in records:
                    w.writerow(record)
                    print(records)"""
        
        
        #https://youtu.be/Bn1n1diGv_0?list=PLCC34OHNcOtoC6GglhF3ncJ5rLwQrLGnV
        # opens popup file window
        def open_file():
            try:
                filepath = fd.askopenfilename(initialdir="booking_invoice_pdf_files",
                                            title="Select a Reciept to open",
                                            filetypes=[("PDF files","*.pdf")])#,("all files","*.*")])#[("text files","*.txt"),("CSV files","*.csv"),("all files","*.*")])
                
                if filepath:
                    #pdf_file=PyPDF2.PdfFileReader(filepath)
                    os.startfile(filepath)
                
                #print(file.read())
            except ValueError:
                messagebox.askretrycancel("FILE NOT FOUND","File did not open\nTry again?.")
                #text=None
            except FileNotFoundError:
                messagebox.showinfo("INFO","Information\nYou closed the dialog box without choseing a file")
                #text=None
                
                
        global print_file
        #,, print file function,,,Basic file printing.
        def print_file():
            #printer_name=win32print.GetDefaultPrinter()
            #status_bar.config(text=printer_name)
            #,,grab filename any file.
            file_to_print=filedialog.askopenfilename(initialdir="C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\",
                                                     title="Open file to print",
                                                     filetypes=[("Text Files","*.txt"),
                                                                ("CSV Files","*.csv"),
                                                                ("All Files","*.*")])
            if file_to_print:
                win32api.ShellExecute(0,"print",file_to_print,None,".",0)
                print("done but your printer is not connected")#prints to termanal
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,,,,,,,,,,dropdown Menu Bars ,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,File Tab,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        menubar = Menu(container)
        filemenu = Menu(menubar, tearoff=0)#,could be True or False.
        menubar.add_cascade(label="File", menu=filemenu)
        #filemenu.add_command(label="New", command=lambda: popupmsg('Not supported just yet!'))
        filemenu.add_command(label="Reciepts", command=open_file)#callback)#lambda: popupmsg('Not supported just yet!'))
        #filemenu.add_command(label="Save", command=lambda:popupmsg('Not supported just yet!'))
        filemenu.add_command(label="Save As", command=lambda:popupmsg('Not supported just yet!'))
        filemenu.add_separator()
        filemenu.add_command(label="Print File", command=print_file)#lambda: popupmsg('Not supported just yet!'))
        filemenu.add_separator()
        #filemenu.add_command(label="Exit", command=quit)
        #,,,Search Tab,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        search_menu = Menu(menubar, tearoff=0)
        global lookup_records
        menubar.add_cascade(label="Search", menu=search_menu)
        search_menu.add_command(label="Search Client",command=lambda:popupmsg('Not supported just yet!'))#lookup_records)
        search_menu.add_command(label="Search Pet",command=lambda: popupmsg('Not supported just yet!'))#lookup_records)
        search_menu.add_separator()
        search_menu.add_command(label="Show all clients",command=lambda: popupmsg('Not supported just yet!'))#query_database)
        #,,,Options Tab,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #options_menu = Menu(menubar, tearoff=0)
        #menubar.add_cascade(label = "Options", menu = options_menu)
        #options_menu.add_command(label = "Primary Colour",command=primary_colour)
        #options_menu.add_command(label = "Secondary Colour",command=secondary_colour)
        #options_menu.add_command(label = "Highlight Colour",command=highlight_colour)
        #,,,Quick Tab,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        quick_form_menu = Menu(menubar, tearoff=0)
        menubar.add_cascade(label = "Quick Forms", menu = quick_form_menu)
        quick_form_menu.add_command(label="New Supplyer form", command=settings_window)
        quick_form_menu.add_command(label="New Client form", command=quick_reg_window)
        quick_form_menu.add_command(label="New Pet form", command=pet_window)
        quick_form_menu.add_command(label="New Booking form", command=book_in_window)#lambda: popupmsg('Not supported just yet!'))
        #,,,Help Tab,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        helpmenu = Menu(menubar, tearoff=0)
        helpmenu.add_command(label="Tutorial", command= tutorial)
        menubar.add_cascade(label="Help", menu=helpmenu)
        Tk.config(self, menu=menubar)
        
        self.frames = {}

        for F in (StartPage, PageOne, PageTwo, PageThree):

            frame = F(container, self)
            self.frames[F] = frame
            frame.grid(row=0, column=0, sticky="nsew")

        self.show_frame(StartPage)

    def show_frame(self, cont):
        frame = self.frames[cont]
        frame.tkraise()
        
#,,,,,Home Page ,,,Anything in TABS,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,      
class StartPage(Frame):#,,All Tab Windows (Services,Products,),,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    def __init__(self, parent, controller):
        Frame.__init__(self,parent)
        #,,Buttons Frame(left),,Center Frame(left),Ribbon Frame(Top),,,,,,,
        can_2=Frame(self,relief=FLAT,bd=0,bg="white")
        can_2.pack(side='left',anchor=W,fill='both')#,expand=True)
        #,,,Ribbon (menu) frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        can_4=Frame(self,relief=FLAT,bd=0,bg="white",height=6)
        can_4.pack(side='top',anchor=W,fill='x')#,expand=True)
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,Frame for Notebook,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        can_3=Frame(self,relief=FLAT)
        can_3.pack(side='left',anchor=W,fill='both',expand=True)  
        #,,,,Home button widget text,,,,
        employee_num_home=StringVar()
        employee_num_home.set("Employee's ")
        pet_num_home=StringVar()
        pet_num_home.set("Pets's ")
        client_num_home=StringVar()
        client_num_home.set("Client's ")
        arrival_num_home=StringVar()
        arrival_num_home.set("Arrival's ")
        checkout_num_home=StringVar()
        checkout_num_home.set("Checkout's ")
        occupancy_num_home=StringVar()
        occupancy_num_home.set("Kennel Occupancy ")
        
        #,,, connect to database,,,,,,,,,,.
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,get company name from Database (should have only 1 row),,,,,,,
        c.execute("SELECT business_name FROM business_details")
        record=c.fetchall()
        for name in record:
                company_name=name[0]#makes result into a variable.
       
        #,,,commit the changes
        conn.commit()
        #,,, close the connection
        conn.close()
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create database for products,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create a products table.
        c.execute("""CREATE TABLE if not exists products_two(
            product_name text,
            supplier text,
            category text,
            product_id int,
            cost_price REAL,
            retail_price REAL,
            inbound_qt int,
            stock_level int,
            min_stock_level int,
            max_stock_level int,
            barcode text,
            date_purchased int)
            """)

        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        #global tab_0
        
        
        def talk_product():
            global engine
            if p4_ent.get()=="" :
                engine=pyttsx3.init()
                engine.say(f'''you need to select a product first before you can delete a record,
                           you can search using the product barcode''')
                #engine.say(my_entry.get())
                #engine.say(my_entry2.get())
                #engine.say(my_entry3.get())
            else:
                engine=pyttsx3.init()
                engine.say(f''' CAUTION,, you are removing, product reference number{p4_ent.get()}, {p1_ent.get()}, with barcode, {p9_ent.get()}''')
            engine.runAndWait()
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #create notebook tabs for can_3 frame.,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        notebook=ttk.Notebook(can_3)
        tab_0=Frame(notebook,background="white")#'#e1e4fa')#,background="#a2bee8")#new frame for tab_0
        tab_1=Frame(notebook)#,background='#e6f5f3')#,background="yellow")#new frame for tab_1
        tab_2=Frame(notebook)#,background='#e6f5f3')#,background="gray")#new frame for tab_2
        tab_3=Frame(notebook)#,background='#e6f5f3')#,background="black")#new frame for tab_3
        tab_4=Frame(notebook)#,background='#e6f5f3')#,background="black")#new frame for tab_3
        tab_5=Frame(notebook)#,background='#e6f5f3')#,background="black")#new frame for tab_3
        tab_6=Frame(notebook)#,background='#e6f5f3')#,background="black")#new frame for tab_3
        #self.imagrObj=PhotoImage(file='save.ico')
        notebook.add(tab_0,text=f'{company_name}')
        notebook.add(tab_1,state="hidden",text="INVENTIRY")
        notebook.add(tab_2,state="hidden",text="SERVICES")
        notebook.add(tab_3,state="hidden",text="PET INFORMATION")#image=self.imagrObj,compound="left")
        notebook.add(tab_4,state="hidden",text="SELL SOMETHING")
        notebook.add(tab_5,state="hidden",text="REPORTS")
        notebook.add(tab_6,state="hidden",text="EMPLOYEES")
        notebook.pack(expand=True,fill="both")
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create Services table in database,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create Service table and headings.,,,,,,,,,,,,,,,,,,,,,,,
        c.execute("""CREATE TABLE if not exists service(
            service_name text NOT NULL,
            description text,
            rate DECIMAL(3,2) NOT NULL,
            department_id text NOT NULL,
            FOREIGN KEY ('department_id')
            REFERENCES department (department_id)
            )""")

        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create Departments table in database,,,,,,,,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create department table and headings.,,,,,,,,,,,,,,,,,,,,,,,
        c.execute('''CREATE TABLE if not exists department(
            department_id integer,
            department_name text NOT NULL,
            PRIMARY KEY("department_id" AUTOINCREMENT)
            )''')

        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,takes Department names from Departments data base table and appends to combo box,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        global department_list
        department_list=[]
        sql6=("SELECT rowid,  department_name FROM department")
        #,,,get row using ID,,,,,,,
        c.execute(sql6)
        ids=c.fetchall()
        for i in ids:
            department_list.append(str(i[0])+" "+i[1])
            
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,takes Service names from Services data base table and appends to combo box,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        global service_list_get
        service_list_get=[]
        sql7=("SELECT rowid,  service_name FROM service")
        #,,,get row using ID,,,,,,,
        c.execute(sql7)
        ids=c.fetchall()
        for i in ids:
            service_list_get.append(str(i[1])+" ")#+i[2])
        
        
        # ,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,takes names from Supplier data base and appends to combo box,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        global sup_names
        sup_names=[]
        sql4=("SELECT rowid,  supplier_name, last_name FROM suppliers")
        #,,,get row using ID,,,,,,,
        c.execute(sql4)
        ids=c.fetchall()
        for i in ids:
            sup_names.append(str(i[1]))#+" "+i[2])
        
        
        global query_product_database
        #,,,refreshes the table,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def query_product_database():
            #,,,clear the treeview
            for record in product_tree.get_children():
                product_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute("SELECT rowid, * FROM products_two")
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            for record in records:
                if count % 2==0:
                    product_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12]), tags=("evenrow",))
                else:
                    product_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12]), tags=("oddrow",))
                #,,,increment counter adds 1 to each record entered (ID number),,,,
                count +=1   
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            #print("Table view done")
            p1_ent.delete(0,END)
            p2_ent.delete(0,END)
            p3_ent.delete(0,END)
            p4_ent.delete(0,END)
            p5_ent.delete(0,END)
            p6_ent.delete(0,END)
            p7_ent.delete(0,END)
            p8_ent.delete(0,END)
            p9_ent.delete(0,END)
            #min_ent.delete(0,END)
            #max_ent.delete(0,END)
            #stock_level_ent.delete(0,END)
            stock_date_ent.delete(0,END)
            prod_help.set("")
            
            
        #,,,update product Details,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def update_product_details():
            if p1_ent.get()=="":
                messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            elif p2_ent.get()=="":
                messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            elif p3_ent.get()=="":
                messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            elif p4_ent.get()=="":
                messagebox.showerror("Missing Data!","ID cannot be blank!.\nSellect a product to update.\nthen amend field's as needed.")
            elif p5_ent.get()=="":
                messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            elif p6_ent.get()=="":
                messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            elif p7_ent.get()=="":
                messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            #elif p8_ent.get()=="":
                #messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            elif p9_ent.get()=="":
                messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            elif stock_date_ent.get()=="":
                messagebox.showerror("Missing Data!","Sellect a product to update.\nthen amend field's as needed.")
            else:
                #,,, Connect to database.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record.(UPDATE) sercice 
                c.execute(""" UPDATE products_two SET 
                          product_name=:product_name, 
                          supplier=:supplier, 
                          category=:category, 
                          product_id=:product_id,
                          cost_price=:cost_price,
                          retail_price=:retail_price,
                          inbound_qt=:inbound_qt,
                          stock_level=:stock_level,
                          barcode=:barcode,
                          date_purchased=:date_purchased
                    
                          WHERE oid = :oid """,
                            {
                            "product_name":p1_ent.get(),
                            "supplier":p2_ent.get(),
                            "category":p3_ent.get(),
                            "product_id":p4_ent.get(),
                            "cost_price":p5_ent.get(),
                            "retail_price":p6_ent.get(),
                            "inbound_qt":p7_ent.get(),
                            "stock_level":p8_ent.get(),
                            "barcode":p9_ent.get(),
                            "date_purchased":stock_date_ent.get(),
                            "oid":p4_ent.get(),
                            })
                p1_ent.delete(0,END)
                p2_ent.delete(0,END)
                p3_ent.delete(0,END)
                p4_ent.delete(0,END)
                p5_ent.delete(0,END)
                p6_ent.delete(0,END)
                p7_ent.delete(0,END)
                p8_ent.delete(0,END)
                p9_ent.delete(0,END)
                stock_date_ent.delete(0,END)
                last_sale_ent.delete(0,END)
                #,,,commit the changes
                conn.commit()
                #,,, close the connection
                conn.close()
                messagebox.showinfo("UPDATED", "Inventory record has been\nSuccessfully Updated")
            query_product_database()
            
                
        #,,Delete a Product from Database,,,,,,,,,,,,,,,,,,,,,
        def remove_product(*args):
            #talk_product()
            #app.bell()
            if p1_ent.get()=="" or p2_ent.get()=="" or p3_ent.get()=="" or p4_ent.get()=="" \
                or p5_ent.get()=="" or p6_ent.get()=="" or p7_ent.get()=="" or p9_ent.get()=="":
                #talk_product()
                messagebox.showerror("ERROR","Missing entry?\nChoose a record to delete")
            else:
                talk_product()
                response = messagebox.askyesno("WARNING!!", f"You are Deleting {p1_ent.get()}\nBarcode: {p8_ent.get()}\nAre You Sure?!.")
                #,,,add logic for message box.
                if response==1:
                    #,,,desegnate selections
                    x=product_tree.selection()
                    #,,,create list of ID's.
                    ids_to_delete=[]
                    #,,,add selections to ids_to_delete list.
                    for record in x:
                        ids_to_delete.append(product_tree.item(record,"values")[3])
                    #,,,Delete from treeview.
                    for record in x:
                        product_tree.delete(record)
                    #,,,create database.    
                    conn=sqlite3.connect("tree_crm.db")
                    #,,,create a cursor.
                    c=conn.cursor()
                    #,,,delete everything from table
                    c.executemany("DELETE FROM products_two WHERE oid = ?",[(a,)for a in ids_to_delete])
                    #,,,commit the changes
                    conn.commit()
                    #,,,close the connection
                    conn.close()
                    #,,,clear entries from boxes if filled.
                    clear_product_view()
                    query_product_database()
                    messagebox.showinfo("SELECTION DELETED", "Records have been Deleted")
                    

        #,,,Products frames and buttons,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        top_label_frame1=Frame(tab_1,bd=1)  
        top_label_frame1.pack(pady=0,padx=0,anchor=W,fill=X)
        heading=Label(top_label_frame1,text="INVENTORY & PRODUCTS",font="bold 20",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=17)
        button_label_frame=Frame(tab_1,bg="light blue",bd=1)  
        button_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        box_frame3=Frame(tab_1)  
        box_frame3.pack(pady=0,padx=0,anchor=W,fill=X)
        
        
        def help_prod_info():
            a=(f'Select REFRESH to display all records\n')
            b=(f"then select a record to display it's information.\n\n")
            c=(f'Select SEARCH and enter the product Barcode to find an entry.\n\n')
            d=(f"Select ADD PRODUCT to add an entry.\n\n")
            e=(f'WARNING... DELETE will permanently remove from database.\n')
            f=(f'To remove a record, select a record in view, click DELETE RECORD.\n')
            g=(f'(you will have the option to continue or cancel the deletion.)\n\n')
            h=(f'To amend an error, select a record, amend in the Product Overview\nthen select UPDATE DETAILS.')
            
            answer=a+b+c+d+e+f+g+h
            prod_help.set(answer)
        
        
        prod_help=StringVar()
        prod_help.set("")
        prod_help_lbl=Label(box_frame3,textvariable=prod_help,font="Verdana, 12",fg="#347083",justify=LEFT)#,bg="#e6f5f3")#,width=50,height=25,background="yellow")
        prod_help_lbl.grid(row=1,column=0,padx=20,sticky=W)
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create Product Tree view frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        prod_label_frame=LabelFrame(box_frame3,text="",bd=0,font=20)#,bg="#e6f5f3")
        prod_label_frame.grid(row=0,column=0,columnspan=2,padx=10,pady=0,sticky=W)#pack(padx=30,anchor=W)
        product_tree_frame=Frame(prod_label_frame)  
        product_tree_frame.grid(row=0,column=0,padx=5,pady=5)#pack(pady=10,padx=10)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,Product overview frame and entries,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        supplier_frame=LabelFrame(box_frame3,text="Product overview",font=20)#,bg="#e6f5f3")
        supplier_frame.grid(row=0,rowspan=4,column=2,padx=10,pady=30,sticky=N)
        
        #,,Labels and entry boxes in Product Overview frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        p1_lbl=Label(supplier_frame,text='Product Name:')#,bg="#e6f5f3")#,width=50,height=25,background="yellow")
        p1_lbl.grid(row=0,column=0,padx=5,pady=5,sticky=E)
        p1_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p1_ent.grid(row=0,column=1,padx=5,pady=5)
        p2_lbl=Label(supplier_frame,text='Supplier:')#,width=50,height=25,background="yellow")
        p2_lbl.grid(row=1,column=0,padx=5,pady=5,sticky=E)
        p2_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p2_ent.grid(row=1,column=1,padx=5,pady=5)
        p3_lbl=Label(supplier_frame,text='Department:')#,width=50,height=25,background="yellow")
        p3_lbl.grid(row=2,column=0,padx=5,pady=5,sticky=E)
        p3_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p3_ent.grid(row=2,column=1,padx=5,pady=5)
        p4_lbl=Label(supplier_frame,text='ID:')#,width=50,height=25,background="yellow")
        p4_lbl.grid(row=3,column=0,padx=5,pady=5,sticky=E)
        p4_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p4_ent.grid(row=3,column=1,padx=5,pady=5)
        p5_lbl=Label(supplier_frame,text='Single Unit Cost: £')#,width=50,height=25,background="yellow")
        p5_lbl.grid(row=4,column=0,padx=5,pady=5,sticky=E)
        p5_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p5_ent.grid(row=4,column=1,padx=5,pady=5)
        p6_lbl=Label(supplier_frame,text='Retail Price: £')#,width=50,height=25,background="yellow")
        p6_lbl.grid(row=5,column=0,padx=5,pady=5,sticky=E)
        p6_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p6_ent.grid(row=5,column=1,padx=5,pady=5)
        p7_lbl=Label(supplier_frame,text='Inbound Qt:')#,width=50,height=25,background="yellow")
        p7_lbl.grid(row=6,column=0,padx=5,pady=5,sticky=E)
        p7_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p7_ent.grid(row=6,column=1,padx=5,pady=5)
        p8_lbl=Label(supplier_frame,text='Current Inventory:')#,width=50,height=25,background="yellow")
        p8_lbl.grid(row=7,column=0,padx=5,pady=5,sticky=E)
        p8_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p8_ent.grid(row=7,column=1,padx=5,pady=5)
        p9_lbl=Label(supplier_frame,text='Barcode:')#,width=50,height=25,background="yellow")
        p9_lbl.grid(row=8,column=0,padx=5,pady=5,sticky=E)
        p9_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")
        p9_ent.grid(row=8,column=1,padx=5,pady=5)
        
        lbl_13=Label(supplier_frame,text='Purchased:')#,width=50,height=25,background="yellow")
        lbl_13.grid(row=9,column=0,padx=5,pady=5,sticky=E)#side="left")
        stock_date_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")#,relief=FLAT)
        stock_date_ent.grid(row=9,column=1,padx=5,pady=5,sticky=E)
        lbl_14=Label(supplier_frame,text='Last Sold:')#,width=50,height=25,background="yellow")
        lbl_14.grid(row=10,column=0,padx=5,pady=5,sticky=E)#side="left")
        last_sale_ent=Entry(supplier_frame,relief=FLAT,fg="blue",bg="systembuttonface")#,relief=FLAT)
        last_sale_ent.grid(row=10,column=1,padx=5,pady=5,sticky=E)
        

        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,    
        #,,,Create the product Tree view,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,Create Treeview Scroll for products,,,
        tree_scroll_y=Scrollbar(product_tree_frame)
        tree_scroll_y.pack(side=RIGHT,fill=Y)
        product_tree=ttk.Treeview(product_tree_frame,height=3,yscrollcommand=tree_scroll_y.set,selectmode="browse")
        product_tree.pack()
        #,,,Configure the product scrollbar,,,,,,,,,,,,,
        tree_scroll_y.config(command=product_tree.yview)
        #tree_scroll2.config(command=product_tree.xview)
        #,,,Define the product treeview Columns,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        product_tree["columns"]=("Product","Supplier","Category","ID","Cost Price","Retail Price","In Qt","Stock Level","Min Level","Max Level","Barcode","Purchase Date",)
        #displaycolumns=("Product","Supplier","Category","ID","Cost Price","Retail Price","In Qt","Stock Level","Barcode","Purchase Date",)#show only columns you need to see.
        
        #,,,Format the product treeview Columns,,,
        product_tree.column("#0",width=0,stretch=NO)#,minwidth=25
        product_tree.column("Product",anchor=W,width=130,minwidth=25)
        product_tree.column("Supplier",anchor=W,width=130)
        product_tree.column("Category",anchor=W,width=90)
        product_tree.column("ID",anchor=CENTER,width=50,minwidth=25)
        product_tree.column("Cost Price",anchor=E,width=50)
        product_tree.column("Retail Price",anchor=E,width=50)
        product_tree.column("In Qt",anchor=CENTER,width=50)
        product_tree.column("Stock Level",anchor=CENTER,width=50)
        product_tree.column("Min Level",anchor=CENTER,width=50)
        product_tree.column("Max Level",anchor=CENTER,width=50)
        product_tree.column("Barcode",anchor=W,width=80)
        product_tree.column("Purchase Date",anchor=W,width=80)
        
        #,,,Create Headings for product treeview,,,,,,,,,,,,,,,,,,,,,
        product_tree.heading("#0",text="",anchor=W)
        product_tree.heading("Product",text="Product",anchor=W)
        product_tree.heading("Supplier",text="Supplier",anchor=CENTER)
        product_tree.heading("Category",text="Category",anchor=CENTER)
        product_tree.heading("ID",text="ID",anchor=CENTER)
        product_tree.heading("Cost Price",text="Cost",anchor=CENTER)
        product_tree.heading("Retail Price",text="Retail",anchor=CENTER)
        product_tree.heading("In Qt",text="In Qt",anchor=CENTER)
        product_tree.heading("Stock Level",text="Qt",anchor=CENTER)
        product_tree.heading("Min Level",text="Min",anchor=CENTER)
        product_tree.heading("Max Level",text="Max",anchor=CENTER)
        product_tree.heading("Barcode",text="Barcode",anchor=W)
        product_tree.heading("Purchase Date",text="Purchase Date",anchor=W)
            #product_tree.insert(parent='',index='end',iid=0,text='',values=('Eye drops','Med','Smiths','1','2.99','9.99','24','5','15','30','BF01234'))
            #add child to parent,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
            #product_tree.insert(parent='0',index='end',iid=4,text='Child',values=('Test child','somthing','Smiths','1.2','11.77','19.99','24','5','15','20','child001'))
            #product_tree.move('4', '0', '0')
        #,,, create striped rows in treeview,,,
        product_tree.tag_configure("oddrow",background="white")
        product_tree.tag_configure("evenrow",background="lightblue")
        query_product_database()
        
        #,,,Clear Product view entry boxes and inventory boxes,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def clear_product_view():
            #product_tree.forget
            #,,,Clear entry boxes.
            p1_ent.delete(0,END)
            p2_ent.delete(0,END)
            p3_ent.delete(0,END)
            p4_ent.delete(0,END)
            p5_ent.delete(0,END)
            p6_ent.delete(0,END)
            p7_ent.delete(0,END)
            p8_ent.delete(0,END)
            #min_ent.delete(0,END)
            #max_ent.delete(0,END)
            #stock_level_ent.delete(0,END)
            p9_ent.delete(0,END)
            stock_date_ent.delete(0,END)
            query_product_database()
        
        
        #,,, Select a Record in tree view,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def select_product_record(e):
            #code to ignore header being clicked.
            region_clicked=product_tree.identify_region(e.x, e.y)
            if region_clicked not in ('product_tree','cell'):
                return
                #print("Clicked Header")
            else:
                #button_del_product["state"]="normal"
                #prod_save_button["state"]="disabled"#,,disables save button.
                #,,,Clear entry boxes to enter tree data.
                p1_ent.delete(0,END)
                p2_ent.delete(0,END)
                p3_ent.delete(0,END)
                p4_ent.delete(0,END)
                p5_ent.delete(0,END)
                p6_ent.delete(0,END)
                p7_ent.delete(0,END)
                #min_ent.delete(0,END)
                #max_ent.delete(0,END)
                #stock_level_ent.delete(0,END)
                p8_ent.delete(0,END)
                p9_ent.delete(0,END)
                stock_date_ent.delete(0,END)
                #,,,Grab entry in treeview,,,,,,,,,,,,,,,,,,
                selected=product_tree.focus()
                #,,,Grab record values,,,,,,,,,,,,,,,,,,,,,,
                values=product_tree.item(selected,"values")
                #,,,output data from tree view to entry boxes.
                p1_ent.insert(0,values[0])
                p2_ent.insert(0,values[1])
                p3_ent.insert(0,values[2])
                p4_ent.insert(0,values[3])
                p5_ent.insert(0,format(float(values[4]),'.2f'))
                p6_ent.insert(0,format(float(values[5]),'.2f'))
                p7_ent.insert(0,values[6])
                p8_ent.insert(0,values[7])
                #min_ent.insert(0,values[8])
                #max_ent.insert(0,values[9])
                #stock_level_ent.insert(0,values[7])
                p9_ent.insert(0,values[10])
                stock_date_ent.insert(0,values[11])
        
        #,,, Bind the tree when clicked on.
        product_tree.bind("<ButtonRelease-1>",select_product_record)
        query_product_database()
        
        #,,,Product search popup window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        global lookup_product_records
        def lookup_product_records():
            global search_entry_product, search
            search=Toplevel(self)
            search.title("Search Product Records")
            search.geometry("400x200",)
            search.resizable(False,False)
            #,,,create lable frame
            search_frame=LabelFrame(search,text="Barcode")
            search_frame.pack(padx=10,pady=10)
            #,,,add entry box
            search_entry_product=Entry(search_frame,font=("Helvetica",18))
            search_entry_product.pack(padx=20,pady=20,)
            #,,,focus cursor on start.
            search_entry_product.focus()
            #,,,add a button
            search_button1=Button(search,text="Search Records",command=search_products)
            search_button1.pack(padx=20,pady=20)
            clear_product_view()
        
        
        def search_products():
            get_product=search_entry_product.get()
            #,,,close the search box
            search.destroy()
            #,,,clear the treeview
            for record in product_tree.get_children():
                 product_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            c.execute("SELECT rowid, * FROM products_two WHERE barcode like ? OR oid like ?", ((get_product,get_product,)))
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0
            
            for record in records:
                if count % 2==0:
                    product_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12]), tags=("evenrow",))
                else:
                    product_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12]), tags=("oddrow",))
                #,,,increment counter,,,,
                count +=1 
            if not records:
                query_product_database()
                #petinfo.set("NO results found for that search, check for spelling errors or try another name.")
                records=talk_search_pet()  
            #,,,commit the changes
            conn.commit()
                #,,,close the connection
            conn.close() 
            
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,    
        #,,,Home Tab Main Side (Home Page Frames), TAB 1,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,list of greetings on home page on start up.
        frase_list=["O.M.G!!,,, I can't believe\n",
                    "Shit,,,, I can't believe\n",
                    "Unbelievable,, I can't believe\n"]
        fraze=random.choice(frase_list)#https://youtu.be/XKHEtdqhLK8
        today= date.today()
        f_today=today.strftime('%A \n%B %d \n%Y ')
        day_name=today.strftime('%A')
        #print(day_name)
        if day_name =="Friday":
            today_label=Label(tab_0,text=f"Shake your tits,,,\nIt's \n{f_today}",font="helvetica, 70",bg="white",fg="orange",justify=LEFT)
            today_label.grid(row=0,column=0,padx=80,pady=20,sticky=W)#pack(pady=20)
        else:
            today_label=Label(tab_0,text=f"{fraze}It's \n{f_today}",font="helvetica, 70",bg="white",fg="orange",justify=LEFT)
            today_label.grid(row=0,column=0,padx=80,pady=20,sticky=W)#pack(pady=20)
        
        
        def count_home_pets():#https://youtu.be/gOLYhfCQXL4
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            cli=c.execute("SELECT * FROM customer_pets ")
            count=cli.fetchall()
            for x in count:
                booking_ref=(x[0])
            
            conn.commit()
            conn.close()
            count=str(len(count))
            pet_num_home.set("Registered Pet's\n\n\n""Total Pets "+ "(" + count + ")")
        
        
        def count_home_client():
            #https://youtu.be/gOLYhfCQXL4
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            cli=c.execute("SELECT * FROM clients ")
            count=cli.fetchall()
            for x in count:
                booking_ref=(x[0])
            
            conn.commit()
            conn.close()
            count=str(len(count))
            client_num_home.set("Registered Client's "+ "\n\n\n" +"Total Client's "+"(" + count + ")")
        
        
        def checkout_home_num_get():
            #https://youtu.be/jngHjpBiiKo
            #https://youtu.be/tlmOm35lsps
            dt=date.today()
            tomorrows_date=dt + timedelta(days=1)
            today_date=datetime.strftime(dt,'%d/%m/%Y')#%x=todays date#change the format to (2022/07/13 19-29-37) string format.
            tomorrows_date=datetime.strftime(tomorrows_date,'%d/%m/%Y')#%x=todays date#change the format to (2022/07/13 19-29-37) string format.
            day_3=dt + timedelta(days=2)
            day_4=dt + timedelta(days=3)
            day_5=dt + timedelta(days=4)
            day_6=dt + timedelta(days=5)
            day_7=dt + timedelta(days=6)
            day_3=datetime.strftime(day_3,'%d/%m/%Y')#string format.
            day_4=datetime.strftime(day_4,'%d/%m/%Y')#string format.
            day_5=datetime.strftime(day_5,'%d/%m/%Y')#string format.
            day_6=datetime.strftime(day_6,'%d/%m/%Y')#string format.
            day_7=datetime.strftime(day_7,'%d/%m/%Y')#string format.
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            #,,,find bookings between ???..(you can use calendar pick box for dates).
            arrival=c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(today_date,))# order by in_date desc")#desc
            
            returndates=arrival.fetchall()
            for x in returndates:
                booking_ref=(x[0])
                first_name=(x[1])
                last_name=(x[2])
                check_in=(x[8])
                check_out=(x[9])
                total_days=(x[10])
                #print(f"{booking_ref} {first_name} {last_name}, check in date: {check_in} check out date: {check_out} for {total_days} days.") 
            
            c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(tomorrows_date,))# order by in_date desc")#desc
            checkoutdate=c.fetchall()
            for x in checkoutdate:
                booking_ref=(x[0])
                first_name=(x[1])
                last_name=(x[2])
                check_in=(x[8])
                check_out=(x[9])
                total_days=(x[10])
                #print(f"{booking_ref} {first_name} {last_name}, check in date: {check_in} check out date: {check_out} for {total_days} days.")
            checkout_tomorrow=str(len(checkoutdate))
            d3=c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(day_3,))#order by in_date desc")#desc
            next_day_3=d3.fetchall()
            d4=c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(day_4,))#order by in_date desc")#desc
            next_day_4=d4.fetchall()
            d5=c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(day_5,))#order by in_date desc")#desc
            next_day_5=d5.fetchall()
            d6=c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(day_6,))#order by in_date desc")#desc
            next_day_6=d6.fetchall()
            d7=c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(day_7,))#order by in_date desc")#desc
            next_day_7=d7.fetchall()
            day_3=int(len(next_day_3))
            day_4=int(len(next_day_4))
            day_5=int(len(next_day_5))
            day_6=int(len(next_day_6))
            day_7=int(len(next_day_7))
            checkout_today=str(len(returndates))
            out_today=str(len(checkoutdate))
            total_week_arrive=str(int(checkout_today)+int(checkout_tomorrow)+day_3+day_4+day_5+day_6+day_7)#total for next 2 days
            checkout_num_home.set("Expected Checkout's\n\nToday           " + "(" + checkout_today +")"+ "\nTomorrow      " + "(" + out_today +")"+ "\nNext 7 day's  " + "(" + total_week_arrive +")")
            conn.commit()
            conn.close()
        
        
        def employ_home_num_get():
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            #,,,get row from Database,,,,,,,
            c.execute("SELECT rowid, * FROM employees WHERE date_left = 'No' ")
            count=c.fetchall()
            for x in count:
                booking_ref=(x[11])
            
            c.execute("SELECT * FROM employees WHERE position = 'Volenteer' AND date_left = 'No'")
            title=c.fetchall()
            for x in title:
                booking_ref=(x[11])
            
            conn.commit()
            conn.close()
            count=str(len(count))
            title=str(len(title))
            employee_num_home.set("Active Employee's"+"\n\nTotal Staff   " +"(" + count + ")"+"\nVolenteer's  "+"(" + title + ")")
        
        
        def arrival_home_num_get():
            global total_week_arrive
            #https://youtu.be/jngHjpBiiKo
            #https://youtu.be/tlmOm35lsps
            dt=date.today()
            todays_date=datetime.strftime(dt,'%d/%m/%Y')#string format.
            tomorrows_date=dt + timedelta(days=1)
            tomorrows_date=datetime.strftime(tomorrows_date,'%d/%m/%Y')#string format.
            day_3=dt + timedelta(days=2)
            day_4=dt + timedelta(days=3)
            day_5=dt + timedelta(days=4)
            day_6=dt + timedelta(days=5)
            day_7=dt + timedelta(days=6)
            day_3=datetime.strftime(day_3,'%d/%m/%Y')#string format.
            day_4=datetime.strftime(day_4,'%d/%m/%Y')#string format.
            day_5=datetime.strftime(day_5,'%d/%m/%Y')#string format.
            day_6=datetime.strftime(day_6,'%d/%m/%Y')#string format.
            day_7=datetime.strftime(day_7,'%d/%m/%Y')#string format.
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            
            #,,,find bookings between ???..(you can use calendar pick box for dates).
            #c.execute("SELECT DISTINCT (enclosure) FROM kennel_bookings WHERE in_date == ?",(todays_date,))#order by in_date desc")#desc
            #kennels=c.fetchall()
            #ken_total=int(len(kennels))#gets the number of kennels in use(but need between dates?????)
            #print(str(ken_total)+" Kennels in use today")
            #for x in kennels:
                #print("kennel "+ (str(x[0]))+ " is occupied today")#shows kennels(numbers)
            
            arrival=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(todays_date,))
            returndates=arrival.fetchall()
            for x in returndates:
                booking_ref=(x[0])
                first_name=(x[1])
                last_name=(x[2])
                check_in=(x[8])
                check_out=(x[9])
                total_days=(x[10])
                #print(f"{booking_ref} {first_name} {last_name}, check in date: {check_in} check out date: {check_out} for {total_days} days.") 
            arrive_today=str(len(returndates))
            c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(tomorrows_date,))# order by in_date desc")#desc
            checkoutdate=c.fetchall()
            for x in checkoutdate:
                booking_ref=(x[0])
                first_name=(x[1])
                last_name=(x[2])
                check_in=(x[8])
                check_out=(x[9])
                total_days=(x[10])#use this total to mark kennel filled days.
            arrive_tomorrow=str(len(checkoutdate))
            d3=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_3,))#order by in_date desc")#desc
            next_day_3=d3.fetchall()
            d4=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_4,))#order by in_date desc")#desc
            next_day_4=d4.fetchall()
            d5=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_5,))#order by in_date desc")#desc
            next_day_5=d5.fetchall()
            d6=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_6,))#order by in_date desc")#desc
            next_day_6=d6.fetchall()
            d7=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_7,))#order by in_date desc")#desc
            next_day_7=d7.fetchall()
            day_3=int(len(next_day_3))
            day_4=int(len(next_day_4))
            day_5=int(len(next_day_5))
            day_6=int(len(next_day_6))
            day_7=int(len(next_day_7))
            total_week_arrive=str(int(arrive_today)+int(arrive_tomorrow)+day_3+day_4)#+day_5+day_6+day_7)#total for next 7 days
            arrival_num_home.set("Expected Arrivals\n\n  Today           " + "(" + arrive_today +")"+ "\n  Tomorrow      " + "(" + arrive_tomorrow +")"+ "\n  Next 7 day's  " + "(" + total_week_arrive +")")
            conn.commit()
            conn.close()
        
        
        def occupancy_home_num_get():
            dt=date.today()
            today_date=datetime.strftime(dt,'%d/%m/%Y')#%x=todays date#change the format to (2022/07/13 19-29-37) string format.
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            #,,,get row from Database (should have only 1 row),,,,,,,
            kennel_data=c.execute("SELECT DISTINCT (enclosure) FROM kennel_bookings WHERE in_date = ? ",(today_date,))
            kennels=kennel_data.fetchall()
            ken_total=int(len(kennels))#gets the number of kennels in use(but need between dates?????)
            inuse=(str(ken_total))
        
            c.execute("SELECT * FROM general_settings")
            records=c.fetchall()
            for record in records:
                occupancy_amount=(record[1])
                num_ken=int(occupancy_amount)
                used_ken=int(inuse)
                available_ken=(num_ken-used_ken)
                #print(kennels)    
                
            conn.commit()
            conn.close()  
            occupancy_num_home.set("Total "+pen_name +"'s          "+" ("+str(occupancy_amount)+")\n\n"+pen_name+"'s "+"available    ""("+str(available_ken)+")\n"+ pen_name+"'s in use         ""("+inuse+")")
        
        
        
        def pet_home_hover(e):
            count_home_pets()
            
        def pet_home_hover_leave(e):
            pet_num_home.set("Pet's ")
            
        def client_home_hover(e):
            count_home_client()
            
        def client_home_hover_leave(e):
            client_num_home.set("Client's ")    
            
        
        def employ_home_hover(e):
            employ_home_num_get()
            
        def employ_home_hover_leave(e):
            employee_num_home.set("Employee's ")
            
   
        def arrival_home_hover(e):
            arrival_home_num_get()
            
        def arrival_home_hover_leave(e):
            arrival_num_home.set("Arrival's ")
               
        def checkout_home_hover(e):
            checkout_home_num_get()
            
        def checkout_home_hover_leave(e):
            checkout_num_home.set("Checkout's ")
        
        def occupancy_home_hover(e):
            occupancy_home_num_get()
            
        def occupancy_home_hover_leave(e):
            occupancy_num_home.set("Kennel Occupancy ")
        
        
        
        #,,,Home screen widgets,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,    
        home_repots_frame1=Frame(tab_0,bg="white")  
        home_repots_frame1.place(y=240,x=1054)
        home_box1_reports=HoverButton(home_repots_frame1,textvariable=pet_num_home,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#975ac7",fg="white",height=5,width=18,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Manager only, Password Needed"))
        home_box1_reports.grid(row=0,column=0,padx=4,pady=3,ipady=5,ipadx=5)
        home_box1_reports.bind("<Enter>",pet_home_hover)
        home_box1_reports.bind("<Leave>",pet_home_hover_leave)
        
        home_repots_frame2=Frame(tab_0,bg="white")  
        home_repots_frame2.place(y=360,x=728) 
        home_box2_reports=HoverButton(home_repots_frame2,textvariable=arrival_num_home,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#d97f43",fg="white",height=5,width=18,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Manager only, Password Needed"))
        home_box2_reports.grid(row=1,column=0,padx=4,pady=3,ipady=5,ipadx=5)
        home_box2_reports.bind("<Enter>",arrival_home_hover)
        home_box2_reports.bind("<Leave>",arrival_home_hover_leave)
        home_box3_reports=HoverButton(home_repots_frame2,textvariable=checkout_num_home,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#eb757d",fg="white",height=5,width=30,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Manager only, Password Needed"))
        home_box3_reports.grid(row=1,column=1,padx=4,pady=3,ipady=5,ipadx=5)
        home_box3_reports.bind("<Enter>",checkout_home_hover)
        home_box3_reports.bind("<Leave>",checkout_home_hover_leave)
        
        home_repots_frame3=Frame(tab_0,bg="white")  
        home_repots_frame3.place(y=480,x=520)
        home_box4_reports=HoverButton(home_repots_frame3,textvariable=client_num_home,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#05aff7",fg="white",height=5,width=18,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Manager only, Password Needed"))
        home_box4_reports.grid(row=2,column=0,padx=4,pady=3,ipady=5,ipadx=5)
        home_box4_reports.bind("<Enter>",client_home_hover)
        home_box4_reports.bind("<Leave>",client_home_hover_leave)
        
        home_box5_reports=HoverButton(home_repots_frame3,textvariable=occupancy_num_home,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#a67289",fg="white",height=5,width=30,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Manager only, Password Needed"))
        home_box5_reports.grid(row=2,column=1,padx=4,pady=3,ipady=5,ipadx=5)
        home_box5_reports.bind("<Enter>",occupancy_home_hover)
        home_box5_reports.bind("<Leave>",occupancy_home_hover_leave)
        
        home_box6_reports=HoverButton(home_repots_frame3,textvariable=employee_num_home,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#2d8734",fg="white",height=5,width=18,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Manager only, Password Needed"))
        home_box6_reports.grid(row=2,column=2,padx=4,pady=3,ipady=5,ipadx=5)
        #name_lab=Label(tab_0,text=company_name,bg="white",font="helvetica, 40",fg="#e2f723",justify=CENTER)
        #name_lab.grid(row=1,column=0,sticky=S)
        home_box6_reports.bind("<Enter>",employ_home_hover)
        home_box6_reports.bind("<Leave>",employ_home_hover_leave)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,,,,,,,,,,,,,,,,,,,,,Inventory & Products Tab main screen,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        serviceinfo=StringVar()
        prod_lbl_frame=LabelFrame(box_frame3,text="Add new product details",font=20)#,bg="#e6f5f3")
        prod_lbl_frame.grid(row=1,column=0,padx=10,pady=0,ipadx=0,ipady=0,sticky=N)#fill="x",expand="yes"
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,product overview command buttons,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        comand_frame=LabelFrame(supplier_frame,text="")#,bg="#e6f5f3")
        comand_frame.grid(row=14,columnspan=2,column=0,padx=3,pady=3)
        
        clear_button=HoverButton(comand_frame,text='CLEAR OVERVIEW',cursor="hand2",activebackground = '#e2f723',fg="black",command=query_product_database)#lambda: popupmsg('Not supported just yet!'))
        clear_button.grid(row=0,column=0,padx=3,pady=2)#,ipadx=5,ipady=5)
        update_button=HoverButton(comand_frame,text='UPDATE DETAILS',cursor="hand2",activebackground = '#e2f723',fg="black",command=update_product_details)#lambda: popupmsg('Not supported just yet!'))
        update_button.grid(row=0,column=1,padx=3,pady=2)#,ipadx=5,ipady=5)
        
        
        button1_product=HoverButton(button_label_frame,text="",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT)
        button1_product.grid(row=0,column=0,padx=1,pady=1,ipady=5,ipadx=10)
        button_search_product=HoverButton(button_label_frame,text="SEARCH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lookup_product_records)#lambda:popupmsg("Not yet completed"))
        button_search_product.grid(row=0,column=1,padx=1,pady=1,ipady=5,ipadx=10)
        button_add_product=HoverButton(button_label_frame,text="ADD PRODUCT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=new_product_window)#lambda:popupmsg("Not yet completed"))
        button_add_product.grid(row=0,column=2,padx=1,pady=1,ipady=5,ipadx=10)
        button_del_product=HoverButton(button_label_frame,text="DELETE",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=remove_product)#lambda:popupmsg("Not yet completed"))
        button_del_product.grid(row=0,column=3,padx=1,pady=1,ipady=5,ipadx=10)
        button_update_product=HoverButton(button_label_frame,text="REFRESH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=query_product_database)#lambda:popupmsg("Not yet completed"))
        button_update_product.grid(row=0,column=4,padx=1,pady=1,ipady=5,ipadx=10)
        button_help_product=HoverButton(button_label_frame,text="HELP",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=help_prod_info)#lambda:popupmsg("Not yet completed"))
        button_help_product.grid(row=0,column=5,padx=1,pady=1,ipady=5,ipadx=10)
        button_exit_product=HoverButton(button_label_frame,text="EXIT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:notebook.tab(tab_1,state="hidden"))
        button_exit_product.grid(row=0,column=6,padx=1,pady=1,ipady=5,ipadx=10)
        
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,SERVICES TAB 2,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        
        #,,,refreshes the table,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def query_services_database():
            #,,,clear the treeview
            for record in services_tree.get_children():
                services_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute("SELECT rowid, * FROM service")
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            for record in records:
                if count % 2==0:
                    services_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[3],record[4]), tags=("evenrow",))
                else:
                    services_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[3],record[4]), tags=("oddrow",))
                #,,,increment counter adds 1 to each record entered (ID number),,,,
                count +=1   
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            
            
        #,,,Clear entry boxes.
        def clear_service_box():
            #services_tree.forget()#removes the whole tree view.
            service_save_button["state"]="normal"
            sev_name.delete(0,END)
            descrip.delete(0,END)
            rate.delete(0,END)
            type_combo.set('')
            serv_ref.delete(0,END)
            servicetitle.set(" Add New Service")
            serviceinfo.set("")
            query_services_database()
            
            
        def help_services():
            a=(f'')
            b=(f'')
            c=(f'Select a service to delete or update.\n\n')
            d=(f'The SEARCH function will help.\n\n')
            e=(f'update as needed then click UPDATE or DELETE\n\n')
            f=(f'WARNING... DELETE will permanently remove from database.\n')
            g=(f'(you will have the option to continue or cancel the deletion.)\n\n')
            h=(f'')
            ans=c+d+e+f+g
            serviceinfo.set(ans)
            
        
        def find_service(*args):
            lookup_service=search_entry_service.get()
            #,,,close the search box
            search.destroy()
            service_save_button["state"]="disabled"
            #servicetitle.set(" Make changes or update")
            #,,,clear the treeview
            for record in services_tree.get_children():
                 services_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            c.execute("SELECT rowid, * FROM service WHERE service_name like ? ",(lookup_service,))
            records=c.fetchall()
            global count
            count = 0 #,,,start of ID
            #,,,Add Record Entry Boxes,,,
            for record in records:
                if count % 2==0:
                    services_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[3],record[4]), tags=("evenrow",))
                else:
                    services_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[3],record[4]), tags=("oddrow",))
                #,,,increment counter,,,,
                count +=1 
            if not records:
                serviceinfo.set("NO results found for that search, check for spelling errors or try another name.")
                query_services_database()
                records=talk_search_pet()
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            
        
        def search_service():
            global search_entry_service, search
            serv_ref.delete(0,END)
            sev_name.delete(0,END)
            descrip.delete(0,END)
            rate.delete(0,END)
            type_combo.set('')
            search=Toplevel(self)
            search.title("Search Services Records")
            search.geometry("400x200",)
            search.resizable(False,False)
            #,,,create lable frame
            search_frame=LabelFrame(search,text="Enter Service Name:")
            search_frame.pack(padx=10,pady=10)
            #,,,add entry box
            search_entry_service=Entry(search_frame,font=("Helvetica",18))
            search_entry_service.bind("<Return>",find_service)#enter key binding(dont forget *args).
            search_entry_service.pack(padx=20,pady=20,)
            #,,,focus cursor on start.
            search_entry_service.focus()
            #,,,add a button
            search_button1=Button(search,text="Search Records",command=find_service)
            search_button1.pack(padx=20,pady=20)
            serviceinfo.set("")
        
    
        top_label_frame1=Frame(tab_2,bd=1)  
        top_label_frame1.pack(pady=0,padx=0,anchor=W,fill=X)
        heading=Label(top_label_frame1,text="SERVICES",font="bold 20",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=17)
        button_label_frame=Frame(tab_2,bg="light blue",bd=1)  
        button_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        box_frame4=Frame(tab_2)#,bg="green",bd=1)  
        box_frame4.pack(pady=0,padx=0,anchor=W,fill=X)
        box_frame5=Frame(tab_2)#,bg="pink",bd=1)  
        box_frame5.pack(pady=0,padx=0,anchor=W,fill=X)
        box_frame6=Frame(tab_2)#,bg="blue",bd=1)  
        box_frame6.pack(pady=0,padx=0,anchor=W,fill=X)
        
        servicetitle=StringVar()
        servicetitle.set(" Add New Service")
        service_option=Label(box_frame4,textvariable=servicetitle,font=20,fg="#347083")
        service_option.grid(row=0,column=0,columnspan=3,padx=20,pady=(20,0))
        
        details_frame=LabelFrame(box_frame4,text="",font=20,bd=0)#,bg="red")#,fg="#e6f5f3")
        details_frame.grid(row=1,column=0,columnspan=6,padx=20,pady=5,ipadx=5,ipady=5,sticky=W)
        
        services_tree_frame=LabelFrame(box_frame6,bd=0)#,fg="#e6f5f3")
        services_tree_frame.grid(row=0,column=0,padx=20)
        
        #,,,Create Treeview Scroll for products,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        tree_scroll_y=Scrollbar(services_tree_frame)
        tree_scroll_y.pack(side=RIGHT,fill=Y)
        #,,,Create the services Tree view,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        services_tree=ttk.Treeview(services_tree_frame,height=3,yscrollcommand=tree_scroll_y.set,selectmode="browse")
        services_tree.pack()
        #,,,Configure the services scrollbar,,,,,,,,,,,,,
        tree_scroll_y.config(command=services_tree.yview)
        
        #,,,Define the services treeview Columns,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        services_tree["columns"]=("id","service name","description","rate","department_id",)
        #displaycolumns=("service name","description","rate","department_id",)#show only columns you need to see.
        
        #,,,Format the services treeview Columns,,,
        services_tree.column("#0",width=0,stretch=NO)#,minwidth=25
        services_tree.column("id",anchor=W,width=50,minwidth=50)
        services_tree.column("service name",anchor=W,width=150,minwidth=75)
        services_tree.column("description",anchor=W,width=250,minwidth=75)
        services_tree.column("rate",anchor=CENTER,width=90,minwidth=75)
        services_tree.column("department_id",anchor=CENTER,width=100,minwidth=75)
        
        #,,,Create Headings for services treeview,,,,,,,,,,,,,,,,,,,,,
        services_tree.heading("#0",text="",anchor=W)
        services_tree.heading("id",text="ID",anchor=W)
        services_tree.heading("service name",text="Service",anchor=W)
        services_tree.heading("description",text="Discription",anchor=CENTER)
        services_tree.heading("rate",text="Price (£)",anchor=CENTER)
        services_tree.heading("department_id",text="Department",anchor=CENTER)
        
        #,,, create striped rows in services treeview,,,
        services_tree.tag_configure("oddrow",background="white")
        services_tree.tag_configure("evenrow",background="lightblue")
        query_services_database()
        
        #,,, Select a Record in service tree view,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def select_service_record(e):
            #code to ignore header being clicked.
            region_clicked=services_tree.identify_region(e.x, e.y)
            if region_clicked not in ('services_tree','cell'):
                return
            else:
                #button_del_product["state"]="normal"
                #,,,Clear entry boxes to enter tree data.
                service_save_button["state"]="normal"
                sev_name.delete(0,END)
                descrip.delete(0,END)
                rate.delete(0,END)
                type_combo.set('')
                serv_ref.delete(0,END)
                servicetitle.set(" Add New Service")
                serviceinfo.set("")
                #,,,Grab entry in treeview,,,,,,,,,,,,,,,,,,
                selected=services_tree.focus()
                #,,,Grab record values,,,,,,,,,,,,,,,,,,,,,,
                service_values=services_tree.item(selected,"values")
                #,,,output data from tree view to entry boxes.
                sev_name.insert(0,service_values[1])
                descrip.insert(0,service_values[2])
                rate.insert(0,format(float(service_values[3]),'.2f'))
                type_combo.insert(0,service_values[4])
                serv_ref.insert(0,service_values[0])
                service_save_button["state"]="disabled"
                servicetitle.set(" Update or Delete this Service")
            
        
        def double_click(e):
            region=services_tree.identify_region(e.x,e.y)
            if region not in ("cell"):#("tree","cell")
                return#Do NOTHING
            #print(region)
            #print("You double clicked")
            column=services_tree.identify_column(e.x)#identy column id.
            column_index = int(column[1:])-1
            #print(column)
            #print(column_index)
            selected_iid=services_tree.focus()#identy row id.
            #print(selected_iid)
            selected_values=services_tree.item(selected_iid)#identify all values in row.
            #print(selected_values)
            if column == "#0":
                selected_text = selected_values.get("text")
            else:
                selected_text= selected_values.get("values")[column_index]
            #print(selected_text)
            column_box = services_tree.bbox(selected_iid, column)#column will be (#0,#1 ect)
            #print(column_box)#this gives the cordinants ie(51, 43, 150, 25)(x, y,width,height)top left corner of cell clicked.
            entry_edit=ttk.Entry(services_tree,width=column_box[2])
            #record the column index and item iid
            entry_edit.editing_column_index=column_index
            entry_edit.editing_item_iid = selected_iid
            entry_edit.insert(0,selected_text)
            entry_edit.select_range(0, tk.END)
            entry_edit.focus()
            
            entry_edit.bind("<FocusOut>",on_focus_out)
            entry_edit.bind("<Return>",on_enter_pressed)
            
            entry_edit.place(x=column_box[0],
                             y=column_box[1],
                             width=column_box[2],
                             height=column_box[3])
            
        
            
        def on_enter_pressed(event):
            new_text = event.widget.get()
            #such as i002
            selected_iid = event.widget.editing_column_index
            #such as -1(tree column)
            column_index = event.widget.editing_item_iid
            if column_index == -1:
                services_tree.item(selected_iid, text= new_text)
            else:
                current_values = services_tree.item(selected_iid).get("values")
                current_values[column_index] = new_text
                services_tree.item(selected_iid, values=current_values)
                
        
        def on_focus_out(event):
            event.widget.destroy()
            
                
            event.widget.destroy()
            
        #https://youtu.be/n5gItcGgIkk  #This is a good video for Treeview cells
        #,,, Bind the tree when clicked on.
        services_tree.bind("<ButtonRelease-1>",select_service_record)
        services_tree.bind("<Double-1>",double_click)
        
            
        

        label_1=Label(details_frame,text='Service name:')#,bg="#e6f5f3")
        label_1.grid(row=0,column=0,padx=5,pady=5,sticky=W)
        sev_name=Entry(details_frame,width=28)
        sev_name.grid(row=0,column=1,sticky=W)
        label_2=Label(details_frame,text='Description:')#,bg="#e6f5f3")
        label_2.grid(row=1,column=0,padx=5,pady=5,sticky=W)
        descrip=Entry(details_frame,width=28)
        descrip.grid(row=1,column=1,sticky=W)
        label_3=Label(details_frame,text='Retail Rate: £')#,bg="#e6f5f3")
        label_3.grid(row=2,column=0,padx=5,pady=5,sticky=W)
        rate=Entry(details_frame,width=8)
        rate.grid(row=2,column=1,sticky=W)
        
        
        info_frame=Frame(box_frame4)  
        info_frame.grid(row=0,column=1)
        #serviceinfo.set("There is a hidden Label here!.\nClick HELP to view.")
        lb1=Label(box_frame4,textvariable=serviceinfo,font="Verdana, 12",fg="#347083",justify=LEFT)
        lb1.grid(row=1,column=7,padx=20,pady=0,sticky=W)
        
        '''
        #,,check Button, on services window,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def check_button_action():
            print(check_button_var.get())
            #default button states are 0 and 1.
            if check_button_var.get()==1:
                check_button.configure(fg="red")
            else:
                check_button.configure(fg="black")        
        
        check_button_var=tk.IntVar()
        check_button=tk.Checkbutton(details_frame,text="VAT inclusive",bd=1,variable=check_button_var,command=check_button_action)
        check_button.grid(row=4,column=0,padx=5,pady=5,sticky=W)'''
        
        
        def talk_service():
            global engine
            if serv_ref.get()=="" :
                engine=pyttsx3.init()
                engine.say(f'''you need to select a service to delete,
                           you can also choose from a search result using the service name''')
            else:
                engine=pyttsx3.init()
                engine.say(f''' CAUTION,, you are deleteing service number{serv_ref.get()}, {sev_name.get()}, is this correct ?.''')
            engine.runAndWait()
        
        
         #,,Delete a Product from Database,,,,,,,,,,,,,,,,,,,,,
        def remove_service():
            #talk_service()
            #app.bell()
            if  serv_ref.get()=="":
                talk_service()
                messagebox.showerror("ERROR","Select a service record to delete")
            else:
                talk_service()
                response = messagebox.askyesno("WARNING!!", "This will DELETE the Selected service\nfrom the Database\nAre You Sure?!.")
                #,,,add logic for message box.
                if response==1:
                    #,,,connect to database.    
                    conn=sqlite3.connect("tree_crm.db")
                    #,,,create a cursor.
                    c=conn.cursor()
                    #,,,delete from table
                    c.executemany("DELETE FROM service WHERE oid = ? ", serv_ref.get())
                    #,,,commit the changes
                    conn.commit()
                    #,,,close the connection
                    conn.close()
                    #,,,clear entries from boxes if filled.
                    clear_service_box()
                    messagebox.showinfo("SELECTION DELETED", "Records have been Deleted")
                    query_services_database()
        
        
        def save_service():
            if sev_name.get()=="":
                popupmsg("Looks like you forgot somthing!\n\nService name needed.")
            elif descrip.get()=="":
                popupmsg("Looks like you forgot somthing!\n\nService discription needed\n(Please keep it short!).")
            elif rate.get()=="":
                popupmsg("More information needed\nYou forgot the Service rate.")
            elif type_combo.get()=="":
                popupmsg("Department needed!\n\nNOT in the list?\nadd a new Department in then SYSTEM SETTINGS.")
            else:
                #,,, Connect to database.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record.
                c.execute("INSERT OR IGNORE INTO service VALUES(:service_name,:description,:rate,:department)",
                        {
                        "service_name":sev_name.get().capitalize(),
                        "description":descrip.get().capitalize(),
                        "rate":rate.get(),
                        "department":type_combo.get(),
                        #"vat":check_button_var.get(),
                        })
                
                #,,,commit the changes
                conn.commit()
                #,,, close the connection
                conn.close()
                query_services_database()
                
                #,,,Clear entry boxes.
                sev_name.delete(0,END)
                descrip.delete(0,END)
                rate.delete(0,END)
                type_combo.set('')
                #check_button.configure(fg="black")
                popupmsg("New service has been saved.")
            
        
        def service_show():
            serviceinfo.set("")
            #,,, search database or connect to one.
            #employ_info.set("")
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row from Database,,,,,,,
            c.execute("SELECT rowid, * FROM service")
            records=c.fetchall()
            global count
            count = 0 #,,,start of ID
            for record in records:
                count=count +1
                #print(records)
            num = 7
            for record in records:
                #test=float(record[3])
                serv_price=format(float(record[3]),'.2f')
                intro=Label(box_frame5,text=f'Number of Services offered {count}\n',font="Verdana, 12",fg="#347083")#,justify="left")
                intro.grid(row=6,column=0,columnspan=6,padx=5,pady=5,sticky=W)
                id=Label(box_frame5,text=f'Service ref: {record[0]}',font="Verdana, 12",fg="#347083")#,justify="left")
                id.grid(row=num,column=0,columnspan=2,padx=5,sticky=W)
                servname1=Label(box_frame5,text=f' {record[1]}',font="Verdana, 12",fg="#347083")#,justify="left")
                servname1.grid(row=num,column=2,padx=5,sticky=W)
                name2=Label(box_frame5,text=record[2],font="Verdana, 12",fg="#347083")#,justify="left")
                name2.grid(row=num,column=3,padx=10,sticky=W)
                phone=Label(box_frame5,text=f'£ {serv_price}',font="Verdana, 12",fg="#347083")#,justify="left")
                phone.grid(row=num,column=4,padx=2,sticky=E)
                position=Label(box_frame5,text=record[4],font="Verdana, 12",fg="#347083")#,justify="left")
                position.grid(row=num,column=5,padx=10,sticky=E)
                
                num=num+1
        
            
            conn.commit()
            conn.close()
            
            sev_name.delete(0,END)
            descrip.delete(0,END)
            rate.delete(0,END)
            type_combo.set('')
            serv_ref.delete(0,END)
            #serviceinfo.set()
        
        
        def clear_service_list():
            pass
        
        
        #,,,update service Details,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def update_service_details():#,,updates Business Details boxes.
            serviceinfo.set("")
            if serv_ref.get()=="":
                messagebox.showerror("Missing Data!","View services then\nSearch for A service by name\nto populate fields.")
            elif sev_name.get()=="":
                messagebox.showerror("Missing Data!","Service name needed.")
            elif descrip.get()=="":
                messagebox.showerror("Missing Data!","Description needed.")
            elif rate.get()=="":
                messagebox.showerror("Missing Data!","rate needed.")
            else:
                #,,, Connect to database.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record.(UPDATE) sercice 
                c.execute(""" UPDATE service SET 
                          service_name=:service_name, 
                          description=:description, 
                          rate=:rate, 
                          department_id=:department_id
                    
                          WHERE oid = :oid """,
                            {
                            "service_name":sev_name.get().capitalize(),
                            "description":descrip.get().capitalize(),
                            "rate":rate.get(),
                            "department_id":type_combo.get(),
                            "oid":serv_ref.get(),
                            })
                sev_name.delete(0,END)
                descrip.delete(0,END)
                rate.delete(0,END)
                type_combo.set('')
                serv_ref.delete(0,END)
                #,,,commit the changes
                conn.commit()
                #,,, close the connection
                conn.close()
                query_services_database()
                messagebox.showinfo("UPDATED", "Service record has been\nSuccessfully Updated")
                service_save_button["state"]="normal"
                
                
        label_1=Label(details_frame,text='Department:')#,bg="#e6f5f3")
        label_1.grid(row=3,column=0,padx=5,pady=5,sticky=W)
        type_combo=ttk.Combobox(details_frame,width=25)#,state="readonly")
        type_combo['values']=department_list
        type_combo.current()
        type_combo.grid(row=3,column=1,padx=5,pady=5,sticky=W)
        label_ref=Label(details_frame,text='Service ref:')#,bg="#e6f5f3")
        label_ref.grid(row=4,column=0,padx=5,pady=5,sticky=W)
        serv_ref=Entry(details_frame,width=20,relief=FLAT,bg="systembuttonface")
        serv_ref.grid(row=4,column=1,sticky=W)
        
        #,,,Service details Buttons,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        button_frame=LabelFrame(details_frame,text="",bd=0)#,bg="#e6f5f3")
        button_frame.grid(row=5,column=0,columnspan=2,padx=5,pady=5)#fill="x",expand="yes"
        global clear_button_Service
        clear_button_Service=HoverButton(button_frame,text='CLEAR',cursor="hand2",activebackground= '#e2f723',fg="black",command=clear_service_box)#clear_service_box)#lambda: popupmsg('Not supported just yet!'))
        clear_button_Service.grid(row=0,column=0,padx=(70,3),pady=3)
        global service_save_button
        service_save_button=HoverButton(button_frame,text='SAVE',cursor="hand2",activebackground= '#e2f723',fg="black",command=save_service)#lambda: popupmsg('Not supported just yet!'))
        service_save_button.grid(row=0,column=1,padx=3,pady=3)
        service_update_button=HoverButton(button_frame,text='UPDATE',cursor="hand2",activebackground= '#e2f723',fg="black",command=update_service_details)#lambda: popupmsg('Not supported just yet!'))
        service_update_button.grid(row=0,column=2,padx=3,pady=3)
        '''
        #,,,Notes Frame Services Tab 2,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def delete_all():
            notes_list.delete(1.0,END)
        
        #,,open file function
        def open_file():
            #,,delete previus text
            notes_list.delete("1.0",END)
            #,,Grab file name
            text_file=filedialog.askopenfilename(initialdir="C:/gui/",title="Langton & Roberts, Open a File",filetypes=(("Text Files","*.txt"),("HTML Files","*.htlm"),("Python Files","*.py"),("All Files","*.*")))
            #,,check to see if there is a file name
            if text_file:
                #,,make filename global so we can access it later.
                global open_status_name
                open_status_name=text_file
            #,,update status bars.
            name=text_file
            status_bar.config(text=f'{name}        ')
            #name=name.replace("C:/gui/","")
            #root.title(f"{name}-TextPad!")
            #,,open the file
            text_file=open(text_file,'r')
            stuff=text_file.read()
            #,,add file to text box.
            notes_list.insert(END,stuff)
            #,,close the opened file.
            text_file.close()
                
            
        def save_notes():
            text_file=filedialog.asksaveasfilename(initialdir="C:/gui/",
                                                   title="Save Notes, Langton & Roberts.",
                                                   filetypes=(("Text Files", "*.txt"),))
            #,,saves the file
            text_file=open(text_file,'w')
            text_file.write(notes_list.get(1.0,END))
            text_file.close()
            notes_list.delete(1.0,END)
            #status_bar.config(text=f'Saved:{open_status_name}')
            status_bar.config(text=Langton_and_Roberts)#f'{name}        ')
            popupmsg("Notes have been\nsaved to notes folder")
            
           
        notes_frame=LabelFrame(tab_2,text="Note Pad",font=20)#,bg="#e6f5f3")
        notes_frame.grid(row=1,column=0,columnspan=1,padx=30,pady=0,sticky=NW)
        notes_list=Text(notes_frame,width=36,height=4,font="helvetica 12")
        notes_list.grid(row=0,column=0,columnspan=4,padx=10,pady=10)
        button_frame=LabelFrame(notes_frame,text="Button comands")
        button_frame.grid(row=1,column=0,padx=5,pady=5)
        button_frame=LabelFrame(notes_frame,text="Note Pad comands")
        button_frame.grid(row=1,column=3,padx=5,pady=5,sticky=E)
       
        #btn1=Button(button_frame,text="Delete All",command=delete_all)#,,bg="red"
        #btn1.grid(row=0,column=0,padx=5,pady=5)
        btn2=Button(button_frame,text="Clear",cursor="hand2",command=delete_all)#lambda: popupmsg('Not supported just yet!'))
        btn2.grid(row=0,column=1,padx=5,pady=5)
        btn3=Button(button_frame,text="Save",cursor="hand2",command=save_notes)#lambda: popupmsg('Not supported just yet!'))
        btn3.grid(row=0,column=2,padx=5,pady=5)
        btn4=Button(button_frame,text="open",cursor="hand2",command=open_file)#lambda: popupmsg('Not supported just yet!'))
        btn4.grid(row=0,column=3,padx=5,pady=5)
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        '''
        button_1_Service=HoverButton(button_label_frame,text="",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT)
        button_1_Service.grid(row=0,column=0,padx=1,pady=1,ipady=5,ipadx=10)
        #button_refresh_Service=HoverButton(button_label_frame,text="REFRESH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=query_services_database)#lambda:popupmsg("Not yet completed"))
        #button_refresh_Service.grid(row=0,column=1,padx=1,pady=1,ipady=5,ipadx=10)
        #button_update_Service=HoverButton(button_label_frame,text="UPDATE",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=update_service_details)#lambda:popupmsg("Not yet completed"))
        #button_update_Service.grid(row=0,column=2,padx=1,pady=1,ipady=5,ipadx=10)
        button_search_Service=HoverButton(button_label_frame,text="SEARCH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=search_service)
        button_search_Service.grid(row=0,column=3,padx=1,pady=1,ipady=5,ipadx=10)
        button_del_Service=HoverButton(button_label_frame,text="DELETE",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=remove_service)#lambda:popupmsg("Not yet completed"))
        button_del_Service.grid(row=0,column=4,padx=1,pady=1,ipady=5,ipadx=10)
        button_help_Service=HoverButton(button_label_frame,text="HELP!",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=help_services)
        button_help_Service.grid(row=0,column=5,padx=1,pady=1,ipady=5,ipadx=10)
        button_exit_Service=HoverButton(button_label_frame,text="EXIT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:notebook.tab(tab_2,state="hidden"))
        button_exit_Service.grid(row=0,column=6,padx=1,pady=1,ipady=5,ipadx=10)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,TAB 3 Pet Information,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        top_label_frame1=Frame(tab_3,bd=1)  
        top_label_frame1.pack(pady=0,padx=0,anchor=W,fill=X)
        heading=Label(top_label_frame1,text="PET INFOMATION",font="bold 20",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=17)
        button_label_frame=Frame(tab_3,bg="light blue",bd=1)  
        button_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        box_frame3=Frame(tab_3)  
        box_frame3.pack(pady=10,padx=30,anchor=W,fill=X)
        petinfo=StringVar()
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create Pet table in database,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create Pet table.,,,,,,,,,,,,,,,,,,,,,,,
        c.execute("""CREATE TABLE if not exists customer_pets(
            pet_owner text,
            pet_name text,
            pet_age integer,
            status text,
            pet_type text,
            pet_dob integar,
            breed text,
            size text,
            micro_chip text,
            chip_number integer,
            gender text,
            weight real,
            colour text,
            temper text,
            reg_date integer DEFAULT CURRENT_TIMESTAMP,
            health text,
            attributes text,
            photo BLOB)
            """)
        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        
        
        
        def talk_pet():
            engine=pyttsx3.init()
            engine.say(f''' WARNING,, you have chosen to delete {secret_pet_name.get()}, with system ID {secret_id.get()} from the database.,
                        if this is correct, select yes,
                        alternatively, if this was an error in any way, select no..''')
            engine.runAndWait()
        
        
        def talk():
            global engine
            if secret_id.get()=="" :
                engine=pyttsx3.init()
                engine.say(f'''you need to select a pet first before you can delete a record,
                           you can search using the pet name or the pet ID''')
                #engine.say(my_entry.get())
                #engine.say(my_entry2.get())
                #engine.say(my_entry3.get())
            else:
                engine.say(f''' WARNING,, you have chosen to delete {secret_pet_name.get()}, with system ID {secret_id.get()} from the database.,
                        if this is correct, select yes,
                        alternatively, if this was an error in any way, select no..''')
            engine.runAndWait()
        
        
        #,,Delete a Product from Database,,,,,,,,,,,,,,,,,,,,,
        def remove_pet_info():
            talk_pet()
            app.bell()
            if secret_id.get()=="" :
                messagebox.showerror("ERROR","Select a record to delete")
            else:
                response = messagebox.askyesno("WARNING!!", "This will DELETE Selected\nfrom the Database\n\nAre You Sure?!.")
            #,,,add logic for message box.
            if response==1:
                #,,,desegnate selections
                x=pet_tree.selection()
                #,,,create list of ID's.
                ids_to_delete=[]
                #,,,add selections to ids_to_delete list.
                for record in x:
                    ids_to_delete.append(pet_tree.item(record,"values")[3])
                #,,,Delete from treeview.
                for record in x:
                    pet_tree.delete(record)
                #,,,create database.    
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,delete everything from table
                c.executemany("DELETE FROM customer_pets WHERE oid = ?",[(a,)for a in ids_to_delete])
                #,,,commit the changes
                conn.commit()
                #,,,close the connection
                conn.close()
                #,,,clear entries from boxes if filled.
                clear_product_view()
                messagebox.showinfo("SELECTION DELETED", "Records have been Deleted")
        
        
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create Pet Info Treeview frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        pet_label_frame_1=LabelFrame(tab_3,text="",bd=0)#,width=500,height=250) # Pads (space) inside of frame.
        pet_label_frame_1.pack(padx=20,pady=0,anchor=W) # Pads (space) outside of frame.
        tree_frame_pet=Frame(pet_label_frame_1)  
        tree_frame_pet.pack(pady=5,padx=5)
        dog_photo=StringVar()
        dog_photo.set("")
        #,,,Create Treeview Scroll,,,
        global pet_tree
        pet_tree_scroll=Scrollbar(tree_frame_pet)
        pet_tree_scroll.pack(side=RIGHT,fill=Y)
        #,,,Create the Treeview,,,,
        pet_tree=ttk.Treeview(tree_frame_pet,height=5,yscrollcommand=pet_tree_scroll.set,selectmode="extended")
        pet_tree.pack()
        #,,,Configure the scrollbar,,,
        pet_tree_scroll.config(command=pet_tree.yview)
        #,,,Define the Columns,,,
        pet_tree["columns"]=("Owner","Name","Age","Pet ID","Status","Type","DOB","Breed","Size","Micro chip","Chip ID","Gender","Weight","Colour","Temper","Reg date","Health","Attributes",)
        #,,,Format the Columns,,,
        pet_tree.column("#0",width=0,stretch=NO)
        pet_tree.column("Owner",anchor=W,width=120)
        pet_tree.column("Name",anchor=W,width=90)
        pet_tree.column("Age",anchor=CENTER,width=40)
        pet_tree.column("Pet ID",anchor=CENTER,width=40)
        pet_tree.column("Status",anchor=W,width=60)
        pet_tree.column("Type",anchor=W,width=40)
        pet_tree.column("DOB",anchor=W,width=70)
        pet_tree.column("Breed",anchor=W,width=70)
        pet_tree.column("Size",anchor=CENTER,width=50)
        pet_tree.column("Micro chip",anchor=CENTER,width=80)
        pet_tree.column("Chip ID",anchor=W,width=80)
        pet_tree.column("Gender",anchor=W,width=50)
        pet_tree.column("Weight",anchor=CENTER,width=50)
        pet_tree.column("Colour",anchor=W,width=70)
        pet_tree.column("Temper",anchor=W,width=70)
        pet_tree.column("Reg date",anchor=W,width=80)
        pet_tree.column("Health",anchor=W,width=70)
        pet_tree.column("Attributes",anchor=W,width=100)
        #,,,Create Headings,,,,,,,,,,,,,,,,,,,,,,,,
        pet_tree.heading("#0",text="",anchor=W)
        pet_tree.heading("Owner",text="Owner",anchor=W)
        pet_tree.heading("Name",text="Name",anchor=CENTER)
        pet_tree.heading("Age",text="Age",anchor=CENTER)
        pet_tree.heading("Pet ID",text="Pet ID",anchor=CENTER)
        pet_tree.heading("Status",text="Status",anchor=CENTER)
        pet_tree.heading("Type",text="Type",anchor=CENTER)
        pet_tree.heading("DOB",text="DOB",anchor=CENTER)
        pet_tree.heading("Breed",text="Breed",anchor=CENTER)
        pet_tree.heading("Size",text="Size",anchor=CENTER)
        pet_tree.heading("Micro chip",text="Micro chip",anchor=CENTER)
        pet_tree.heading("Chip ID",text="Chip ID",anchor=CENTER)
        pet_tree.heading("Gender",text="Gender",anchor=CENTER)
        pet_tree.heading("Weight",text="Weight",anchor=CENTER)
        pet_tree.heading("Colour",text="Colour",anchor=CENTER)
        pet_tree.heading("Temper",text="Temper",anchor=CENTER)
        pet_tree.heading("Reg date",text="Reg date",anchor=CENTER)
        pet_tree.heading("Health",text="Health",anchor=CENTER)
        pet_tree.heading("Attributes",text="Attributes",anchor=CENTER)
        
        #,,, create striped rows in treeview,,,,,,,,,,,,,,,,,,,,,
        pet_tree.tag_configure("oddrow",background="white")
        pet_tree.tag_configure("evenrow",background="lightblue")
        
        def query_pet_database():
            #,,,clear the treeview
            for record in pet_tree.get_children():
                pet_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute("SELECT rowid, * FROM customer_pets")
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            
            for record in records:
                if count % 2==0:
                    pet_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14],record[15],record[16],record[17]), tags=("evenrow",))
                else:
                    pet_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14],record[15],record[16],record[17]), tags=("oddrow",))
                #,,,increment counter adds 1 to each record entered (ID number),,,,
                count +=1   
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            petinfo.set("")
            dog_photo.set("")
        query_pet_database()
            #fn_entry.focus() #forget
        
        
        #creates popup box for search.
        def lookup_pet():
            button_del_pet["state"]="disabled"
            global pet_search_entry, search
            search=Toplevel(self)
            search.title("Search Pet Records")
            search.geometry("400x200",)
            search.resizable(False,False)
            #,,,create lable frame
            search_frame=LabelFrame(search,text="Pet Name")
            search_frame.pack(padx=10,pady=10)
            #,,,add entry box
            pet_search_entry=Entry(search_frame,font=("Helvetica",18))
            pet_search_entry.pack(padx=20,pady=20,)
            #,,,focus cursor on start.
            pet_search_entry.focus()
            #,,,add a button
            search_button1=Button(search,text="Search Records",command=search_pet)
            search_button1.pack(padx=20,pady=20)
        
        
        def talk_search_pet():
            engine=pyttsx3.init()
            engine.say(f'''NO results found, check for spelling errors or try another name.''')
            engine.runAndWait()
        
        
        def search_pet():
            lookup_pet=pet_search_entry.get()
            #,,,close the search box
            search.destroy()
            #,,,clear the treeview
            for record in pet_tree.get_children():
                 pet_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            c.execute("SELECT rowid, * FROM customer_pets WHERE pet_name like ? ",(lookup_pet,))
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0
            
            for record in records:
                if count % 2==0:
                    pet_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14],record[15],record[16],record[17]), tags=("evenrow",))
                else:
                    pet_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14],record[15],record[16],record[17]), tags=("oddrow",))
                #,,,increment counter,,,,
                count +=1 
            if not records:
                query_pet_database()
                petinfo.set("NO results found for that search, check for spelling errors or try another name.")
                records=talk_search_pet()  
                #,,,commit the changes
            conn.commit()
                #,,,close the connection
            conn.close()
            
        
        def help_pet_info():
            a=(f'Click SHOW ALL to display all records\n')
            b=(f"then select a record to display it's information.\n\n")
            c=(f'Click SEARCH and enter the pet name.\n\n')
            d=(f"Click NEW PET FORM to register a new pet.\n\n")
            e=(f'WARNING... Deleteing will permanently remove from database.\n')
            f=(f'To remove a record, select a record in view, click DELETE RECORD.\n')
            g=(f'(you will have the option to continue or cancel the deletion.)\n\n')
            h=(f'If you need a reminder, click HELP!.')
            
            answer=a+b+c+d+e+f+g+h
            petinfo.set(answer)
        
        #,,, Select Record from pet treeviw(bind),,,,,,,,,,,,,,,,,,
        def select_pet_record(e):
            global pet_photo_label
            #code to ignore header being clicked.
            region_clicked=pet_tree.identify_region(e.x, e.y)
            if region_clicked not in ('pet_tree','cell'):
                return
                #print("Clicked Header")
            else:
                button_del_pet["state"]="normal"
                button_update["state"]="normal"
                #client_name_entry.delete(0,END)
                #pet_name_entry.delete(0,END)
                #,,,Grabe entry on click
                selected=pet_tree.focus()
                #,,,Grab record values.
                values=pet_tree.item(selected,"values")
                #,,,variable names for output to boxes/labels.
                pet_owner=values[0]
                pet_name=values[1]
                pet_age=values[2]
                pet_id=values[3]
                status=values[4]
                pet_type=values[5]
                pet_dob=values[6]
                breed=values[7]
                size=values[8]
                micro_chip=values[9]
                chip_number=values[10]
                gender=values[11]
                weight=values[12]
                colour=values[13]
                temper=values[14]
                reg_date=values[15]
                health=values[16]
                attributes=values[17]
                #pet_photo=values[18]#pet photo
                secret_id.delete(0,END)
                secret_id.insert(0,values[3])
                secret_pet_name.delete(0,END)
                secret_pet_name.insert(0,values[1])
                #fn_entry.focus()
                #print(pet_owner,pet_name)
                a=(f'System ID: {pet_id}.\n')
                a1=(f"The owner of {pet_name} is {pet_owner}.\n\n")
                a2=(f"{pet_name} is a {colour} {breed} {pet_type}\n")
                a3=(f"Weighing about {weight}kg\n")
                a4=(f'{pet_age} year old {gender}, born {pet_dob}\n\n')
                a5=(f'{pet_name} is {health} with a {temper} temper.\n\n')
                a6=(f'Helpful info:- {attributes}')
                #e=(f'Pick-up service requested: {pick_up}, Drop off service requested: {drop_off}.\n\n')
                #f=(f'TOTAL BALANCE for the accommodation only £{total_price}\n')
                g=a+a1+a2+a3+a4+a5+a6
                petinfo.set(g)
                #dog_photo.set("PHOTO OF DOG\nPLACE HEAR")
                #dog_photo.set(pet_photo)#Shows pet photo on screen.
                #,,,Show an Image,(pet photo),,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
                pet_image=(Image.open("michell_photo.jpg"))#("michell_photo.jpg"))
                width, height=pet_image.size
                width_new=int(width/3)
                height_new=int(width/3)
                pet_image_resized=pet_image.resize((width_new,height_new))#this will resizes the image keep ratio.
                pet_image=pet_image.resize((100,100))#resizes the image.
                pet_imagePH=ImageTk.PhotoImage(pet_image)
                pet_photo_label=tk.Label(tab_3,image=pet_imagePH)
                pet_photo_label.image=pet_imagePH#this must be put to show the image.
                pet_photo_label.place(x=400,y=315)#grid(row=0,column=1,padx=20,pady=20)
        
        #,,, Bind the tree,,,,when clicked,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        pet_tree.bind("<ButtonRelease-1>",select_pet_record)
        
        
        def convert_pick():#converts file to save to db.
            filename="michell_photo.jpg"
            with open(filename, 'rb') as file:
                photo=file.read()
            return photo
        
        def get_photo():
            photo=convert_pick()
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute(f"UPDATE customer_pets SET photo = ? WHERE oid= 4",([photo]))#WORKING
            #c.execute("INSERT INTO pet_images(datestamp, dog_name, photo) VALUES (?,?,?)",(datestamp, dog_name, pet_photo))
            conn.commit()
            c.close()
            
        #get_photo()
        
        #https://youtu.be/vvq6xRziKns
        def play_video():
            cap = cv2.VideoCapture("michelle_video_clip.mp4")#Name of the file/file path needed.
            #if (cap.isOpened()==False):
            #    print("Error Opening Clip.")    
            while(cap.isOpened()):
                ret,frame = cap.read()
                #if ret == True:
                frame = cv2.resize(frame,(600,700))#size of window.
                cv2.imshow("video play",frame)
                if (cv2.waitKey(23) & 0xFF==ord('q')):
                    break    
            cap.release()
            cv2.destroyWindow()
        
        #https://youtu.be/vvq6xRziKns
        #https://youtu.be/Bn1n1diGv_0?list=PLCC34OHNcOtoC6GglhF3ncJ5rLwQrLGnV
        # opens popup file window
        def open_dog_mp4_file():
            try:
                filepath = fd.askopenfilename(initialdir="C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\video_clips_mp4",
                                            title="Select a Video Clip to open.",
                                            filetypes=[("Video Clips","*.mp4")])#,("all files","*.*")])#[("text files","*.txt"),("CSV files","*.csv"),("all files","*.*")])
                
                if filepath:
                    #pdf_file=PyPDF2.PdfFileReader(filepath)
                    os.startfile(filepath)
                
                #print(file.read())
            except ValueError:
                messagebox.askretrycancel("FILE NOT FOUND","File did not open\nTry again?.")
                #text=None
            except FileNotFoundError:
                messagebox.showinfo("INFO","Information\nYou closed the dialog box without choseing a file")
                #text=None
        
        
        #,,,Buttons Pet infomation,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        button1_pet=HoverButton(button_label_frame,text="",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT)
        button1_pet.grid(row=0,column=1,padx=1,pady=1,ipady=5,ipadx=5)
        button_new_pet_form=HoverButton(button_label_frame,text="NEW PET FORM",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=pet_window)
        button_new_pet_form.grid(row=0,column=2,padx=1,pady=1,ipady=5,ipadx=10)
        button_del_pet=HoverButton(button_label_frame,text="DELETE",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,state="disabled",command=remove_pet_info)#lambda:popupmsg("Not yet completed"))
        button_del_pet.grid(row=0,column=3,padx=1,pady=1,ipady=5,ipadx=10)
        button_show=HoverButton(button_label_frame,text="SHOW ALL",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=query_pet_database)
        button_show.grid(row=0,column=4,padx=1,pady=1,ipady=5,ipadx=10)
        button_update=HoverButton(button_label_frame,text="UPDATE RECORD",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,state="disabled",command=lambda:popupmsg("Not yet completed"))
        button_update.grid(row=0,column=5,padx=1,pady=1,ipady=5,ipadx=10)
        button_search=HoverButton(button_label_frame,text="SEARCH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lookup_pet)
        button_search.grid(row=0,column=6,padx=1,pady=1,ipady=5,ipadx=10)
        button_clear=HoverButton(button_label_frame,text="CLEAR/REFRESH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=query_pet_database)
        button_clear.grid(row=0,column=7,padx=1,pady=1,ipady=5,ipadx=10)
        button_help=HoverButton(button_label_frame,text="HELP!",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=help_pet_info)
        button_help.grid(row=0,column=8,padx=1,pady=1,ipady=5,ipadx=10)
        button_exit=HoverButton(button_label_frame,text="PLAY CLIPS",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=open_dog_mp4_file)#play_video)#lambda:popupmsg("Not yet completed"))
        button_exit.grid(row=0,column=9,padx=1,pady=1,ipady=5,ipadx=10)
        button_exit=HoverButton(button_label_frame,text="EXIT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:notebook.tab(tab_3,state="hidden"))
        button_exit.grid(row=0,column=10,padx=1,pady=1,ipady=5,ipadx=10)
        
        info_frame=Frame(tab_3)  
        info_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        #petinfo.set("There is a hidden Label here!.\nClick on a record in the treeview.")
        lb1=Label(info_frame,textvariable=petinfo,font="Verdana, 12",fg="#347083",justify=LEFT)
        lb1.grid(row=0,column=0,padx=20,pady=20)
        
        pet_photo_label=Label(tab_3,textvariable=dog_photo,fg="red",justify=LEFT)
        pet_photo_label.place(x=400,y=315)#grid(row=0,column=3,padx=0,pady=0)
        
        secret_id=Entry(info_frame,width=6,relief=FLAT,bg="systembuttonface",fg="systembuttonface")
        secret_id.grid(row=0,column=1,padx=0,pady=0)
        secret_pet_name=Entry(info_frame,width=10,relief=FLAT,bg="systembuttonface",fg="systembuttonface")
        secret_pet_name.grid(row=0,column=2,padx=0,pady=0)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,TAB 4 Sell Something Frames,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        top_label_frame1=Frame(tab_4,bd=1)  
        top_label_frame1.pack(pady=0,padx=0,anchor=W,fill=X)
        heading=Label(top_label_frame1,text="SELL SOMETHING (Reciept's & Invoice's.)",font="bold 20",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=17)
        button_label_frame=Frame(tab_4,bg="light blue",bd=1)  
        button_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        box_frame3=Frame(tab_4)  
        box_frame3.pack(pady=30,padx=30,anchor=W,fill=X)
        
        invoice_name_label=Label(box_frame3,text="Client Name")
        invoice_name_label.grid(row=0,column=0)
        client_name_ent=Entry(box_frame3)
        client_name_ent.grid(row=0,column=1)
        invoice_prod_label=Label(box_frame3,text="Product")
        invoice_prod_label.grid(row=1,column=0)
        prod_inv=Entry(box_frame3)
        prod_inv.grid(row=1,column=1)
        invoice_unit_label=Label(box_frame3,text="Unit")
        invoice_unit_label.grid(row=2,column=0)
        unit_inv=Entry(box_frame3)
        unit_inv.grid(row=2,column=1)
        invoice_price_label=Label(box_frame3,text="Price")
        invoice_price_label.grid(row=3,column=0)
        price_inv=Entry(box_frame3)
        price_inv.grid(row=3,column=1)
        
        #with open('bookings.csv','a',newline='') as csvfile:#shows all the content of the file.
            #writer=csv.writer(csvfile)
            
       # with open("bookings.csv","r") as f: #(with) is a context manager,,so you do not need to close the file after you open it.
                #content = f.read()
                #print("------------this is the full list of the bookings--------------")
                #print(content)
                #print("--------------------------------------------------------")            
        #add data to CSV file.
        def addbooking():
            bookings=[]
            data=[]#data will be the list we use to search the column where booking ref is.
            
            if client_name_ent.get()=="":
                messagebox.showerror("Missing Data!","Name needed.")
            elif prod_inv.get()=="":
                messagebox.showerror("Missing Data!","Product or Service name needed.")
            elif unit_inv.get()=="":
                messagebox.showerror("Missing Data!","Unit needed.")
            elif price_inv.get()=="":
                messagebox.showerror("Missing Data!","Price needed.")
            else:
                with open('bookings.csv') as csvfile:
                    reader=csv.reader(csvfile)
                    for row in reader:
                        data.append(row)
                    
                booking_ref= random.randint(0,1000)
                booking_ref=str(booking_ref)
                
                col1=[x[0] for x in data]
                #print(col1)
                #if the booking ref exists in column 1 (0) then do this.
                if booking_ref in col1:
                    booking_ref=int(booking_ref)
                    
                    #code to make new booking ref.
                    booking_ref2=random.randint(1,1000)
                    new_booking_ref=booking_ref * booking_ref2
                    new_booking_ref=str(new_booking_ref)
                    bookings.append(new_booking_ref)

                    print('The booking ref is: ',new_booking_ref)
                    client=client_name_ent.get()#,this will show as file name.
                    product=prod_inv.get()
                    unit=unit_inv.get()
                    unit=int(unit)
                    price=price_inv.get()
                    price=float(price)

                    bookings.append(client)
                    bookings.append(product)
                    bookings.append(unit)
                    bookings.append(price)
                    
                    print('Infomation added.')

                    with open('bookings.csv','a',newline='') as csvfile:
                        writer=csv.writer(csvfile)
                        writer.writerow(bookings)
                #if the booking ref is not in column 1 (0) continue as normal    
                else:
                    print('The booking ref is: ',booking_ref)
                    client=client_name_ent.get()#,this will show as file name.
                    product=prod_inv.get()
                    unit=unit_inv.get()
                    unit=int(unit)
                    price=price_inv.get()
                    price=float(price)

                    bookings.append(booking_ref)
                    bookings.append(client)
                    bookings.append(product)
                    bookings.append(unit)
                    bookings.append(price)
                    client_name_ent.delete(0,END)
                    prod_inv.delete(0,END)
                    unit_inv.delete(0,END)
                    price_inv.delete(0,END)
                    messagebox.showinfo("CSV FILE GENERATED", "YOUR CSV FILE WAS GENERATED SUCESSFULLY")
                    #print('Infomation added.')
                    with open('bookings.csv','a',newline='') as csvfile:
                        writer=csv.writer(csvfile)
                        writer.writerow(bookings)
        
        
        #makes invoice as txt file.
        #look into booking ref.py
        def make_client_invoice():
            if client_name_ent.get()=="":
                messagebox.showerror("Missing Data!","Name needed.")
            elif prod_inv.get()=="":
                messagebox.showerror("Missing Data!","Product or Service name needed.")
            elif unit_inv.get()=="":
                messagebox.showerror("Missing Data!","Unit needed.")
            elif price_inv.get()=="":
                messagebox.showerror("Missing Data!","Price needed.")
            else:
                client=client_name_ent.get()#,this will show as file name.
                unit=unit_inv.get()
                unit=int(unit)
                product=prod_inv.get()
                price=price_inv.get()
                price=float(price)
                document = Document()
                document.add_picture('C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\network.png',width=Inches(1))
                document.add_heading('Invoice',0)
                #this adds paragraph 1
                p1=document.add_paragraph('Dear ')
                p1.add_run(client).bold=True
                p1.add_run(',')
                #this adds paragraph 2
                p2=document.add_paragraph('Thank you for your recent booking of ')
                p2.add_run(str(unit)).bold=True
                p2.add_run(' days ')
                p2.add_run(product).bold=True
                p2.add_run(',')
                [document.add_paragraph('')for _ in range(2)]#adds 2 empty lines.
                table=document.add_table(rows=1, cols=4)
                hdr_cells=table.rows[0].cells
                hdr_cells[0].text='Product Name'
                hdr_cells[1].text='Units'
                hdr_cells[2].text='Unit Price'
                hdr_cells[3].text='Total Price'
                
                for i in range(4):
                    hdr_cells[i].paragraphs[0].runs[0].font.bold=True
                    
                row_cells=table.add_row().cells
                row_cells[0].text=product
                row_cells[1].text=f'{unit}'
                row_cells[2].text=f'{price:,.2f}'
                row_cells[3].text=f'{unit*price:,.2f}'
                
                [document.add_paragraph('')for _ in range(4)]#,adds empty spaces.
                
                document.add_paragraph('We appreciate your business and look forward to other bookings in the future.\nfor bank transfers please use.\nsort: 01-23-45 acc: 123456789.')
                document.add_paragraph('Sincerely')
                document.add_paragraph('Mr Test Invoice')
                document.save(f'{client}.docx')
                client_name_ent.delete(0,END)
                prod_inv.delete(0,END)
                unit_inv.delete(0,END)
                price_inv.delete(0,END)
                messagebox.showinfo("INVOICE GENERATED", "YOUR INVOICE WAS GENERATED SUCESSFULLY")


        #,,,Buttons Sell Something,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        button1_marketing=HoverButton(button_label_frame,text="",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT)#,command=lambda:popupmsg("Not yet completed"))
        button1_marketing.grid(row=0,column=1,padx=1,pady=1,ipady=5,ipadx=5)
        button_invoice_marketing=HoverButton(button_label_frame,text="INVOICE",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=make_client_invoice)#addbooking)#lambda:popupmsg("Not yet completed"))#make_client_invoice
        button_invoice_marketing.grid(row=0,column=2,padx=1,pady=1,ipady=5,ipadx=10)
        button_help_marketing=HoverButton(button_label_frame,text="HELP!",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:popupmsg("Not yet completed"))
        button_help_marketing.grid(row=0,column=3,padx=1,pady=1,ipady=5,ipadx=10)
        button_exit_marketing=HoverButton(button_label_frame,text="EXIT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:notebook.tab(tab_4,state="hidden"))
        button_exit_marketing.grid(row=0,column=5,padx=1,pady=1,ipady=5,ipadx=10)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,TAB 5 REPORTS,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        
        top_label_frame1=Frame(tab_5,bd=1)  
        top_label_frame1.pack(anchor=W,fill=X)
        heading=Label(top_label_frame1,text="REPORTS",font="bold 20",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=17)
        button_label_frame=Frame(tab_5,bg="light blue",bd=1)  
        button_label_frame.pack(anchor=W,fill=X)
        
        #,,,load Top Left logo/Image for all note pages,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #box_frame3=Frame(tab_5)  
        #box_frame3.pack_propagate(False)
        #box_frame3.pack(pady=30,padx=30,anchor=W,fill=X)
        networkpeople=(Image.open("network.png"))
        networkpeople=networkpeople.resize((180,180))#resizes the image.
        networkpeoplePH=ImageTk.PhotoImage(networkpeople)
        network_label=tk.Label(box_frame3,image=networkpeoplePH)
        network_label.image=networkpeoplePH#this must be put to show the image.
        global arrival_num
        employee_num=StringVar()
        employee_num.set("Employee's")
        arrival_num=StringVar()
        arrival_num.set("Arrivals")
        checkout_num=StringVar()
        checkout_num.set("Checkout's")
        penn=StringVar()
        boarding_num=StringVar()
        boarding_num.set("Boarding Income")
        vpenn=StringVar()
        client_num=StringVar()
        client_num.set("Client's")
        pet_num=StringVar()
        pet_num.set("Pet's")
        
        
        def client_count():
            #https://youtu.be/gOLYhfCQXL4
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            cli=c.execute("SELECT * FROM clients ")
            count=cli.fetchall()
            for x in count:
                booking_ref=(x[0])
            
            conn.commit()
            conn.close()
            count=str(len(count))
            client_num.set("Registered Client's "+ "\n\n\n" +"Total Client's "+"(" + count + ")")
            
        
        def sales_income():# needs try except...(TypeError)
            #https://youtu.be/gOLYhfCQXL4
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            t_sales=c.execute("SELECT SUM(total) AS total_sales FROM kennel_bookings ")
            for x in t_sales:
                total_sales=format(x[0],'.2f')
                income=('Total Boarding Income: £' + str(total_sales))
                #print(income)
                
            p_sales=c.execute("SELECT SUM(total) AS total_sales FROM kennel_bookings WHERE care_option = 'Premium'")
            for x in p_sales:
                prem_sales=format(x[0],'.2f')
                premium_income=('Premium Care Income: £' + str(prem_sales))
                #print(premium_income)
                        
            stand_sales=c.execute("SELECT SUM(total) AS total_sales FROM kennel_bookings WHERE care_option = 'Standard'")
            for x in stand_sales:
                stand_sales=format(x[0],'.2f')
                standard_income=('Standard Care Income: £' + str(stand_sales))
                #print(standard_income)
                
            l_sales=c.execute("SELECT SUM(total) AS total_sales FROM kennel_bookings WHERE care_option = 'Luxury'")
            for x in l_sales:
                lux_sales=format(x[0],'.2f')
                luxury_income=('Luxury Care Income: £' + str(lux_sales))
                #print(luxury_income)
                boarding_num.set("Total Boarding Income £"+ total_sales + "\n\nStandard £"+ stand_sales+ "\nPremium £"+ prem_sales+ "\nLuxury    £"+ lux_sales)

            conn.commit()
            conn.close()  
        
        
        def between_dates():
            #https://youtu.be/jngHjpBiiKo
            #https://youtu.be/tlmOm35lsps
            dt=date.today()
            tomorrows_date=dt + timedelta(days=1)
            todays_date=datetime.strftime(dt,'%d/%m/%Y')#string format.
            tomorrows_date=datetime.strftime(tomorrows_date,'%d/%m/%Y')#string format.
            
            day_3=dt + timedelta(days=2)
            day_4=dt + timedelta(days=3)
            day_5=dt + timedelta(days=4)
            day_6=dt + timedelta(days=5)
            day_7=dt + timedelta(days=6)
            day_3=datetime.strftime(day_3,'%d/%m/%Y')#string format.
            day_4=datetime.strftime(day_4,'%d/%m/%Y')#string format.
            day_5=datetime.strftime(day_5,'%d/%m/%Y')#string format.
            day_6=datetime.strftime(day_6,'%d/%m/%Y')#string format.
            day_7=datetime.strftime(day_7,'%d/%m/%Y')#string format.
            
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            #,,,find bookings between ???..(you can use calendar pick box for dates).
            arrival=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(todays_date,))#order by in_date desc")#desc

            returndates=arrival.fetchall()
            for x in returndates:
                booking_ref=(x[0])
                first_name=(x[1])
                last_name=(x[2])
                check_in=(x[8])
                check_out=(x[9])
                total_days=(x[10])
                #print(f"{booking_ref} {first_name} {last_name}, check in date: {check_in} check out date: {check_out} for {total_days} days.") 
            
            c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(tomorrows_date,))# order by in_date desc")#desc
            checkoutdate=c.fetchall()
            for x in checkoutdate:
                booking_ref=(x[0])
                first_name=(x[1])
                last_name=(x[2])
                check_in=(x[8])
                check_out=(x[9])
                total_days=(x[10])
                #print(f"{booking_ref} {first_name} {last_name}, check in date: {check_in} check out date: {check_out} for {total_days} days.")     
            
            
            d3=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_3,))#order by in_date desc")#desc
            next_day_3=d3.fetchall()
            d4=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_4,))#order by in_date desc")#desc
            next_day_4=d4.fetchall()
            d5=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_5,))#order by in_date desc")#desc
            next_day_5=d5.fetchall()
            d6=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_6,))#order by in_date desc")#desc
            next_day_6=d6.fetchall()
            d7=c.execute("SELECT * FROM kennel_bookings WHERE in_date = ?",(day_7,))#order by in_date desc")#desc
            next_day_7=d7.fetchall()
            day_3=int(len(next_day_3))
            day_4=int(len(next_day_4))
            day_5=int(len(next_day_5))
            day_6=int(len(next_day_6))
            day_7=int(len(next_day_7))
            conn.commit()
            conn.close()
            
            arrive_today=str(len(returndates))
            arrive_tomorrow=str(len(checkoutdate))
            total_week_arrive=str(int(arrive_today)+int(arrive_tomorrow)+day_3+day_4+day_5+day_6+day_7)#total for next 2 days
            arrival_num.set("Expected Arrivals\n\n   Today           " + "(" + arrive_today +")"+ "\n   Tomorrow      " + "(" + arrive_tomorrow +")"+ "\n   Next 7 day's  " + "(" + total_week_arrive +")")
                
        
        
        
        def bookin_hover(e):
            between_dates()
            
        def bookin_hover_leave(e):
            arrival_num.set("Arrivals")
            
        
        def checkout_dates():
            #https://strftime.org
            #https://youtu.be/jngHjpBiiKo
            #https://youtu.be/tlmOm35lsps
            #https://youtu.be/C37Q-4ar6cc  #good info on datetime/timedelta.
            dt=date.today()
            tomorrows_date=dt + timedelta(days=1)
            today_date=datetime.strftime(dt,'%d/%m/%Y')#%x=todays date#change the format to (2022/07/13 19-29-37) string format.
            tomorrows_date=datetime.strftime(tomorrows_date,'%d/%m/%Y')#%x=todays date#change the format to (2022/07/13 19-29-37) string format.
            
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            #,,,find bookings between ???..(you can use calendar pick box for dates).
            arrival=c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(today_date,))# order by in_date desc")#desc
            returndates=arrival.fetchall()
            
            for x in returndates:
                booking_ref=(x[0])
                first_name=(x[1])
                last_name=(x[2])
                check_in=(x[8])
                check_out=(x[9])
                total_days=(x[10])
                #print(f"{booking_ref} {first_name} {last_name}, check in date: {check_in} check out date: {check_out} for {total_days} days.") 
            
            c.execute("SELECT * FROM kennel_bookings WHERE out_date = ?",(tomorrows_date,))# order by in_date desc")#desc
            checkoutdate=c.fetchall()
            for x in checkoutdate:
                booking_ref=(x[0])
                first_name=(x[1])
                last_name=(x[2])
                check_in=(x[8])
                check_out=(x[9])
                total_days=(x[10])
                #print(f"{booking_ref} {first_name} {last_name}, check in date: {check_in} check out date: {check_out} for {total_days} days.")     
            conn.commit()
            conn.close()
            #print("In today " + str(len(returndates)))#gives the total of checkouts 
            #print("Out " + str(len(checkoutdate)))#gives the total of checkouts 
            arrive_today=str(len(returndates))
            out_today=str(len(checkoutdate))
            checkout_num.set("Expected Checkout's\n\n\nToday          " + "(" + arrive_today +")"+ "\nTomorrow     " + "(" + out_today +")")
        
        
        def employ_num_get():
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            #,,,get row from Database,,,,,,,
            c.execute("SELECT rowid, * FROM employees WHERE date_left = 'No' ")
            count=c.fetchall()
            for x in count:
                booking_ref=(x[11])
            
            c.execute("SELECT * FROM employees WHERE position = 'Volenteer' AND date_left = 'No'")
            title=c.fetchall()
            for x in title:
                booking_ref=(x[11])
            
            conn.commit()
            conn.close()
            count=str(len(count))
            title=str(len(title))
            employee_num.set("Active Employee's"+"\n\n\nTotal Staff   " +"(" + count + ")"+"\nVolenteer's  "+"(" + title + ")")
            
        
        
        def occupancy_num():
            dt=date.today()
            today_date=datetime.strftime(dt,'%d/%m/%Y')#%x=todays date#change the format to (2022/07/13 19-29-37) string format.
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            #,,,get row from Database (should have only 1 row),,,,,,,
            kennel_data=c.execute("SELECT DISTINCT (enclosure) FROM kennel_bookings WHERE in_date = ? ",(today_date,))
            kennels=kennel_data.fetchall()
            ken_total=int(len(kennels))#gets the number of kennels in use(but need between dates?????)
            inuse=(str(ken_total))
        
            c.execute("SELECT * FROM general_settings")
            records=c.fetchall()
            for record in records:
                occupancy_amount=(record[1])
                num_ken=int(occupancy_amount)
                used_ken=int(inuse)
                available_ken=(num_ken-used_ken)
                #print(kennels)    
                
            conn.commit()
            conn.close()
            penn.set("Total "+pen_name +"'s    "+" ("+str(occupancy_amount)+")\n\n\n"+pen_name+"'s "+"available  ""("+str(available_ken)+")\n"+ pen_name+"'s in use      ""("+inuse+")")
        
        
        def count_pets():#https://youtu.be/gOLYhfCQXL4
            conn=sqlite3.connect("tree_crm.db")
            c=conn.cursor()
            cli=c.execute("SELECT * FROM customer_pets ")
            count=cli.fetchall()
            for x in count:
                booking_ref=(x[0])
            
            conn.commit()
            conn.close()
            count=str(len(count))
            pet_num.set("Registered Pet's\n\n\n""Total Pets "+ "(" + count + ")")
        
        
        def checkout_hover(e):
            checkout_dates()
            
        def checkout_hover_leave(e):
            checkout_num.set("Checkout's")
            
        def boarding_reports_hover(e):
            sales_income()
            
        def boarding_reports_hover_leave(e):
            boarding_num.set("Boarding Income")
            
        def client_hover(e):
            client_count()
            
        def client_hover_leave(e):
            client_num.set("Client's")
               
        def occupancy_hover(e):
            occupancy_num()
            
        def occupancy_hover_leave(e):
            penn.set(pen_name + " Occupancy")
            
        def employ_hover(e):
            employ_num_get()
            
        def employ_hover_leave(e):
            employee_num.set("Employee's ")
        
        def pet_hover(e):
            count_pets()
            
        def pet_hover_leave(e):
            pet_num.set("Pet's ")
        
        def select_6():
            notebook.select(6) 
            
        def select_3():
            notebook.select(3)
        
        #,,takes enclosure name from general settings database table and changes label text in reports,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        c=conn.cursor()
        sql_enc_name=("SELECT enclosure_name FROM general_settings")
        c.execute(sql_enc_name)
        enclosure_name=c.fetchone()
        pen_name=enclosure_name[0]#name of enclosure
        conn.commit()
        conn.close()
        
        penn.set(pen_name + " Occupancy")
        vpenn.set("Income Statment")
        
        reports_frame3=Frame(tab_5)  
        reports_frame3.pack(padx=30,pady=(20,3),anchor=W,fill=X)
        
        #Buttons-reports#https://youtu.be/CJ1TqEGRSe0
        button1_reports=HoverButton(reports_frame3,text="Purchases",font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#ffa742",fg="white",height=6,width=20,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Not yet completed"))
        button1_reports.grid(row=0,column=0,padx=4,pady=3,ipady=5,ipadx=5)
        button2_reports=HoverButton(reports_frame3,text="Purchase Returns",font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#1bb5f7",fg="white",height=6,width=20,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Not yet completed"))
        button2_reports.grid(row=0,column=1,padx=4,pady=3,ipady=5,ipadx=5)
        button3_reports=HoverButton(reports_frame3,text="Sales",font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#529c81",fg="white",height=6,width=20,relief=FLAT,anchor="nw")#,command=lambda:popupmsg("Not yet completed"))
        button3_reports.grid(row=0,column=2,padx=4,pady=3,ipady=5,ipadx=5)
        
        client_reports=HoverButton(reports_frame3,textvariable=client_num,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#298203",fg="white",height=6,width=20,relief=FLAT,anchor="nw",command=lambda:controller.show_frame(PageTwo))
        client_reports.grid(row=0,column=3,padx=4,pady=3,ipady=5,ipadx=5)
        client_reports.bind("<Enter>",client_hover)
        client_reports.bind("<Leave>",client_hover_leave)
        
        pet_reports=HoverButton(reports_frame3,textvariable=pet_num,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#975ac7",fg="white",height=6,width=20,relief=FLAT,anchor="nw",command=select_3)
        pet_reports.grid(row=0,column=4,padx=4,pady=3,ipady=5,ipadx=5)
        pet_reports.bind("<Enter>",pet_hover)
        pet_reports.bind("<Leave>",pet_hover_leave)
        
        reports_frame4=Frame(tab_5)  
        reports_frame4.pack(pady=5,padx=30,anchor=NW,fill=X)
        checkout_reports=HoverButton(reports_frame4,textvariable=checkout_num,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#856fa3",fg="white",height=6,width=20,relief=FLAT,anchor="nw",command=lambda:controller.show_frame(PageThree))
        checkout_reports.grid(row=0,column=0,padx=4,pady=3,ipady=5,ipadx=5)
        checkout_reports.bind("<Enter>",checkout_hover)
        checkout_reports.bind("<Leave>",checkout_hover_leave)
        
        bookin_reports=HoverButton(reports_frame4,textvariable=arrival_num,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#946065",fg="white",height=6,width=31,relief=FLAT,anchor="nw",command=lambda:controller.show_frame(PageThree))
        bookin_reports.grid(row=0,column=1,padx=4,pady=3,ipady=5,ipadx=5)
        bookin_reports.bind("<Enter>",bookin_hover)
        bookin_reports.bind("<Leave>",bookin_hover_leave)
        
        employ_reports=HoverButton(reports_frame4,textvariable=employee_num,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#d97f43",fg="white",height=6,width=20,relief=FLAT,anchor="nw",command=select_6)#,command=lambda:popupmsg("Not yet completed"))
        employ_reports.grid(row=0,column=2,padx=4,pady=3,ipady=5,ipadx=5)
        employ_reports.bind("<Enter>",employ_hover)
        employ_reports.bind("<Leave>",employ_hover_leave)
        
        boarding_reports=HoverButton(reports_frame4,textvariable=boarding_num,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#eb757d",fg="white",height=6,width=32,relief=FLAT,anchor="nw",command=lambda:controller.show_frame(PageThree))
        boarding_reports.grid(row=0,column=3,padx=4,pady=3,ipady=5,ipadx=5)
        boarding_reports.bind("<Enter>",boarding_reports_hover)
        boarding_reports.bind("<Leave>",boarding_reports_hover_leave)
        
        reports_frame5=Frame(tab_5)  
        reports_frame5.pack(pady=5,padx=30,anchor=NW,fill=X)
        button1_reports=HoverButton(reports_frame5,textvariable=vpenn,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#05aff7",fg="white",height=6,width=31,relief=FLAT,anchor="nw",command=lambda:popupmsg("Manager only, Password Needed"))
        button1_reports.grid(row=0,column=0,padx=4,pady=3,ipady=5,ipadx=5)
        
        occupancy_reports=HoverButton(reports_frame5,textvariable=penn,font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#84959c",fg="white",height=6,width=20,relief=FLAT,anchor="nw",command=lambda:controller.show_frame(PageThree))
        occupancy_reports.grid(row=0,column=1,padx=4,pady=3,ipady=5,ipadx=5)
        occupancy_reports.bind("<Enter>",occupancy_hover)
        occupancy_reports.bind("<Leave>",occupancy_hover_leave)
        
        button3_reports=HoverButton(reports_frame5,text="Bank Book",font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#a67289",fg="white",height=6,width=32,relief=FLAT,anchor="nw")
        button3_reports.grid(row=0,column=2,padx=4,pady=3,ipady=5,ipadx=5)
        button4_reports=HoverButton(reports_frame5,text="Cash Book",font=LARGE_FONT,activebackground = '#b9b8ba',cursor="hand2",bg="#2d8734",fg="white",height=6,width=20,relief=FLAT,anchor="nw")
        button4_reports.grid(row=0,column=3,padx=4,pady=3,ipady=5,ipadx=5)
        
        
        
        '''#scrollable frame for content,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #https://youtu.be/VmlgrrXAqb4
        wraper1=LabelFrame(tab_5,bg="blue")
        wraper1.pack(fill=BOTH,expand=1)
        #wraper2=LabelFrame(tab_5,bg="")
        #wraper2.pack(fill=BOTH,expand=1)
        my_canvas=Canvas(wraper1)
        my_canvas.pack(side=LEFT,fill="both",expand="yes")
        yscrollbar=ttk.Scrollbar(wraper1,orient="vertical",command=my_canvas.yview)
        yscrollbar.pack(side=RIGHT,fill="y")
        my_canvas.configure(yscrollcommand=yscrollbar.set)
        my_canvas.bind('<Configure>',lambda e: my_canvas.configure(scrollregion=my_canvas.bbox('all')))
        my_frame=Frame(my_canvas)
        my_canvas.create_window((0,0),window=my_frame, anchor="nw")
        #take number of kennals entered in database and make buttons.
        #,,, connect to db for kennal amount
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,get row from Database (should have only 1 row),,,,,,,
        c.execute("SELECT * FROM general_settings")
        records=c.fetchall()
        kennels=(records[0][1])#gets total enclosures(kennals)
        
        #makes button for each enclosure and places them on screen.
        for button in range(kennels):
            Button(my_frame,text=f'Kennal {button}').grid(row=button,column=0,padx=10,pady=2)
            
        #,,,commit the changes
        conn.commit()
        #,,, close the connection
        conn.close()'''
    
        #,,,Buttons,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        button1_finance=HoverButton(button_label_frame,text="",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT)#,command=lambda:popupmsg("Not yet completed"))
        button1_finance.grid(row=0,column=3,padx=1,pady=1,ipady=5,ipadx=5)
        #button_invoice_finance=HoverButton(button_label_frame,text="INVOICE",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:popupmsg("Not yet completed"))
        #button_invoice_finance.grid(row=0,column=4,padx=1,pady=1,ipady=5,ipadx=10)
        #button_help_finance=HoverButton(button_label_frame,text="HELP!",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:popupmsg("Not yet completed"))
        #button_help_finance.grid(row=0,column=5,padx=1,pady=1,ipady=5,ipadx=10)
        button_exit_finance=HoverButton(button_label_frame,text="EXIT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:notebook.tab(tab_5,state="hidden"))
        button_exit_finance.grid(row=0,column=6,padx=1,pady=1,ipady=5,ipadx=10)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,TAB 6 Employees,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        top_label_frame1=Frame(tab_6,bd=1)  
        top_label_frame1.pack(pady=0,padx=0,anchor=W,fill=X)
        heading=Label(top_label_frame1,text="EMPLOYEES",font="bold 20",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=17)
        button_label_frame=Frame(tab_6,bg="light blue",bd=1)  
        button_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        
        box_frame3=Frame(tab_6)  
        box_frame3.pack(pady=30,padx=30,anchor=W,fill=X)
        dt_stamp=time.strftime("%d%m%Y-%H%M%S")
        leave_num=StringVar()
        leave_num.set("")
        staff_labl=StringVar()
        staff_labl.set("Staff details are used for program interface only.")
        
        #,,,update employee Details,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def update_employee_details():#,,updates Business Details boxes.
            employ_info.set("")
            staff_labl.set("Staff details are used for program interface only.")
            talk_update()
            if employee_fname_ent.get()=="":
                messagebox.showerror("Missing Data!","First name needed.")
            elif employee_lname_ent.get()=="":
                messagebox.showerror("Missing Data!","Last name needed.")
            elif employee_id_ent.get()=="":
                messagebox.showerror("Missing Data!","ID needed.")
            #elif employee_address_ent()=="":
                #messagebox.showerror("Missing Data!","address needed.")
            elif employee_town_ent.get()=="":
                messagebox.showerror("Missing Data!","Town needed.")
            elif employee_city_ent.get()=="":
                messagebox.showerror("Missing Data!","City needed.")
            elif employee_phone_ent.get()=="":
                messagebox.showerror("Missing Data!","Phone number needed.")
            #elif employee_email_ent.get()=="":
                #messagebox.showerror("Missing Data!","Email needed.")
            elif start_date.get()=="":
                messagebox.showerror("Missing Data!","Start date needed.")
            #elif leave_date.get()=="":
                #messagebox.showerror("leave date needed.Enter active if still employed")
            elif position_combo.get()=="":
                messagebox.showerror("Position needed.")
            #elif employee_hours_ent.get()=="":
                #messagebox.showerror("Contracted Hours needed.")
            elif ni_num_ent.get()=="":
                messagebox.showerror("NI Number needed.")
            else:
                #,,, Connect to database.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record.(UPDATE) employees 
                c.execute("""UPDATE employees SET 
                          first_name=:first_name, 
                          last_name=:last_name, 
                          address=:address,
                          town=:town,
                          city=:city,
                          county=:county,
                          postcode=:postcode,
                          phone=:phone,
                          email=:email,
                          start_date=:start_date,
                          date_left=:date_left,
                          position=:position,
                          ni_num=:ni_num,
                          date_stamp=:date_stamp
                          WHERE oid = :id """,
                            {
                            "first_name":employee_fname_ent.get().capitalize(),
                            "last_name":employee_lname_ent.get().capitalize(),
                            "id":employee_id_ent.get(),
                            "address":employee_address_ent.get().title(),
                            "town":employee_town_ent.get().title(),
                            "city":employee_city_ent.get().title(),
                            "county":employee_county_ent.get().title(),
                            "postcode":employee_post_ent.get().upper(),
                            "phone":employee_phone_ent.get(),
                            "email":employee_email_ent.get(),
                            "start_date":start_date.get(),
                            "date_left":leave_date1.get(),
                            "position":position_combo.get().capitalize(),
                            "ni_num":ni_num_ent.get().upper(),
                            "date_stamp":dt_stamp,#time.strftime("%d/%m/%Y-%H:%M:%S")
                            })
                employee_fname_ent.delete(0,END)
                employee_lname_ent.delete(0,END)
                employee_id_ent.delete(0,END)
                employee_address_ent.delete(0,END)
                employee_town_ent.delete(0,END)
                employee_city_ent.delete(0,END)
                employee_county_ent.delete(0,END)
                employee_post_ent.delete(0,END)
                employee_phone_ent.delete(0,END)
                employee_email_ent.delete(0,END)
                start_date.delete(0,END)
                leave_date1.delete(0,END)
                position_combo.delete(0,END)
                #leave_date.delete(0,END)
                ni_num_ent.delete(0,END)
                #,,,commit the changes
                conn.commit()
                #,,, close the connection
                conn.close()
                fetch_current_employees()
                messagebox.showinfo("UPDATED", "Employee settings have been\nSuccessfully Updated")
                save["state"]="normal"
                
                
        
        
        def help_staff():
            #d=(f"For NEW employee, enter.\n\n")
            a=(f"To UPDATE click SEARCH, enter employee's last name in the pop-up box\n")
            b=(f"this will display the employee information to the display below.\n")
            c=(f'Select employee from display, update details if needed.\n\n')
            e=(f'WARNING... DELETE EMPLOYEE will permanently delete from the database.\n')
            f=(f"To remove a record, view the employee's details, click DELETE EMPLOYEE.\n")
            g=(f'(you will have the option to continue or cancel the deletion.)\n\n')
            f1=(f"LEAVERS!,, change Leaver to (Yes) enter date click UPDATE RECORDS.\n\n")
            f2=(f'EXIT simply closes the EMPLOYEE screen.\n\n')
            h=(f'If you need a reminder, click HELP!.')
            
            answer=a+b+c+e+f+g+f1+f2+h
            employ_info.set(answer)
            clear_employee_box()
            
        
        def staff_get():
            #,,, search database or connect to one.
            employ_info.set("")
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row from Database,,,,,,,
            c.execute("SELECT rowid, * FROM employees WHERE date_left = 'No'")
            records=c.fetchall()
            global count
            
            count = 0 #,,,start of ID
            for record in records:
                count=count +1
                #print(records)
            
            num = 2
            for record in records:
                intro=Label(frameemp,text=f'Number of employee records on file {count}\n',font="Verdana, 12",fg="#347083",justify="left")
                intro.grid(row=1,column=0,columnspan=5,padx=5,sticky=N)
                id=Label(frameemp,text=f'ID: {record[0]}',font="Verdana, 12",fg="#347083",justify="left")
                id.grid(row=num,column=0,padx=2,sticky=N)
                name1=Label(frameemp,text=f' {record[1]}',font="Verdana, 12",fg="#347083",justify="left")
                name1.grid(row=num,column=1,padx=2,sticky=N)
                name2=Label(frameemp,text=record[2],font="Verdana, 12",fg="#347083",justify="left")
                name2.grid(row=num,column=2,padx=2,sticky=N)
                phone=Label(frameemp,text=f'Tel: {record[8]}',font="Verdana, 12",fg="#347083",justify="left")
                phone.grid(row=num,column=3,padx=2,sticky=N)
                position=Label(frameemp,text=record[12],font="Verdana, 12",fg="#347083",justify="left")
                position.grid(row=num,column=4,padx=2,sticky=N)
                
                num=num+1
           
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            clear_employee_box()
        
        
        def talk_search_employee():
            engine=pyttsx3.init()
            engine.say(f'''Sorry, NO results found for that search, check for spelling errors or try another name.''')
            engine.runAndWait() 
            #messagebox.showinfo("NO MATCH FOUND", "No Employee match found\ntry another spelling.")
            #employ_info.set("NO results found for that search, check for spelling errors or try another name\nClick REFRESH to see all records.")
        
        def search_records(*args):#args for enter key binding(event)
            employee_fname_ent.delete(0,END)
            employee_lname_ent.delete(0,END)
            employee_id_ent.delete(0,END)
            employee_address_ent.delete(0,END)
            employee_town_ent.delete(0,END)
            employee_city_ent.delete(0,END)
            employee_county_ent.delete(0,END)
            employee_post_ent.delete(0,END)
            employee_phone_ent.delete(0,END)
            employee_email_ent.delete(0,END)
            start_date.delete(0,END)
            leave_date1.delete(0,END)
            position_combo.delete(0,END)
            #leave_date.delete(0,END)
            ni_num_ent.delete(0,END)
            
            lookup_record=search_entry.get()
            #,,,close the search box
            search.destroy()
            #,,,clear the treeview
            for record in employee_tree.get_children():
                 employee_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            c.execute("SELECT rowid, * FROM employees WHERE date_left = 'No' AND last_name like ?",(lookup_record,))
            record=c.fetchall() 
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0
                 
            for record in record:
                #print(record[1])
                if count % 2==0:
                    employee_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[3],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[12],record[13]), tags=("evenrow",))
                else:
                    employee_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[3],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[12],record[13]), tags=("oddrow",))
                #,,,increment counter,,,,
                count +=1 
            if not record:
                fetch_current_employees()
                #clientinfo.set("NO results found for that search, check for spelling errors or try another name.")
                messagebox.showinfo("SORRY","No results found\nCheck spelling")
                #records=talk_search_client()   
            #,,,commit the changes
            conn.commit()
                #,,,close the connection
            conn.close()
            save["state"]="disabled"
            employ_info.set("")
            
           
        
        global lookup_staff_records
        def lookup_staff_records():
            employ_info.set("")
            global search_entry, search
            search=Toplevel(self)
            search.title("Search Records")
            search.geometry("400x200",)
            search.resizable(False,False)
            #,,,create lable frame
            search_frame=LabelFrame(search,text="Employee Last name")
            search_frame.pack(padx=10,pady=10)
            #,,,add entry box
            search_entry=Entry(search_frame,font=("Helvetica",18))
            search_entry.bind("<Return>",search_records)#enter key binding(dont forget *args).
            search_entry.pack(padx=20,pady=20,)
            #,,,focus cursor on start.
            search_entry.focus()
            #,,,add a button
            search_button1=Button(search,text="Search Records",command=search_records)
            search_button1.pack(padx=20,pady=20)
            clear_employee_box()
        
        
        def delete_employee():
            talk()
            if employee_id_ent.get()=="" :
                messagebox.showerror("ERROR","Search for a record to delete\n\nSearch useing last name or employee reference number")
            else:
                app.bell()
                response = messagebox.askyesno("WARNING!!", "This will DELETE the employee from the Database.\n\nAre You Sure?!.\n\nDouble check the details first.")
                #,,,add logic for message box.
                if response==1:
                    #,,,log in database.    
                    conn=sqlite3.connect("tree_crm.db")
                    #,,,create a cursor.
                    c=conn.cursor()
                    #,,,change employee to leaver, hide from current employees.(keep in database)
                    c.executemany("UPDATE employees SET date_left = 'Yes' WHERE oid = ?", employee_id_ent.get())
                    #c.executemany("DELETE FROM employees WHERE oid = ?", employee_id_ent.get())
                    #,,,commit the changes
                    conn.commit()
                        #,,,close the connection
                    conn.close()
                        #,,,clear entries from boxes if filled.
                    clear_employee_box()
                    
                    messagebox.showinfo("SELECTION DELETED", "Records have been Deleted")
        
        
        def talk_update():
            #global engine
            if employee_id_ent.get()=="" :
                engine=pyttsx3.init()
                engine.say(f'''you need to search for an employee first before you can update the record,
                           search using last name or employee reference number''')
                #engine.say(my_entry.get())
                #engine.say(my_entry2.get())
                #engine.say(my_entry3.get())
            else:
                engine=pyttsx3.init()
                engine.say(f''' you have updated, employee reference number {employee_id_ent.get()}, {employee_fname_ent.get()} {employee_lname_ent.get()},
                            ''')
                
            engine.runAndWait()
        
        
        
        def talk():
            #global engine
            if employee_id_ent.get()=="" :
                engine=pyttsx3.init()
                engine.say(f'''you need to search for an employee first before you can delete the record,
                           search using last name or employee reference number''')
                #engine.say(my_entry.get())
                #engine.say(my_entry2.get())
                #engine.say(my_entry3.get())
            else:
                engine=pyttsx3.init()
                engine.say(f''' WARNING,, you have chosen, employee reference number{employee_id_ent.get()}, {employee_fname_ent.get()},{employee_lname_ent.get()}to be deleted from all records,
                        if this is correct, select yes,
                        alternatively, if this was an error in any way, select no..''')
            engine.runAndWait()
        
        
        def clear_employee_box():
            employee_tree.forget
            employee_fname_ent.delete(0,END)
            employee_lname_ent.delete(0,END)
            employee_id_ent.delete(0,END)
            employee_address_ent.delete(0,END)
            employee_town_ent.delete(0,END)
            employee_city_ent.delete(0,END)
            employee_county_ent.delete(0,END)
            employee_post_ent.delete(0,END)
            employee_phone_ent.delete(0,END)
            employee_email_ent.delete(0,END)
            start_date.delete(0,END)
            position_combo.delete(0,END)
            leave_date1.delete(0,END)
            ni_num_ent.delete(0,END)
            leave_num.set("")
            save["state"]="normal"
            staff_labl.set("Staff details are used for program interface only.")
            fetch_current_employees()
            
        
        
        #,,,Buttons,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        button1_employee=HoverButton(button_label_frame,text="",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT)#,command=lambda:popupmsg("Not yet completed"))
        button1_employee.grid(row=0,column=1,padx=1,pady=1,ipady=5,ipadx=5)
        button_search_employee=HoverButton(button_label_frame,text="SEARCH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lookup_staff_records)#lambda:popupmsg("Not yet completed"))
        button_search_employee.grid(row=0,column=2,padx=1,pady=1,ipady=5,ipadx=10)
        button_remove_employee=HoverButton(button_label_frame,text="DELETE EMPLOYEE",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=delete_employee)#lambda:popupmsg("Not yet completed"))
        button_remove_employee.grid(row=0,column=3,padx=1,pady=1,ipady=5,ipadx=10)
        button_refresh_employee=HoverButton(button_label_frame,text="REFRESH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=clear_employee_box)#lambda:popupmsg("Not yet completed"))
        button_refresh_employee.grid(row=0,column=4,padx=1,pady=1,ipady=5,ipadx=10)
        button_help_employee=HoverButton(button_label_frame,text="HELP!",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=help_staff)#lambda:popupmsg("Not yet completed"))
        button_help_employee.grid(row=0,column=5,padx=1,pady=1,ipady=5,ipadx=10)
        button_exit_employee=HoverButton(button_label_frame,text="EXIT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:notebook.tab(tab_6,state="hidden"))
        button_exit_employee.grid(row=0,column=6,padx=1,pady=1,ipady=5,ipadx=10)
        
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create employee database table,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create a table.
        c.execute("""CREATE TABLE if not exists employees(
            first_name text,
            last_name text,
            address text,
            town text,
            city text,
            county text,
            postcode text,
            phone integar,
            email text,
            start_date integar,
            date_left integar,
            position text,
            ni_num text,
            date_stamp integar)
            """)

        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,employee database functions,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,add new employee to database.
        def add_employee():
            if employee_fname_ent.get()=="":
                messagebox.showerror("Missing Data!","First name needed.")
            elif employee_lname_ent.get()=="":
                messagebox.showerror("Missing Data!","Last name needed.")
            elif employee_address_ent.get()=="":
                messagebox.showerror("Missing Data!","Address needed.")
            elif employee_town_ent.get()=="":
                messagebox.showerror("Missing Data!","Town needed.")
            elif employee_city_ent.get()=="":
                messagebox.showerror("Missing Data!","City needed.")
            elif employee_county_ent.get()=="":
                messagebox.showerror("Missing Data!","County needed.")
            elif employee_post_ent.get()=="":
                messagebox.showerror("Missing Data!","Postcode Needed.")
            elif employee_phone_ent.get()=="":
                messagebox.showerror("Missing Data!","Phone number Needed.")
            elif employee_email_ent.get()=="":
                messagebox.showerror("Missing Data!","Email Needed.")
            elif leave_date1.get()=="":
                messagebox.showerror("Error","No MUST be entered for New starter.")
            elif position_combo.get()=="":
                messagebox.showerror("Missing Data!","Job description needed.")
            elif ni_num_ent.get()=="":
                messagebox.showerror("Error","Please enter\nNI Number.")
            else:
                #,,, Connect to database.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record.
                c.execute("INSERT INTO employees VALUES(:first_name,:last_name,:address,:town,:city,:county,:postcode,:phone,:email,:start_date,:date_left,:position,:ni_num,:date_stamp)",
                        {
                        "first_name":employee_fname_ent.get().capitalize(),
                        "last_name":employee_lname_ent.get().capitalize(),
                        #"id":employee_id_ent.get(),
                        "address":employee_address_ent.get().title(),
                        "town":employee_town_ent.get().title(),
                        "city":employee_city_ent.get().title(),
                        "county":employee_county_ent.get().title(),
                        "postcode":employee_post_ent.get().upper(),
                        "phone":employee_phone_ent.get(),
                        "email":employee_email_ent.get(),
                        "start_date":start_date.get(),
                        "date_left":leave_date1.get(),
                        "position":position_combo.get().capitalize(),
                        "ni_num":ni_num_ent.get().upper(),
                        "date_stamp":dt_stamp,#time.strftime("%d/%m/%Y-%H:%M:%S")
                        })
                
                employee_fname_ent.delete(0,END)
                employee_lname_ent.delete(0,END)
                employee_id_ent.delete(0,END)
                employee_address_ent.delete(0,END)
                employee_town_ent.delete(0,END)
                employee_city_ent.delete(0,END)
                employee_county_ent.delete(0,END)
                employee_post_ent.delete(0,END)
                employee_phone_ent.delete(0,END)
                employee_email_ent.delete(0,END)
                start_date.delete(0,END)
                leave_date1.delete(0,END)
                position_combo.delete(0,END)
                ni_num_ent.delete(0,END)
                #,,,commit changes and close connection
                conn.commit()
                conn.close()
                fetch_current_employees()
                messagebox.showinfo("NEW EMPLOYEE SAVED","New employee details have been saved.")
        
        
        
        
        def leave_click(e):
            #global leave_date
            if leave_date1.get()== "Yes":
                leave_num.set("Leave Date:")
                #leave_date.delete(0,END)
                #leave_date=DateEntry(employee_frame,selectmode="day",date_pattern="dd/mm/yyyy")#,state='readonly')
                #leave_date.grid(row=4,column=3,padx=5,pady=5,sticky=W)
                #print(leave_date.winfo_exists())#if widget exists it will return 1.
                #leave_date.configure(bg="white",fg="black",state="normal",bd=1,relief=SUNKEN)
            else:
                leave_num.set("Start Date")
                #leave_date.destroy()
                #print(leave_date.winfo_exists())##if widget exists it will return 0.#https://youtu.be/2_qUokpB1fw?list=PLCC34OHNcOtoC6GglhF3ncJ5rLwQrLGnV
                #leave_date.delete(0,END)
                #leave_date.configure(bg="systembuttonface",fg="gray",bd=0)
                
                
                
        
        #,,,Employee Details frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        frame1=LabelFrame(box_frame3,text="",bd=0)
        frame1.pack(anchor=W)#grid(row=0,column=0)
        employee_frame = LabelFrame(frame1,text=" Employee Form",font = LARGE_FONT,bd=0)
        employee_frame.grid(row=0,column=0,ipadx=3,ipady=3)#,padx=15,pady=5,ipadx=5,sticky=W)
        
        lbl_14_info=Label(employee_frame,textvariable=staff_labl,font=NORM_FONT,fg="brown")
        lbl_14_info.grid(row=0,column=0,columnspan=4,padx=5,pady=5,sticky=W)

        lbl_employee_fname=Label(employee_frame,text="First Name:")
        lbl_employee_fname.grid(row=1,column=0,padx=5,pady=5,sticky=W)
        employee_fname_ent=Entry(employee_frame) 
        employee_fname_ent.grid(row=1,column=1,padx=5,pady=5,sticky=W)
        employee_fname_ent.focus()
        
        lbl_employee_lname=Label(employee_frame,text="Last Name:")
        lbl_employee_lname.grid(row=1,column=2,padx=5,pady=5,sticky=E)
        employee_lname_ent=Entry(employee_frame) 
        employee_lname_ent.grid(row=1,column=3,padx=5,pady=5,sticky=W)
        
        lbl_position=Label(employee_frame,text="Position:")
        lbl_position.grid(row=2,column=0,padx=5,pady=5,sticky=W)
        position_combo = ttk.Combobox(employee_frame)
        position_combo['values']=['Manager', 'Supervisor', 'General', 'Trainee','Volenteer']
        position_combo.current(0)
        position_combo.grid(row=2,column=1,pady=5,padx=5,sticky=W)
        
        lbl_hours=Label(employee_frame,text="Leaver:")
        lbl_hours.grid(row=2,column=2,padx=5,pady=5,sticky=E)
        leave_date1=ttk.Combobox(employee_frame,width=6)#,state='readonly')
        leave_date1['values']=['No', 'Yes']
        leave_date1.current(0) 
        leave_date1.grid(row=2,column=3,padx=5,pady=5,sticky=W)
        leave_date1.bind("<<ComboboxSelected>>",leave_click)
        
        lbl_employee_address=Label(employee_frame,text="Address:")
        lbl_employee_address.grid(row=3,column=0,padx=5,pady=5,sticky=W)
        employee_address_ent=Entry(employee_frame) 
        employee_address_ent.grid(row=3,column=1,padx=5,pady=5,sticky=W)
        
        lbl_employee_start=Label(employee_frame,text="Start Date:")
        lbl_employee_start.grid(row=3,column=2,padx=5,pady=5,sticky=E)
        start_date=DateEntry(employee_frame,selectmode="day",date_pattern="dd/mm/yyyy")#,state='readonly')
        start_date.grid(row=3,column=3,padx=5,pady=5,sticky=W)
        
        lbl_employee_town=Label(employee_frame,text="Town:")
        lbl_employee_town.grid(row=4,column=0,padx=5,pady=5,sticky=W)
        employee_town_ent=Entry(employee_frame) 
        employee_town_ent.grid(row=4,column=1,padx=5,pady=5,sticky=W)
        global leave_date
        lbl_employee_leave=Label(employee_frame,textvariable=leave_num)
        lbl_employee_leave.grid(row=3,column=2,padx=5,pady=5,sticky=E)

        lbl_employee_city=Label(employee_frame,text="City:")
        lbl_employee_city.grid(row=5,column=0,padx=5,pady=5,sticky=W)
        employee_city_ent=Entry(employee_frame) 
        employee_city_ent.grid(row=5,column=1,padx=5,pady=5,sticky=W)
        lbl_employee_county=Label(employee_frame,text="County:")
        lbl_employee_county.grid(row=5,column=2,padx=5,pady=5,sticky=E)
        employee_county_ent=Entry(employee_frame) 
        employee_county_ent.grid(row=5,column=3,padx=5,pady=5,sticky=W)
        lbl_employee_post=Label(employee_frame,text="Postcode:")
        lbl_employee_post.grid(row=6,column=0,padx=5,pady=5,sticky=W)
        employee_post_ent=Entry(employee_frame) 
        employee_post_ent.grid(row=6,column=1,padx=5,pady=5,sticky=W)
        lbl_employee_phone=Label(employee_frame,text="Phone:")
        lbl_employee_phone.grid(row=6,column=2,padx=5,pady=5,sticky=E)
        employee_phone_ent=Entry(employee_frame) 
        employee_phone_ent.grid(row=6,column=3,padx=5,pady=5,sticky=W)
        lbl_employee_email=Label(employee_frame,text="Email:")
        lbl_employee_email.grid(row=7,column=0,padx=5,pady=5,sticky=W)
        employee_email_ent=Entry(employee_frame,width=30) 
        employee_email_ent.grid(row=7,column=1,columnspan=2,padx=5,pady=5,sticky=W)
        lbl_ni_num=Label(employee_frame,text="NI number:")
        lbl_ni_num.grid(row=8,column=0,padx=5,pady=5,sticky=W)
        ni_num_ent=Entry(employee_frame) 
        ni_num_ent.grid(row=8,column=1,padx=5,pady=5,sticky=W)
        lbl_employee_id=Label(employee_frame,text="ID:")
        lbl_employee_id.grid(row=7,column=2,padx=5,pady=5,sticky=E)
        employee_id_ent=Entry(employee_frame,width=5,relief=FLAT,bd=0,bg="systembuttonface")
        employee_id_ent.grid(row=7,column=3,padx=5,pady=5,sticky=W)
        employ_info=StringVar()
        
        #quickview frame frame,,,,,,,,,,,,,,
        frameemp=Frame(frame1,bg="blue")
        frameemp.grid(row=1,column=0,padx=15,sticky=N)
        #employ_info.set("this is it that is that")
        ref_label=Label(frame1,textvariable=employ_info,font="Verdana, 12",fg="#347083",justify="left")
        ref_label.grid(row=0,column=1,padx=15,sticky=N)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,employee comand button frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        employee_button_frame = LabelFrame(employee_frame,text = '')
        employee_button_frame.grid(row=9,column=0,columnspan=4,padx=2,pady=2)
        save=HoverButton(employee_button_frame,text="SAVE",cursor="hand2",activebackground= '#e2f723',fg="black",command=add_employee)#lambda: popupmsg('Not supported just yet!'))
        save.grid(row=0,column=0,padx=2,pady=2)
        clear_all=HoverButton(employee_button_frame,text="CLEAR ALL",cursor="hand2",activebackground= '#e2f723',fg="black",command=clear_employee_box)#lambda: popupmsg('Not supported just yet!'))
        clear_all.grid(row=0,column=1,padx=2,pady=2)
        update_but=HoverButton(employee_button_frame,text="UPDATE RECORD",cursor="hand2",activebackground= '#e2f723',fg="black",command=update_employee_details)#lambda: popupmsg('Not supported just yet!'))
        update_but.grid(row=0,column=2,padx=2,pady=2)
        
        #def show():
            #notebook.add(tab_3,state="normal")
            #notebook.add(tab_1,state="hidden")
        def select_1():
            notebook.select(1)    
        def select_2():
            notebook.select(2)
        def select_3():
            notebook.select(3)
        def select_4():
            notebook.select(4)
        def select_5():
            notebook.select(5) 
        def select_6():
            notebook.select(6) 
            
        
    
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,,Left side Buttons linked to Class Button,,,,,,,,,,,,,,,,,,,,,,,,,,,
        label = Label(can_2,image=networkpeoplePH,font=LARGE_FONT,bg="white",fg="black")
        label.pack(pady=0,padx=0)#grid(row=0,column=0)#pack(pady=10,padx=10)
        button1 = HoverButton(can_2, text="SYSTEM\nSETTINGS",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'white',bg="#e2f723",fg="black",command=lambda: controller.show_frame(PageOne))
        button1.pack(fill=X,pady=2)#grid(row=2,column=2)#pack()
        button2 = HoverButton(can_2, text="CLIENTS",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'orange',bg="#e2f723",fg="black",command=lambda: controller.show_frame(PageTwo))
        button2.pack(fill=X,pady=2)#grid(row=3,column=1)#pack()
        button3 = HoverButton(can_2, text="BOOKING\nRECORDS",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'white',bg="#e2f723",fg="black",command=lambda: controller.show_frame(PageThree))
        button3.pack(fill=X,pady=2)#grid(row=4,column=0)#pack()
        button4 = HoverButton(can_2, text="INVENTORY",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'orange',bg="#e2f723",fg="black",command=select_1)
        button4.pack(fill=X,pady=2)#grid(row=4,column=0)#pack()
        button5 = HoverButton(can_2, text="SERVICES",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'white',bg="#e2f723",fg="black",command=select_2)
        button5.pack(fill=X,pady=2)#grid(row=4,column=0)#pack()
        button6 = HoverButton(can_2, text="PET\nINFORMATION",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'orange',bg="#e2f723",fg="black",command=select_3)
        button6.pack(fill=X,pady=2)#grid(row=4,column=0)#pack()
        button7 = HoverButton(can_2, text="SELL\nSOMETHING",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'white',bg="#e2f723",fg="black",command=select_4)#lambda: controller.show_frame(PageTwo))
        button7.pack(fill=X,pady=2)#grid(row=3,column=1)#pack()
        button8 = HoverButton(can_2, text="REPORTS",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'orange',bg="#e2f723",fg="black",command=select_5)#lambda: controller.show_frame(PageThree))
        button8.pack(fill=X,pady=2)#grid(row=4,column=0)#pack()
        button9 = HoverButton(can_2, text="EMPLOYEES",height=3,relief=FLAT,cursor="hand2",bd=0,activebackground= 'white',bg="#e2f723",fg="black",command=select_6)#lambda: controller.show_frame(PageThree))
        button9.pack(fill=X,pady=2)#grid(row=4,column=0)#pack()
        
        
        
        def close_app():
            response = messagebox.askquestion("CLOSE", "Do you really want to Exit?\nthis will close the program!.")
            #print(response)
            if response == "yes":
                app.quit()
        
        
        #collect records from database and place in tree view.
        def fetch_current_employees():
            #,,,clear the treeview
            for record in employee_tree.get_children():
                employee_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute("SELECT rowid, * FROM employees WHERE date_left = 'No'")
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            
            for record in records:
                if count % 2==0:
                    employee_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[3],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[12],record[13]), tags=("evenrow",))
                else:
                    employee_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[3],record[4],record[5],record[6],record[7],record[8],record[9],record[10],record[12],record[13]), tags=("oddrow",))
                #,,,increment counter adds 1 to each record entered (ID number),,,,
                count +=1   
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            
        
        #,,, Select Record from bookings treeviw(bind),,,,,,,,,,,,,,,,,,
        def select_employee_record(e):
            employ_info.set("")
            save["state"]="disabled"
            staff_labl.set("If Leaver, change Leaver to (Yes), enter date, then UPDATE.")
            employee_fname_ent.delete(0,END)
            employee_lname_ent.delete(0,END)
            employee_id_ent.delete(0,END)
            employee_address_ent.delete(0,END)
            employee_town_ent.delete(0,END)
            employee_city_ent.delete(0,END)
            employee_county_ent.delete(0,END)
            employee_post_ent.delete(0,END)
            employee_phone_ent.delete(0,END)
            employee_email_ent.delete(0,END)
            start_date.delete(0,END)
            leave_date1.delete(0,END)
            position_combo.delete(0,END)
            #leave_date.delete(0,END)
            ni_num_ent.delete(0,END)
            #,,,Grabe entry on click
            selected=employee_tree.focus()
            #,,,Grab record values.
            values=employee_tree.item(selected,"values")
            #,,,variable names for output to boxes/labels.
            id=values[0]
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            c.execute("SELECT rowid, * FROM employees WHERE rowid = ?",(id,))
            record=c.fetchall() 
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            for record in record:
                employee_id_ent.insert(0,record[0])
                employee_fname_ent.insert(0,record[1])
                employee_lname_ent.insert(0,record[2])
                employee_address_ent.insert(0,record[3])
                employee_town_ent.insert(0,record[4])
                employee_city_ent.insert(0,record[5])
                employee_county_ent.insert(0,record[6])
                employee_post_ent.insert(0,record[7])
                employee_phone_ent.insert(0,record[8])
                employee_email_ent.insert(0,record[9])
                start_date.insert(0,record[10])    
                leave_date1.insert(0,record[11])    
                position_combo.insert(0,record[12])
                ni_num_ent.insert(0,record[13])
            
        
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create treeview frames,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        emp_tree_frame_1=LabelFrame(tab_6,text="",fg="#e2f723",font=20,bd=0)#,width=500,height=250) # Pads (space) inside of frame.
        emp_tree_frame_1.place(x=25,y=560,anchor=W) 
        tree_frame=Frame(emp_tree_frame_1)  
        tree_frame.pack(pady=5,padx=5)
        #,,,Create Treeview Scroll,,,
        tree_scroll=Scrollbar(tree_frame)
        tree_scroll.pack(side=RIGHT,fill=Y)
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,Create the Treeview,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        global employee_tree
        employee_tree=ttk.Treeview(tree_frame,height=2,yscrollcommand=tree_scroll.set,selectmode="extended")
        employee_tree.pack()
        #,,,Configure the scrollbar,,,
        tree_scroll.config(command=employee_tree.yview)
        #,,,Define the Columns,,,
        employee_tree["columns"]=("ID","FIRST NAME","LAST NAME","ADDRESS","TOWN","CITY","COUNTY","POSTCODE","PHONE","EMAIL","START DATE","POSITION","NI NUMBER",)
        #,,,Format the Columns,,,
        employee_tree.column("#0",width=0,stretch=NO)
        employee_tree.column("ID",anchor=W,width=10)
        employee_tree.column("FIRST NAME",anchor=W,width=90)
        employee_tree.column("LAST NAME",anchor=W,width=90)
        employee_tree.column("ADDRESS",anchor=W,width=120)
        employee_tree.column("TOWN",anchor=W,width=100)
        employee_tree.column("CITY",anchor=W,width=100)
        employee_tree.column("COUNTY",anchor=W,width=80)
        employee_tree.column("POSTCODE",anchor=W,width=70)
        employee_tree.column("PHONE",anchor=W,width=80)
        employee_tree.column("EMAIL",anchor=W,width=130)
        employee_tree.column("START DATE",anchor=W,width=80)
        employee_tree.column("POSITION",anchor=W,width=80)
        employee_tree.column("NI NUMBER",anchor=W,width=80)
        
        #,,,Create Headings,,,,,,,,,,,,,,,,,,,,,,,,
        employee_tree.heading("#0",text="",anchor=W)
        employee_tree.heading("ID",text="ID",anchor=W)
        employee_tree.heading("FIRST NAME",text="FIRST NAME",anchor=CENTER)
        employee_tree.heading("LAST NAME",text="LAST NAME",anchor=CENTER)
        employee_tree.heading("ADDRESS",text="ADDRESS",anchor=CENTER)
        employee_tree.heading("TOWN",text="TOWN",anchor=CENTER)
        employee_tree.heading("CITY",text="CITY",anchor=CENTER)
        employee_tree.heading("COUNTY",text="COUNTY",anchor=CENTER)
        employee_tree.heading("POSTCODE",text="POSTCODE",anchor=CENTER)
        employee_tree.heading("PHONE",text="PHONE",anchor=CENTER)
        employee_tree.heading("EMAIL",text="EMAIL",anchor=CENTER)
        employee_tree.heading("START DATE",text="START DATE",anchor=CENTER)
        employee_tree.heading("POSITION",text="POSITION",anchor=CENTER)
        employee_tree.heading("NI NUMBER",text="NI NUMBER",anchor=W)
        
        #,,, create striped rows in treeview,,,,,,,,,,,,,,,,,,,,,
        employee_tree.tag_configure("oddrow",background="white")
        employee_tree.tag_configure("evenrow",background="lightblue")
        
        #,,, Bind the tree,,,,when clicked,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        employee_tree.bind("<ButtonRelease-1>",select_employee_record)
        fetch_current_employees()
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,,,,,,, ribbon buttons,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        ribbon_label_1 = Label(can_4, text=" RIBBON \nMENU",font=LARGE_FONT,bg="white",fg="black")
        ribbon_label_1.grid(row=0,rowspan=3,column=0)#pack(pady=5,padx=10)
        new_button1 = HoverButton(can_4, text="  New   ",relief=FLAT,cursor="hand2",bd=0,activebackground="orange",bg="#e2f723",fg="black",height=1,command=lambda: popupmsg('Not supported just yet!'))
        new_button1.grid(row=0,column=5,padx=0,pady=0)#pack()
        save_button1 = HoverButton(can_4, text="  Save  ",relief=FLAT,cursor="hand2",bd=0,activebackground="orange",bg="#e2f723",fg="black",height=1,command=lambda: popupmsg('Not supported just yet!'))
        save_button1.grid(row=1,column=5,padx=0,pady=0)#pack()
        print_button1 = HoverButton(can_4, text="  Print   ",relief=FLAT,cursor="hand2",bd=0,activebackground="orange",bg="#e2f723",fg="black",height=1,command=print_file)#lambda: popupmsg('Not supported just yet!'))
        print_button1.grid(row=2,column=5,pady=0,padx=0)#pack()
        client_button1 = HoverButton(can_4, text="Quick\n  Booking  \nForm",relief=FLAT,activebackground="white",cursor="hand2",bg="#e2f723",bd=0,fg="black",height=5,command=book_in_window)
        client_button1.grid(row=0,rowspan=3,column=3,pady=0,padx=2,)#pack()
        sup_button1 = HoverButton(can_4, text="  New Supplier  \nForm",relief=FLAT,activebackground="orange",cursor="hand2",bg="#e2f723",bd=0,fg="black",height=5, command=settings_window)
        sup_button1.grid(row=0,rowspan=3,column=4,pady=0,padx=2)#pack()
        booking_button1 = HoverButton(can_4, text="New Client\n  Registration  \nForm",relief=FLAT,activebackground="white",cursor="hand2",bg="#e2f723",bd=0,fg="black",height=5, command=quick_reg_window)
        booking_button1.grid(row=0,rowspan=3,column=1,pady=0,padx=2)#pack()
        pet_booking = HoverButton(can_4, text="New Pet\n  Registration  \nForm",relief=FLAT,activebackground="orange",cursor="hand2",bg="#e2f723",bd=0,fg="black",height=5, command=pet_window)
        pet_booking.grid(row=0,rowspan=3,column=2,pady=0,padx=2)#pack()
        search_button1 = HoverButton(can_4, text="New\nProduct\nForm",relief=FLAT,cursor="hand2",bd=0,activebackground="white",bg="#e2f723",fg="black",height=5,command=new_product_window)#lambda: popupmsg('Not supported just yet!'))
        search_button1.grid(row=0,rowspan=3,column=6,padx=2,pady=0)#pack()
        search_button1 = HoverButton(can_4, text="New\nService",relief=FLAT,cursor="hand2",bd=0,activebackground="orange",bg="#e2f723",fg="black",height=5,command=lambda: popupmsg('Not supported just yet!'))
        search_button1.grid(row=0,rowspan=3,column=7,padx=2,pady=0)#pack()
        help_button1 = HoverButton(can_4, text="  HELP  ",relief=FLAT,activebackground="white",cursor="hand2",bg="#e2f723",bd=0,fg="black",height=5, command = tutorial)
        help_button1.grid(row=0,rowspan=3,column=8,pady=0,padx=2)#pack()
        exit_program_button1 = HoverButton(can_4, text="  CLOSE  ",relief=FLAT,activebackground="red",cursor="hand2",bg="#e2f723",bd=0,fg="black",height=5, command = close_app)
        exit_program_button1.grid(row=0,rowspan=3,column=9,pady=0,padx=2)#pack()
       
#,,,,,Sytem Info Page,,(page one),,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,        
class PageOne(Frame):#,,Anything for Page One System Info put hear,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
    def __init__(self, parent, controller):
        Frame.__init__(self, parent)

        #,,, frames left Page One System Info,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        can_2=Frame(self,relief=FLAT,bd=0,bg="white")
        can_2.pack(side='left',anchor=W,fill='both')#,expand=True)
        #,,main frame Right side,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        canvas_3=Frame(self,relief=RAISED,bd=0)
        canvas_3.pack(side='left',anchor=W,fill='both',expand=True)
        #,,Buttons on Left Frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        
        #,,,load an Image,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        networkpeople=(Image.open("network.png"))
        networkpeople=networkpeople.resize((180,180))#resizes the image.
        networkpeoplePH=ImageTk.PhotoImage(networkpeople)
        network_label=tk.Label(can_2,image=networkpeoplePH)
        network_label.image=networkpeoplePH#this must be put to show the image.
        #network_label.grid(row=0,column=0)
        label = Label(can_2, image=networkpeoplePH, font=LARGE_FONT,bg="white")#,background="yellow")
        label.pack()
        
        button1 = HoverButton(can_2, text="Back to Home Page",font="helvetica 12",activebackground = 'orange',bg="#e2f723",height=3,relief=FLAT,bd=0,fg="black",command=lambda: controller.show_frame(StartPage))
        button1.pack(anchor=W,fill=X)
        top_label_frame=Frame(canvas_3,bd=1)  
        top_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        heading=Label(top_label_frame,text="SYSTEM SETTINGS",font="bold 30",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=25)
        button_label_frame=Frame(canvas_3,bg="light blue",bd=1)  
        button_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        box_frame2=Frame(canvas_3,bd=1)  
        box_frame2.pack(pady=30,padx=30,anchor=W,fill=X)
        global vat_persent
        global system_info
        system_info=StringVar()
        vat_num=StringVar()
        vat_num.set("")
        vat_persent=StringVar()
        vat_persent.set("")
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create Business Details table in database,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create Business Details table.,,,,,,,,,,,,,,,,,,,,,,,
        c.execute("""CREATE TABLE if not exists business_details(
            business_name text,
            address text,
            town text,
            city text,
            county text,
            postcode text,
            phone_number integer,
            email text,
            website text,
            vat_number integer)
            """)
        
        
        
        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        
        #,,,Sets new default Business Details,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def set_business_details():#,,updates Business Details boxes.
            if business_name_ent.get()=="":
                messagebox.showerror("Missing Data!","Business name needed.")
            elif address_ent.get()=="":
                messagebox.showerror("Missing Data!","Address needed.")
            elif address_2_ent.get()=="":
                messagebox.showerror("Missing Data!","Town needed.")
            elif city_ent.get()=="":
                messagebox.showerror("Missing Data!","City needed.")
            elif county_ent.get()=="":
                messagebox.showerror("Missing Data!","County needed.")
            elif post_ent.get()=="":
                messagebox.showerror("Missing Data!","Postcode needed.")
            elif phone_ent.get()=="":
                messagebox.showerror("Missing Data!","Phone number needed.")
            elif email_ent.get()=="":
                messagebox.showerror("Missing Data!","Email needed.")
            #elif web_ent.get()=="":
                #popupmsg("website needed.")
            else:
                #,,, Connect to database.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record.(UPDATE) business_details SET
                c.execute("""UPDATE business_details SET 
                          business_name=:business_name, 
                          address=:address,
                          town=:town,
                          city=:city,
                          county=:county,
                          postcode=:postcode,
                          phone_number=:phone_number,
                          email=:email,
                          website=:website,
                          vat_number=:vat_number""",
                            {
                            "business_name":str.title(business_name_ent.get()),
                            "address":str.title(address_ent.get()),
                            "town":str.title(address_2_ent.get()),
                            "city":str.title(city_ent.get()),
                            "county":str.title(county_ent.get()),
                            "postcode":post_ent.get().upper(),
                            "phone_number":phone_ent.get(),
                            "email":email_ent.get(),
                            "website":web_ent.get(),
                            "vat_number":vat_ent.get(),
                            })
                business_name_ent.delete(0,END)
                address_ent.delete(0,END)
                address_2_ent.delete(0,END)
                city_ent.delete(0,END)
                county_ent.delete(0,END)
                post_ent.delete(0,END)
                phone_ent.delete(0,END)
                email_ent.delete(0,END)
                web_ent.delete(0,END)
                vat_ent.delete(0,END)
                #,,,commit the changes
                conn.commit()
                #,,, close the connection
                conn.close()
                messagebox.showinfo("UPDATED", "Your settings have been\nSuccessfully Updated")
                #phrase=['Dog Lick','Ball Buster','Rug Rats','Only on a Monday']
                #return phrase[random.randint(0,4)]+name
                


        #vat_rate_lbl = Label(company_info_frame, text = 'Set VAT rate %:')
        #vat_rate_lbl.grid(row=1, column=2, padx=5, pady=5,sticky=E)
        #vat_rate_ent = Entry(company_info_frame,width=6)
        #vat_rate_ent.grid(row=1,column=3,sticky=W)
        
        def vat_click(e):
            if vat_ent_combo.get()== "Yes":
                vat_num.set("VAT num:")
                vat_ent.delete(0,END)
                vat_ent.configure(bg="white",fg="black",state="normal",bd=1)
                vat_persent.set("VAT % Rate:")
                department_ent.delete(0,END)
                department_ent.configure(bg="white",fg="black",state="normal",bd=1)
            else:
                vat_num.set("")
                vat_ent.delete(0,END)
                vat_ent.configure(bg="systembuttonface",fg="gray",bd=0)
                vat_ent.delete(0,END)
                vat_persent.set("")
                department_ent.delete(0,END)
                department_ent.configure(bg="white",fg="black",state="disabled",bd=0)
        
            
        #,,,Business details frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,   
        company_info_frame = LabelFrame(box_frame2,text = ' Company Details',bd=0,font = LARGE_FONT)
        company_info_frame.grid(row=0,column=0,padx=10,pady=10,ipadx=5,ipady=5)#,padx=15,pady=(10,5),ipadx=10,ipady=2)#pack(side='left',fill=Y,padx=10,pady=10)
        lbl_1 = Label(company_info_frame,text='This information will apear within invoices.',font=NORM_FONT,fg="brown")
        lbl_1.grid(row=0,column=0,columnspan=4,padx=10,pady=10,sticky=W)
        
        #,,,Business info labels,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        lbl_2 = Label(company_info_frame, text = 'Business Name:',fg='#347083')
        lbl_2.grid(row=1, column=0, padx=5, pady=5,sticky=W)
        company_name=StringVar()
        company_name.set("")
        business_name_ent = Entry(company_info_frame, textvariable=company_name,width=30)
        business_name_ent.grid(row=1,column=1,sticky=W)
        vat_lbl = Label(company_info_frame, text = 'Charge VAT:',fg='#347083')
        vat_lbl.grid(row=1, column=2, padx=5, pady=5,sticky=E)
        vat_ent_combo = ttk.Combobox(company_info_frame,width=6,state="readonly")
        vat_ent_combo['values']=['Yes', 'No']
        vat_ent_combo.current(1)
        vat_ent_combo.grid(row=1,column=3,sticky=W)
        vat_ent_combo.bind("<<ComboboxSelected>>",vat_click)
        
        #vat_ent = Entry(company_info_frame)
        #vat_ent.grid(row=1,column=3,sticky=W)
        
        vat_lbl = Label(company_info_frame, textvariable=vat_num,fg='#347083')
        vat_lbl.grid(row=5, column=2, padx=5, pady=5,sticky=E)
        vat_ent = Entry(company_info_frame,bd=0,state="disabled")
        vat_ent.grid(row=5,column=3,sticky=W)
        
        lbl_service = Label(company_info_frame, textvariable=vat_persent,fg='#347083')
        lbl_service.grid(row=2, column=2, padx=5, pady=5,sticky=W)
        department_ent = Entry(company_info_frame,width=10,bd=0,state="disabled")
        department_ent.grid(row=2,column=3,pady=5,sticky=W)
        
        lbl_3 = Label(company_info_frame, text = 'Address:',fg='#347083')
        lbl_3.grid(row=2, column=0, padx=5, pady=5,sticky=W)
        address_ent = Entry(company_info_frame,width=30)
        address_ent.grid(row=2,column=1,sticky=W)
        lbl_5 = Label(company_info_frame, text = 'City:',fg='#347083')
        lbl_5.grid(row=3, column=2, padx=5, pady=5,sticky=E)
        city_ent = Entry(company_info_frame)
        city_ent.grid(row=3,column=1,sticky=W)
        lbl_4 = Label(company_info_frame, text = 'Town:',fg='#347083')
        lbl_4.grid(row=3, column=0, padx=5, pady=5,sticky=W)
        address_2_ent = Entry(company_info_frame)
        address_2_ent.grid(row=3,column=3,sticky=W)
        lbl_6 = Label(company_info_frame, text = 'County:',fg='#347083')
        lbl_6.grid(row=4, column=0, padx=5, pady=5,sticky=W)
        county_ent = Entry(company_info_frame)
        county_ent.grid(row=4,column=1,sticky=W)
        lbl_7 = Label(company_info_frame, text = 'Postcode:',fg='#347083')
        lbl_7.grid(row=4, column=2, padx=5, pady=5,sticky=E)
        post_ent = Entry(company_info_frame)
        post_ent.grid(row=4,column=3,sticky=W)
        
        lbl_8 = Label(company_info_frame, text = 'Phone Number:',fg='#347083')
        lbl_8.grid(row=5, column=0, padx=5, pady=5,sticky=W)
        phone_ent = Entry(company_info_frame)
        phone_ent.grid(row=5,column=1,sticky=W)
        
        lbl_9 = Label(company_info_frame, text = 'Email:',fg='#347083')
        lbl_9.grid(row=6, column=0, padx=5, pady=5,sticky=W)
        email_ent = Entry(company_info_frame,width=30)
        email_ent.grid(row=6,column=1,sticky=W)
        lbl_10 = Label(company_info_frame, text = 'Website:',fg='#347083')
        lbl_10.grid(row=6, column=2, padx=5, pady=5,sticky=E)
        web_ent = Entry(company_info_frame,width=25)
        web_ent.grid(row=6,column=3,sticky=W)
        
        
        
        def show_business_details():#,,populates business details.
            #,,,clear entry boxes to be filled
            business_name_ent.delete(0,END)
            address_ent.delete(0,END)
            address_2_ent.delete(0,END)
            city_ent.delete(0,END)
            county_ent.delete(0,END)
            post_ent.delete(0,END)
            phone_ent.delete(0,END)
            email_ent.delete(0,END)
            web_ent.delete(0,END)
            vat_ent.delete(0,END)
            #,,, connect to database,,,,,,,,,,.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row from Database (should have only 1 row),,,,,,,
            c.execute("SELECT * FROM business_details")
            records=c.fetchall()
            #,,,Add Record to Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            for record in records:
                business_name_ent.insert(0,record[0])
                address_ent.insert(0,record[1])
                city_ent.insert(0,record[3])
                address_2_ent.insert(0,record[2])
                county_ent.insert(0,record[4])
                post_ent.insert(0,record[5])
                phone_ent.insert(0,record[6])
                email_ent.insert(0,record[7])
                web_ent.insert(0,record[8])
                vat_ent.insert(0,record[9])    
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            system_info.set("")
            
        
        def hide_details():#,,clears all general settings boxes.
            business_name_ent.delete(0,END)
            address_ent.delete(0,END)
            city_ent.delete(0,END)
            address_2_ent.delete(0,END)
            county_ent.delete(0,END)
            post_ent.delete(0,END)
            phone_ent.delete(0,END)
            email_ent.delete(0,END)
            web_ent.delete(0,END)
            vat_ent.delete(0,END)
        
        
        def help_system():
            a1=(f'Click SHOW COMPANY DETAILS to display current settings\n')
            b2=(f'update as needed then click UPDATE COMPANY DETAILS\n')
            c3=(f'hide info with HIDE DETAILS.\n\n')
            d4=(f'Click SHOW GENERAL SETTINGS to display current settings\n')
            e5=(f'update as needed then click UPDATE GENERAL SETTINGS\n')
            f6=(f'hide info with HIDE DETAILS.\n\n')
            g7=(f'If you need a reminder, click HELP!.')
            
            h=a1+b2+c3+d4+e5+f6+g7
            system_info.set(h)
            
        #,,,Business Command Button frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        company_button_frame = LabelFrame(company_info_frame,text = '',bd=0)
        company_button_frame.grid(row=9,column=0,columnspan=4,padx=3,pady=3)#pack(side='left',fill=Y,padx=10,pady=10)
        save=HoverButton(company_button_frame,text="UPDATE COMPANY DETAILS",cursor="hand2",activebackground= '#e2f723',fg="black",command=set_business_details)
        save.grid(row=0,column=0,padx=2,pady=2)
        #show_all=Button(company_button_frame,text="Show Current Details",command=show_business_details)#lambda:popupmsg('Not supported just yet!'))
        #show_all.grid(row=0,column=1,padx=2,pady=2)
        delete_all=HoverButton(company_button_frame,text="HIDE DETAILS",cursor="hand2",activebackground= '#e2f723',fg="black",command=hide_details)#lambda: popupmsg('Not supported just yet!'))
        delete_all.grid(row=0,column=2,padx=2,pady=2)
        
        
        #Full= int(1.0)
        #Half= float(0.5)
        #Quarter= float(0.25)
        
        #global employee_combo
        #def add_employee_products():
            #employee_fname_ent.delete(0,END)
            #if employee_ent.get() not in employee_combo['values']:
            #employee_combo["values"] +=(employee_fname_ent.get(),)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create General Settings table in database,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create General Settings table.,,,,,,,,,,,,,,,,,,,,,,,
        c.execute("""CREATE TABLE if not exists general_settings(
            enclosure_name text,
            max_capacity integer,
            charge_by_the text,
            min_day_charge text,
            first_day_at text,
            last_day_at text,
            taxi_rate float,
            standard_rate float,
            luxury_rate float,
            premium_rate float)
            """)

        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        
        #,,,Sets new default General Settings,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def set_new_general_settings():#,,updates general settings boxes.
            if enclosure_ent.get()=="":
                #popupmsg("Enclosure name needed\nFor reports only.")
                messagebox.showerror("Missing Data!","Enclosure name needed.")
            elif capacity_ent.get()=="":
                #popupmsg("Enter enclosures available.")
                messagebox.showerror("Missing Data!","How many Enclosures?.")
            elif standard_ent.get()=="":
                #popupmsg("Care plan 1 (Day rate) needed.")
                messagebox.showerror("Missing Data!","Standard rate 1 needed.")
            elif luxury_ent.get()=="":
                #popupmsg("Care plan 2 (Day rate) needed.")
                messagebox.showerror("Missing Data!","Rate 2 needed\n\nSet to '0' if only Standard rate available.")
            elif premium_ent.get()=="":
                #popupmsg("Care plan 3 (Day rate) needed.")
                messagebox.showerror("Missing Data!","Rate 3 needed\n\nSet to '0' if only Standard rate available.")
            elif taxi_ent.get()=="":
                #popupmsg("Taxi rate needed.")
                messagebox.showerror("Missing Data!","Pet Taxi rate needed.\n\n(Standard fee) pluss a chargable milage fee.")
            else:
                #,,, Connect to database.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record.(UPDATE) general_settings SET
                c.execute("""UPDATE general_settings SET 
                          enclosure_name=:enclosure, 
                          max_capacity=:max,
                          charge_by_the=:length,
                          min_day_charge=:min,
                          first_day_at=:first,
                          last_day_at=:last,
                          taxi_rate=:taxi,
                          standard_rate=:standard,
                          luxury_rate=:luxury,
                          premium_rate=:premium""",
                            {
                            "enclosure":enclosure_ent.get().capitalize(),
                            "max":int(capacity_ent.get()),
                            "length":length_ent_combo.get(),
                            "min":min_charge_combo.get(),
                            "first":first_ent_combo.get(),
                            "last":last_ent_combo.get(),
                            "taxi":float(taxi_ent.get()),
                            "standard":float(standard_ent.get()),
                            "luxury":float(luxury_ent.get()),
                            "premium":float(premium_ent.get()),
                            })
                
                enclosure_ent.delete(0,END)
                capacity_ent.delete(0,END)
                length_ent_combo.current()
                min_charge_combo.current()
                first_ent_combo.current()
                last_ent_combo.current()
                #category_ent.delete(0,END)
                #department_ent.delete(0,END)
                taxi_ent.delete(0,END)
                standard_ent.delete(0,END)
                luxury_ent.delete(0,END)
                premium_ent.delete(0,END)
                
                #,,,commit the changes
                conn.commit()
                #,,, close the connection
                conn.close()
                messagebox.showinfo("UPDATED", "Your settings have been\nSuccessfully Updated")
        
                
        def hide():#,,clears all general settings boxes.
            enclosure_ent.delete(0,END)
            capacity_ent.delete(0,END)
            length_ent_combo.current(0)
            min_charge_combo.current(2)
            first_ent_combo.current(2)
            last_ent_combo.current(2)
            taxi_ent.delete(0,END)
            standard_ent.delete(0,END)
            luxury_ent.delete(0,END)
            premium_ent.delete(0,END)
            
        
        def show_general_settings():#,,populates current settings.
            #,,,clear entry boxes to be filled
            enclosure_ent.delete(0,END)
            capacity_ent.delete(0,END)
            length_ent_combo.delete(0,END)
            min_charge_combo.delete(0,END)
            first_ent_combo.delete(0,END)
            last_ent_combo.delete(0,END)
            taxi_ent.delete(0,END)
            standard_ent.delete(0,END)
            luxury_ent.delete(0,END)
            premium_ent.delete(0,END)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row from Database (should have only 1 row),,,,,,,
            c.execute("SELECT * FROM general_settings")
            records=c.fetchall()
            #,,,Add Record to Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            for record in records:
                enclosure_ent.insert(0,record[0])
                capacity_ent.insert(0,record[1])
                length_ent_combo.insert(0,record[2])
                min_charge_combo.insert(0,record[3])
                first_ent_combo.insert(0,record[4])
                last_ent_combo.insert(0,record[5])
                taxi_ent.insert(0,format(record[6],'.2f'))#when retreaving from DB it returns a string.
                standard_ent.insert(0,format(record[7],'.2f'))
                luxury_ent.insert(0,format(record[8],'.2f'))
                premium_ent.insert(0,format(record[9],'.2f'))
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            system_info.set("")
            
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,(Kennel) Boarding frame,,(General settings),,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        boarding_frame = LabelFrame(box_frame2,text = ' General Settings',bd=0,font = LARGE_FONT)
        boarding_frame.grid(row=0,column=1,padx=10,pady=10,ipadx=5,ipady=5)#,columnspan=4,padx=10,pady=10,ipadx=10,ipady=2,sticky=N)
        lbl_1 = Label(boarding_frame, text = 'Name of Enclosure?\nie "Run","Unit","Kennel":',fg='#347083')
        lbl_1.grid(row=0, column=0, padx=5, pady=5,sticky=W)
        global enclosure_ent
        enclosure_ent = Entry(boarding_frame,width=14)
        enclosure_ent.grid(row=0,column=1,pady=5,sticky=W)
        lbl_2 = Label(boarding_frame, text = 'Number of Enclosures:',fg='#347083')
        lbl_2.grid(row=0, column=2, padx=5, pady=5,sticky=E)
        capacity_ent = Entry(boarding_frame,width=5)
        capacity_ent.grid(row=0,column=3,pady=5,sticky=W)
        lbl_3 = Label(boarding_frame, text = 'Charge by number of:',fg='#347083')
        lbl_3.grid(row=1, column=0, padx=5, pady=5,sticky=W)
        length_ent_combo = ttk.Combobox(boarding_frame,width=10)
        length_ent_combo['values']=['Days', 'Nights', 'Hours', 'Other']
        length_ent_combo.current(0)
        length_ent_combo.grid(row=1,column=1,pady=5,sticky=W)
        lbl_4 = Label(boarding_frame, text = 'Min Day Charge:',fg='#347083')
        lbl_4.grid(row=1, column=2, padx=5, pady=5,sticky=E)
        min_charge_combo = ttk.Combobox(boarding_frame,width=10)
        min_charge_combo['values']=['Quarter', 'Half', 'Full']
        min_charge_combo.current(2)
        min_charge_combo.grid(row=1,column=3,pady=5,sticky=W)
        lbl_5 = Label(boarding_frame, text = 'Charge First Day at:',fg='#347083')
        lbl_5.grid(row=2, column=0, padx=5, pady=5,sticky=W)
        first_ent_combo = ttk.Combobox(boarding_frame,width=10)
        first_ent_combo['values']=['Quarter', 'Half', 'Full']
        first_ent_combo.current(2)
        first_ent_combo.grid(row=2,column=1,pady=5,sticky=W)
        lbl_6 = Label(boarding_frame, text = 'Charge Last Day at:',fg='#347083')
        lbl_6.grid(row=2, column=2, padx=5, pady=5,sticky=E)
        last_ent_combo = ttk.Combobox(boarding_frame,width=10)
        last_ent_combo['values']=['Quarter', 'Half', 'Full']
        last_ent_combo.current(2)
        last_ent_combo.grid(row=2,column=3,pady=5,sticky=W)
        
        lbl_category = Label(boarding_frame, text = 'Cash Book Balance:',fg='#347083')
        lbl_category.grid(row=3, column=0, padx=5, pady=5,sticky=W)
        category_ent = Entry(boarding_frame,width=10,state="disabled")
        category_ent.grid(row=3,column=1,pady=5,sticky=W)
        
        lbl_rate_1 = Label(boarding_frame, text = 'Standard Rate 1:  £',fg='#347083')
        lbl_rate_1.grid(row=3, column=2, padx=5, pady=5,sticky=E)
        standard_ent = Entry(boarding_frame,width=7)
        standard_ent.grid(row=3,column=3,sticky=W)
        lbl_rate_2 = Label(boarding_frame, text = 'Luxury Rate 2:  £',fg='#347083')
        lbl_rate_2.grid(row=4, column=2, padx=5, pady=5,sticky=E)
        luxury_ent = Entry(boarding_frame,width=7)
        luxury_ent.grid(row=4,column=3,sticky=W)
        lbl_rate_3 = Label(boarding_frame, text = 'Premium Rate 3:  £',fg='#347083')
        lbl_rate_3.grid(row=5, column=2, padx=5, pady=5,sticky=E)
        premium_ent = Entry(boarding_frame,width=7)
        premium_ent.grid(row=5,column=3,sticky=W)
        lbl_taxi = Label(boarding_frame, text = 'Pet taxi rate: £',fg='#347083')
        lbl_taxi.grid(row=6, column=2, padx=5, pady=5,sticky=E)
        taxi_ent = Entry(boarding_frame,width=7)
        taxi_ent.grid(row=6,column=3,pady=5,sticky=W)
        
        boarding_button_frame = LabelFrame(boarding_frame,text = '',bd=0)
        boarding_button_frame.grid(row=7,column=0,columnspan=4,padx=3,pady=3)
        save=HoverButton(boarding_button_frame,text="UPDATE GENERAL SETTINGS",cursor="hand2",activebackground= '#e2f723',fg="black",command=set_new_general_settings)#lambda: popupmsg('Not supported just yet!'))
        save.grid(row=0,column=0,padx=2,pady=2)
        button_show_comp=HoverButton(button_label_frame,text="SHOW COMPANY DETAILS",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=show_business_details)
        button_show_comp.grid(row=0,column=0,padx=1,pady=1,ipady=10,ipadx=20)
        button_gen_set=HoverButton(button_label_frame,text="SHOW GENERAL SETTINGS",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=show_general_settings)
        button_gen_set.grid(row=0,column=1,padx=1,pady=1,ipady=10,ipadx=20)
        button_h=HoverButton(button_label_frame,text="HELP!",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=help_system)#lambda:popupmsg("Not yet completed"))
        button_h.grid(row=0,column=2,padx=1,pady=1,ipady=10,ipadx=20)
        hide_button2=HoverButton(boarding_button_frame,text="HIDE DETAILS",cursor="hand2",activebackground= '#e2f723',fg="black",command=hide)#lambda: popupmsg('Not supported just yet!'))
        hide_button2.grid(row=0,column=2,padx=2,pady=2)
        
        info_help=Frame(canvas_3)  
        info_help.pack(pady=0,padx=0,anchor=W,fill=X)
        #system_info.set("There is a hidden Label here!.")
        lb1=Label(info_help,textvariable=system_info,font="Verdana, 12",fg="#347083",justify=LEFT)
        lb1.grid(row=0,column=0,padx=40,pady=0)
                                    
#,,,,,Client Database,,(Page Two),,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
class PageTwo(Frame):#,,Anything for Client Database put hear,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,

    def __init__(self, parent, controller):
        Frame.__init__(self, parent)
        
        #,,,left and right frames,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        can_2=Frame(self,relief=FLAT,bd=0,bg="white")
        can_2.pack(side='left',anchor=W,fill='both')#,expand=True)
        can_3=Frame(self,relief=FLAT,bd=0)
        can_3.pack(side='left',anchor=W,fill='both',expand=True)
        
        top_label_frame=Frame(can_3,bd=1)  
        top_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        heading=Label(top_label_frame,text="CLIENT RECORDS",font="bold 30",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=25)
        button_label_frame=Frame(can_3,bg="light blue",bd=1)  
        button_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        clientinfo=StringVar()
        
        #,,,create database,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        conn=sqlite3.connect("tree_crm.db")
        #,,,create a cursor.
        c=conn.cursor()
        #,,,create a table in Database.
        c.execute("""CREATE TABLE if not exists clients(
            title text,
            first_name text,
            last_name text,
            id integer,
            address text,
            town text,
            city text,
            county text,
            postcode text,
            phone integar,
            email text)
            """)
        #,,,commit the changes
        conn.commit()
        #,,,close the connection
        conn.close()
        
        
        def query_database():
            #,,,clear the treeview
            for record in my_tree.get_children():
                my_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute("SELECT rowid, * FROM clients ORDER by last_name ASC")
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            
            for record in records:
                if count % 2==0:
                    my_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[5],record[6],record[7],record[8],record[9],record[10],record[11]), tags=("evenrow",))
                else:
                    my_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[5],record[6],record[7],record[8],record[9],record[10],record[11]), tags=("oddrow",))
                #,,,increment counter adds 1 to each record entered (ID number),,,,
                count +=1   
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
        
        
        def talk_search_client():
            engine=pyttsx3.init()
            engine.say(f'''NO results found for that search, check for spelling errors.''')
            engine.runAndWait()
        
            
        def search_records(*args):
            lookup_record=search_entry.get()
            #,,,close the search box
            search.destroy()
            #,,,clear the treeview
            for record in my_tree.get_children():
                 my_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            c.execute("SELECT rowid, * FROM clients WHERE last_name like ?",(lookup_record,))
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0
                 
            for record in records:
                if count % 2==0:
                    my_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[5],record[6],record[7],record[8],record[9],record[10],record[11]), tags=("evenrow",))
                else:
                    my_tree.insert(parent="",index="end",iid=count,text="",values=(record[1],record[2],record[3],record[0],record[5],record[6],record[7],record[8],record[9],record[10],record[11]), tags=("oddrow",))
                #,,,increment counter,,,,
                count +=1 
            if not records:
                query_database()
                #clientinfo.set("NO results found for that search, check for spelling errors or try another name.")
                messagebox.showinfo("SORRY","No results found\nCheck spelling")
                records=talk_search_client()   
            #,,,commit the changes
            conn.commit()
                #,,,close the connection
            conn.close()
            
            
        
        global lookup_records
        def lookup_records():
            global search_entry, search
            search=Toplevel(self)
            search.title("Search Records")
            search.geometry("400x200",)
            search.resizable(False,False)
            #,,,create lable frame
            search_frame=LabelFrame(search,text="Last Name")
            search_frame.pack(padx=10,pady=10)
            #,,,add entry box
            search_entry=Entry(search_frame,font=("Helvetica",18))
            search_entry.bind("<Return>",search_records)#enter key binding(dont forget *args).
            search_entry.pack(padx=20,pady=20,)
            #,,,focus cursor on start.
            search_entry.focus()
            #,,,add a button
            search_button1=Button(search,text="Search Records",command=search_records)
            search_button1.pack(padx=20,pady=20)
            clientinfo.set("")
            clear_entries()
        
        #,,, Select Record from treeviw(bind),,,,,,,,,,,,,,,,,,
        def select_record(e):
            #code to ignore header being clicked.
            #identify the region/area clicked in treeview.
            #https://youtu.be/n5gItcGgIkk
            region_clicked=my_tree.identify_region(e.x, e.y)
            if region_clicked not in ('my_tree','cell'):
                return#returns NONE(nothing)
                #print("Clicked Header")
            else:
                add_button["state"]="disabled"
                #,,,Clear entry boxes.
                title_entry.delete(0,END)
                fn_entry.delete(0,END)
                ln_entry.delete(0,END)
                id_entry.delete(0,END)
                address_entry.delete(0,END)
                town_entry.delete(0,END)
                city_entry.delete(0,END)
                county_entry.delete(0,END)
                postcode_entry.delete(0,END)
                phone_entry.delete(0,END)
                email_entry.delete(0,END)
                #,,,Grab entry
                selected=my_tree.focus()
                #,,,Grab record values.
                values=my_tree.item(selected,"values")
                #,,,output to boxes.
                title_entry.insert(0,values[0])
                fn_entry.insert(0,values[1])
                ln_entry.insert(0,values[2])
                id_entry.insert(0,values[3])
                address_entry.insert(0,values[4])
                town_entry.insert(0,values[5])
                city_entry.insert(0,values[6])
                county_entry.insert(0,values[7])
                postcode_entry.insert(0,values[8])
                phone_entry.insert(0,values[9])
                email_entry.insert(0,values[10])
                #fn_entry.focus()
           
        #,,, Update record,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        def update_record():
            if title_entry.get()=="" or fn_entry.get()=="" or ln_entry.get()=="" \
                or address_entry.get()=="" or town_entry.get()=="" or city_entry.get()=="" \
                or county_entry.get()=="" or postcode_entry.get()=="" or phone_entry.get()=="" \
                or email_entry.get()=="":
                messagebox.showinfo("ERROR","Sellect a record\nto update")
            else:
                messagebox.showinfo("UPDATED", "Record has been UPDATED!")
            #,,,grab record number
            selected=my_tree.focus()
            #,,,update record
            my_tree.item(selected,text="",
                        values=(title_entry.get(),
                        fn_entry.get(),
                        ln_entry.get(),
                        id_entry.get(),
                        address_entry.get(),
                        town_entry.get(),
                        city_entry.get(),
                        county_entry.get(),
                        postcode_entry.get(),
                        phone_entry.get(),
                        email_entry.get()))
            #,,,update database
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            c.execute("""UPDATE clients SET
                title = :title,
                first_name = :first,
                last_name = :last,
                address = :address,
                town = :town,
                city = :city,
                county = :county,
                postcode = :postcode,
                phone = :phone,
                email = :email
        
                WHERE oid = :oid """,
                {"title":title_entry.get(),
                "first":fn_entry.get(),
                "last":ln_entry.get(),
                "address":str.title(address_entry.get()),
                "town":str.title(town_entry.get()),
                "city":city_entry.get(),
                "county":str.title(county_entry.get()),
                "postcode":str.upper(postcode_entry.get()),
                "phone":phone_entry.get(),
                "email":email_entry.get(),
                "oid":id_entry.get()
                })
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            clear_entries()
            #messagebox.showinfo("UPDATED", "Record has been UPDATED!")
                
        #,,,add new record to database.
        def add_record():
            if title_entry.get()=="":
                messagebox.showerror("Error","Title needed.")
            elif fn_entry.get()=="":
                messagebox.showerror("Error","First name needed.")
            elif ln_entry.get()=="":
                messagebox.showerror("Error","Please enter\nLast name.")
            elif address_entry.get()=="":
                messagebox.showerror("Error","Please enter\nAddress.")
            elif town_entry.get()=="":
                messagebox.showerror("Error","Please enter\nTown.")
            elif city_entry.get()=="":
                messagebox.showerror("Error","Please enter\ncity.")
            elif county_entry.get()=="":
                messagebox.showerror("Error","County needed.")
            elif postcode_entry.get()=="":
                messagebox.showerror("Error","Please enter\nPostcode.")
            elif phone_entry.get()=="":
                messagebox.showerror("Error","Phone number needed.")
            elif email_entry.get()=="":
                messagebox.showerror("Error","Email needed.")
            else:
                #,,, Create a database or connect to one that executes.
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,add new record to clients (option 1).
                #client_info=["title_entry.get().capitalize()","fn_entry.get().capitalize()","ln_entry.get().capitalize()","id_entry.get()","address_entry.get().capitalize()","town_entry.get().capitalize()","city_entry.get().capitalize()","county_entry.get().capitalize()","postcode_entry.get().upper()","phone_entry.get()","email_entry.get()"]        
                #c.execute("INSERT INTO clients (title,first,last,id,address,town,city,county,postcode,phone,email) VALUES(?,?,?,?,?,?,?,?,?,?,?)",(client_info[0],client_info[1],client_info[2],client_info[3],client_info[4],client_info[5],client_info[6],client_info[7],client_info[8],client_info[9],client_info[10]))
                #,,,add new record to clients (option 2)
                c.execute("INSERT INTO clients VALUES(:title,:first,:last,:id,:address,:town,:city,:county,:postcode,:phone,:email)",
                        {
                        "title":title_entry.get().capitalize(),
                        "first":fn_entry.get().capitalize(),
                        "last":ln_entry.get().capitalize(),
                        "id":id_entry.get(),
                        "address":address_entry.get().title(),
                        "town":town_entry.get().capitalize(),
                        "city":city_entry.get().capitalize(),
                        "county":county_entry.get().capitalize(),
                        "postcode":postcode_entry.get().upper(),
                        "phone":phone_entry.get(),
                        "email":email_entry.get(),
                        })
                #,,,commit the changes
                conn.commit()
                #,,, close the connection
                conn.close()
                #,,,Clear entry boxes.
                title_entry.delete(0,END)
                fn_entry.delete(0,END)
                ln_entry.delete(0,END)
                id_entry.delete(0,END)
                address_entry.delete(0,END)
                town_entry.delete(0,END)
                city_entry.delete(0,END)
                county_entry.delete(0,END)
                postcode_entry.delete(0,END)
                phone_entry.delete(0,END)
                email_entry.delete(0,END)
                #,,,clear treeview table.
                my_tree.delete(*my_tree.get_children())
                query_database()
                messagebox.showinfo("Added", "Record has been Added!")
            
        
        def highlight_colour():
            highlight_colour=colorchooser.askcolor()[1]
            #,,,update treeview colour
            #,,,Change selected colour,,,
            if highlight_colour:
                style.map("Treeview",
                    background=[("selected",highlight_colour)])
        
        
        
        
        
        def remove_many(*args):
            talk()
            if title_entry.get()=="" or fn_entry.get()=="" or ln_entry.get()=="" or address_entry.get()=="" or town_entry.get()=="" or city_entry.get()=="" or county_entry.get()=="" or postcode_entry.get()=="" or phone_entry.get()=="" or email_entry.get()=="":
                messagebox.showerror("ERROR","Choose a record to delete")
            else:
                app.bell()
                response = messagebox.askyesno("WARNING!!", "This will DELETE Selected\nfrom the Database\nAre You Sure?!.")
                #,,,add logic for message box.
                if response==1:
                    #,,,desegnate selections
                    x=my_tree.selection()
                    #,,,create list of ID's.
                    ids_to_delete=[]
                    #,,,add selections to ids_to_delete list.
                    for record in x:
                        ids_to_delete.append(my_tree.item(record,"values")[3])
                    #,,,Delete from treeview.
                    for record in x:
                        my_tree.delete(record)
                    #,,,create database.    
                    conn=sqlite3.connect("tree_crm.db")
                    #,,,create a cursor.
                    c=conn.cursor()
                    #,,,delete everything from table
                    c.executemany("DELETE FROM clients WHERE oid = ?",[(a,)for a in ids_to_delete])
                    #,,,commit the changes
                    conn.commit()
                        #,,,close the connection
                    conn.close()
                        #,,,clear entries from boxes if filled.
                    clear_entries()
                    
                    messagebox.showinfo("SELECTION DELETED", "Records have been Deleted")
        
        
        def talk():
            global engine
            if id_entry.get()=="" :
                engine=pyttsx3.init()
                engine.say(f''' Select a client first. before you can delete a record,
                           you can search the database using the clients last name or ID number.''')
                #engine.say(my_entry.get())
                #engine.say(my_entry2.get())
                #engine.say(my_entry3.get())
            else:
                engine=pyttsx3.init()
                engine.say(f''' WARNING, you have chosen, {fn_entry.get()},{ln_entry.get()} client ID number {id_entry.get()}, to be deleted from all records,
                        if this is correct, select yes, if this was an Error, select no..''')
            engine.runAndWait()
        
        #,,,clear entry
        def clear_entries():
            my_tree.forget
            #,,,re activate add button when clicked.
            add_button["state"]="normal"
            #,,,Clear entry boxes.
            title_entry.delete(0,END)
            fn_entry.delete(0,END)
            ln_entry.delete(0,END)
            id_entry.delete(0,END)
            address_entry.delete(0,END)
            town_entry.delete(0,END)
            city_entry.delete(0,END)
            county_entry.delete(0,END)
            postcode_entry.delete(0,END)
            phone_entry.delete(0,END)
            email_entry.delete(0,END)
            my_tree.delete(*my_tree.get_children())
            clientinfo.set("")
            query_database()
            #fn_entry.focus() #forget
            
        
        def help_client():
            a=(f'Select a record in view to show information. You can amend the details as needed then click UPDATE RECORD.\n')
            b=(f'Search by selecting SEARCH CLIENTS, enter the last name in the popup box.\n')
            c=(f'')
            d=(f"Click CLEAR/REFRESH to clear the entry boxes and refresh client list.\n\n")
            e=(f'''WARNING... Delete will permanently remove the client - to delete, select a record, click DELETE RECORD, (you will have the option to continue or cancel the deletion).\n\n''')
            g=(f'To add a new client, enter details and click ADD CLIENT.')
            h=(f'If you need a reminder, click HELP!.')
            
            answer=a+b+d+e+g
            clientinfo.set(answer)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create Treeview frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        label_frame_1=LabelFrame(can_3,fg="black",bd=0,font=20)#,width=500,height=250) # Pads (space) inside of frame.
        label_frame_1.pack(padx=30,pady=30,anchor=W) # Pads (space) outside of frame.
        tree_frame=Frame(label_frame_1)  
        tree_frame.pack(padx=5,anchor=W)
        #,,,Create Treeview Scroll,,,
        tree_scroll=Scrollbar(tree_frame)
        tree_scroll.pack(side=RIGHT,fill=Y)
        #,,,Create the Treeview,,,,
        my_tree=ttk.Treeview(tree_frame,height=5,yscrollcommand=tree_scroll.set,selectmode="extended")
        my_tree.pack()
        #,,,Configure the scrollbar,,,
        tree_scroll.config(command=my_tree.yview)
        #,,,Define the Columns,,,
        my_tree["columns"]=("Title","First Name","Last Name","ID","Address","Town","City","County","Postcode","Phone","Email",)
        #,,,Format the Columns,,,
        my_tree.column("#0",width=0,stretch=NO)
        my_tree.column("Title",anchor=W,width=35)
        my_tree.column("First Name",anchor=W,width=130)
        my_tree.column("Last Name",anchor=W,width=130)
        my_tree.column("ID",anchor=CENTER,width=50)
        my_tree.column("Address",anchor=W,width=130)
        my_tree.column("Town",anchor=W,width=130)
        my_tree.column("City",anchor=W,width=120)
        my_tree.column("County",anchor=W,width=120)
        my_tree.column("Postcode",anchor=CENTER,width=100)
        my_tree.column("Phone",anchor=W,width=100)
        my_tree.column("Email",anchor=W,width=130)
        #,,,Create Headings,,,,,,,,,,,,,,,,,,,,,,,,
        my_tree.heading("#0",text="",anchor=W)
        my_tree.heading("Title",text="Title",anchor=W)
        my_tree.heading("First Name",text="First Name",anchor=CENTER)
        my_tree.heading("Last Name",text="Last Name",anchor=CENTER)
        my_tree.heading("ID",text="ID",anchor=CENTER)
        my_tree.heading("Address",text="Address",anchor=CENTER)
        my_tree.heading("Town",text="Town",anchor=CENTER)
        my_tree.heading("City",text="City",anchor=CENTER)
        my_tree.heading("County",text="County",anchor=CENTER)
        my_tree.heading("Postcode",text="Post code",anchor=CENTER)
        my_tree.heading("Phone",text="Phone",anchor=CENTER)
        my_tree.heading("Email",text="Email",anchor=CENTER)
        #,,, create striped rows in treeview,,,,,,,,,,,,,,,,,,,,,
        my_tree.tag_configure("oddrow",background="white")
        my_tree.tag_configure("evenrow",background="lightblue")
        
        #,,,add some style,,,
        style=ttk.Style() 
        #,,,pick a theme,,,
        style.theme_use("default")
        #,,,configure the Treevew Colour,,,
        style.configure("Treeview",
            background="#D3D3D3",
            foreground="black",
            rowheight=25,
            fielbackground="#D3D3D3")

        #,,,Change selected colour in treeview,,,,,,
        global background
        style.map("Treeview",
            background=[("selected","#347083")]) 
        
        #,,, Bind the tree,,,,when clicked,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        my_tree.bind("<ButtonRelease-1>",select_record)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,add record Entry frame and Boxes,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        label_frame_2=LabelFrame(can_3,text="    Client Information",bd=0,fg="#347083")
        label_frame_2.pack(padx=30,anchor=W)#,fill="x",expand="yes")
        
        title_lab=Label(label_frame_2,text="Title:")
        title_lab.grid(row=0,column=0,padx=10,pady=10,sticky=W)
        title_entry=Entry(label_frame_2,width=7)
        title_entry.grid(row=0,column=1,padx=0,pady=0,sticky=W)
        fn_label=Label(label_frame_2,text="First Name:")
        fn_label.grid(row=0,column=2,padx=10,pady=10,sticky=E)
        fn_entry=Entry(label_frame_2)
        fn_entry.grid(row=0,column=3,padx=0,pady=0,sticky=W)
        ln_label=Label(label_frame_2,text="Last Name:")
        ln_label.grid(row=0,column=4,padx=10,pady=10,sticky=E)
        ln_entry=Entry(label_frame_2)
        ln_entry.grid(row=0,column=5,padx=0,pady=0,sticky=W)
        id_label=Label(label_frame_2,text="ID:")
        id_label.grid(row=0,column=6,padx=10,pady=10,sticky=E)
        id_entry=Entry(label_frame_2,relief="flat",bg="#f0f0f0")
        id_entry.grid(row=0,column=7,padx=0,pady=0,sticky=W)
        address_label=Label(label_frame_2,text="Address:")
        address_label.grid(row=1,column=0,padx=10,pady=10,sticky=W)
        address_entry=Entry(label_frame_2)
        address_entry.grid(row=1,column=1,padx=0,pady=0,sticky=W)
        town_label=Label(label_frame_2,text="Town:")
        town_label.grid(row=1,column=2,padx=10,pady=10,sticky=E)
        town_entry=Entry(label_frame_2)
        town_entry.grid(row=1,column=3,padx=0,pady=0,sticky=W)  
        city_label=Label(label_frame_2,text="City:")
        city_label.grid(row=1,column=4,padx=10,pady=10,sticky=E)
        city_entry=Entry(label_frame_2)
        city_entry.grid(row=1,column=5,padx=3,pady=3,sticky=W)
        county_label=Label(label_frame_2,text="County:")
        county_label.grid(row=1,column=6,padx=10,pady=10,sticky=W)
        county_entry=Entry(label_frame_2)
        county_entry.grid(row=1,column=7,padx=3,pady=3,sticky=W)  
        postcode_label=Label(label_frame_2,text="Postcode:")
        postcode_label.grid(row=1,column=8,padx=10,pady=10,sticky=E)
        postcode_entry=Entry(label_frame_2,width=9)
        postcode_entry.grid(row=1,column=9,padx=(0,20),pady=3,sticky=E)
        phone_label=Label(label_frame_2,text="Phone:")
        phone_label.grid(row=2,column=0,padx=10,pady=10,sticky=W)
        phone_entry=Entry(label_frame_2,width=15)
        phone_entry.grid(row=2,column=1,padx=3,pady=3,sticky=W)
        email_label=Label(label_frame_2,text="Email:")
        email_label.grid(row=2,column=2,padx=10,pady=10,sticky=E)
        email_entry=Entry(label_frame_2,width=30)
        email_entry.grid(row=2,column=3,padx=3,pady=3,sticky=W)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,add Buttons to database frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        label_frame_2=LabelFrame(label_frame_1,text="Command Buttons")
        label_frame_2.pack(padx=0,pady=10)#fill="x",expand="yes"

        update_button=HoverButton(button_label_frame,text="UPDATE RECORD",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=update_record)
        update_button.grid(row=0,column=0,padx=0,pady=0,ipady=10,ipadx=30)
        add_button=HoverButton(button_label_frame,text="ADD CLIENT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=add_record)
        add_button.grid(row=0,column=1,padx=0,pady=0,ipady=10,ipadx=20)
        search_button1=HoverButton(button_label_frame,text="SEARCH CLIENTS",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lookup_records)
        search_button1.grid(row=0,column=2,padx=0,pady=0,ipady=10,ipadx=20)
        delete_button=HoverButton(button_label_frame,text="DELETE RECORD",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=remove_many)
        delete_button.grid(row=0,column=3,padx=0,pady=0,ipady=10,ipadx=20)
        clear_button=HoverButton(button_label_frame,text="CLEAR/REFRESH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=clear_entries)
        clear_button.grid(row=0,column=4,padx=0,pady=0,ipady=10,ipadx=20)
        help_button=HoverButton(button_label_frame,text="HELP",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=help_client)#lambda:popupmsg("Not yet completed"))
        help_button.grid(row=0,column=5,padx=0,pady=0,ipady=10,ipadx=20)
        
        info_frame=Frame(can_3)  
        info_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        #clientinfo.set("There is a hidden Label here!.\nClick on a record in the treeview.")
        lb1=Label(info_frame,textvariable=clientinfo,font="Verdana, 12",fg="#347083",justify=LEFT)
        lb1.grid(row=0,column=0,padx=30,pady=20)
        #lb2=Label(info_frame,text="",font="Verdana, 142",fg="#347083",justify=LEFT)
        #lb2.pack()#grid(row=0,column=1,padx=0,pady=20)

        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,Client database buttons (left side),,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,load an Image,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        networkpeople=(Image.open("network.png"))
        networkpeople=networkpeople.resize((180,180))#resizes the image.
        networkpeoplePH=ImageTk.PhotoImage(networkpeople)
        network_label=tk.Label(can_2,image=networkpeoplePH)
        network_label.image=networkpeoplePH#this must be put to show the image.
        #network_label.grid(row=0,column=0)
        label = Label(can_2, image=networkpeoplePH, font=LARGE_FONT,bg="white")#,background="yellow")
        label.pack()
        button1 = HoverButton(can_2, text="Back to Home Page",cursor="hand2",font="helvetica 12",activebackground= 'orange',bg="#e2f723",height=3,relief=FLAT,bd=0,fg="black",command=lambda: controller.show_frame(StartPage))
        button1.pack(anchor=W,fill=X)
        #button2 = Button(can_2, text="Visit Page One",width=20,relief=RAISED,bd=4,command=lambda: controller.show_frame(PageOne))
        #button2.pack(pady=5)
        #button3 = Button(can_2, text="Visit Page Three",width=20,relief=RAISED,bd=4,command=lambda: controller.show_frame(PageThree))
        #button3.pack(pady=5)
        query_database()#,,,shows tree table on start up.
        
#,,,,,Booking records,,(page Three),,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
class PageThree(Frame):#,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,

    def __init__(self, parent, controller):
        Frame.__init__(self, parent)
        
        #,,,left side frame,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        can_2=Frame(self,relief=FLAT,bd=0,bg="white")
        can_2.pack(side='left',anchor=W,fill='both')#,expand=True)
        can_3=Frame(self,relief=FLAT,bd=0)
        can_3.pack(side='left',anchor=W,fill='both',expand=True)
        global bookinginfo
        bookinginfo=StringVar()
        #,,,load an Image,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        networkpeople=(Image.open("network.png"))
        networkpeople=networkpeople.resize((180,180))#resizes the image.
        networkpeoplePH=ImageTk.PhotoImage(networkpeople)
        network_label=tk.Label(can_2,image=networkpeoplePH)
        network_label.image=networkpeoplePH#this must be put to show the image.
        #network_label.grid(row=0,column=0)
        label = Label(can_2, image=networkpeoplePH, font=LARGE_FONT,bg="white")#,background="yellow")
        label.pack()
        button1 = HoverButton(can_2, text="Back to Home Page",font="helvetica 12",activebackground = 'orange',bg="#e2f723",cursor="hand2",height=3,relief=FLAT,bd=0,fg="black",command=lambda: controller.show_frame(StartPage))
        button1.pack(anchor=W,fill=X)
        
        
        #https://youtu.be/Bn1n1diGv_0?list=PLCC34OHNcOtoC6GglhF3ncJ5rLwQrLGnV
        # opens popup file window
        def open_file():
            try:
                filepath = fd.askopenfilename(initialdir="booking_invoice_files",
                                            title="Select a Reciept to open",
                                            filetypes=[("all files","*.*")])#[("text files","*.txt"),("CSV files","*.csv"),("all files","*.*")])
                
                if filepath:
                    #pdf_file=PyPDF2.PdfFileReader(filepath)
                    os.startfile(filepath)
                
                #print(file.read())
            except ValueError:
                messagebox.askretrycancel("FILE NOT FOUND","File did not open\nTry again?.")
                #text=None
            except FileNotFoundError:
                messagebox.showinfo("INFO","Information\nYou closed the dialog box without choseing a file")
                #text=None
        
        
        button_reciept = HoverButton(can_2, text="Search Receipts",font="helvetica 12",activebackground = 'orange',bg="#e2f723",cursor="hand2",height=3,relief=FLAT,bd=0,fg="black",command=open_file)
        button_reciept.pack(anchor=W,pady=2,fill=X)
        
        
        #,,,treeview select cell,,,,,
        def select_record():
            item = calendar_tree.selection()
            for i in item:
                print('You selected: ',calendar_tree.item(i,'values')[14])#gets the cell text(14=booking ref/1=name).
        
        #collect records from database and place in tree view.
        def fetch_bookings():
            button_1["state"]="disabled"
            button_6["state"]="disabled"
            #,,,clear the treeview
            for record in calendar_tree.get_children():
                calendar_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute("SELECT rowid, * FROM kennel_bookings ORDER by in_date DESC")
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
            
            for record in records:
                if count % 2==0:
                    calendar_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[4],record[3],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14]), tags=("evenrow",))
                else:
                    calendar_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[4],record[3],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14]), tags=("oddrow",))
                #,,,increment counter adds 1 to each record entered (ID number),,,,
                count +=1   
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()
            bookinginfo.set("")
            
        #collect records from database and place in tree view.
        '''def fetch_date():
            #,,,clear the treeview
            for record in calendar_tree.get_children():
                calendar_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            #,,,get row using ID,,,,,,,
            c.execute("SELECT rowid, * FROM kennel_bookings WHERE in_date = '21/6/2022' ")
            records=c.fetchall()
            #print(records)
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0 #,,,start of ID
           
            for record in records:
                if count % 2==0:
                    calendar_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[4],record[3],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14]), tags=("evenrow",))
                else:
                    calendar_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[4],record[3],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14]), tags=("oddrow",))
                #,,,increment counter adds 1 to each record entered (ID number),,,,
                count +=1   
            #,,,commit the changes
            conn.commit()
            #,,,close the connection
            conn.close()'''  

            
        def search_bookings():
            lookup_bookings=search_entry.get()
            #,,,close the search box
            search.destroy()
            bookinginfo.set("")
            
            #,,,clear the treeview
            for record in calendar_tree.get_children():
                 calendar_tree.delete(record)
            #,,, Create a database or connect to one.
            conn=sqlite3.connect("tree_crm.db")
            #,,,create a cursor.
            c=conn.cursor()
            c.execute("SELECT rowid, * FROM kennel_bookings WHERE pet_name like ? OR date_stamp = ?",(lookup_bookings,lookup_bookings,))
            records=c.fetchall()
            #,,,Add Record Entry Boxes,,,
            global count
            count = 0
            
            for record in records:
                if count % 2==0:
                    calendar_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[4],record[3],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14]), tags=("evenrow",))
                else:
                    calendar_tree.insert(parent="",index="end",iid=count,text="",values=(record[0],record[1],record[2],record[4],record[3],record[5],record[6],record[7],record[8],record[9],record[10],record[11],record[12],record[13],record[14]), tags=("oddrow",))
                #,,,increment counter,,,,
                count +=1   
            if not records:
                fetch_bookings()
                bookinginfo.set("NO results found for that search, check for spelling errors,\nClick REFRESH or SHOW ALL to see all records.")
                records=talk_search()
                
                #,,,commit the changes
            conn.commit()
                #,,,close the connection
            conn.close()
            
            
        #creates popup box for search.
        def lookup_bookings():
            button_1["state"]="disabled"
            button_6["state"]="disabled"
            global search_entry, search
            search=Toplevel(self)
            search.title("Search Booking Records")
            search.geometry("400x200",)
            search.resizable(False,False)
            #,,,create lable frame
            search_frame=LabelFrame(search,text="Pet Name or booking ref")
            search_frame.pack(padx=10,pady=10)
            #,,,add entry box
            search_entry=Entry(search_frame,font=("Helvetica",18))
            search_entry.pack(padx=20,pady=20,)
            #,,,focus cursor on start.
            search_entry.focus()
            #,,,add a button
            search_button1=Button(search,text="Search Records",command=search_bookings)
            search_button1.pack(padx=20,pady=20)
            
        
        #,,, Select Record from bookings treeviw(bind),,,,,,,,,,,,,,,,,,
        def select_bookings_record(e):
            #code to ignore header being clicked.
            region_clicked=calendar_tree.identify_region(e.x, e.y)
            if region_clicked not in ('calendar_tree','cell'):
                return
                #print("Clicked Header")
            else:
                button_1["state"]="normal"
                button_6["state"]="normal"
                #,,,Grabe entry on click
                selected=calendar_tree.focus()
                #,,,Grab record values.
                values=calendar_tree.item(selected,"values")
                #,,,variable names for output to boxes/labels.
                id=values[0]
                client_name=values[1]
                pet_name=values[2]
                staff=values[3]
                enclosure=values[4]
                pet_type=values[5]
                pick_up=values[6]
                drop_off=values[7]
                care_option=values[8]
                in_date=values[9]
                out_date=values[10]
                units=values[11]
                unite_price=values[12]
                unite_price=float(unite_price)
                total_price=values[13]
                total_price=float(total_price)
                booking_ref=values[14]
                #fn_entry.focus()
                a=(f"Booking ref: {booking_ref}.\n")
                a1=(f"Loged by staff member: {staff}.\n")
                a2=(f"System ID: {id}\n\n")
                b=(f'{client_name} requested accommodation for {pet_name} ({pet_type}) to stay {in_date} untill {out_date}\n')
                c=(f'{care_option} care plan was chosen at £{unite_price:.2f} + vat per day for a total of {units} days.\n')
                d=(f'{pet_name} was to be alocated enclosure {enclosure} on arrival.\n')
                e=(f'Pick-up service requested: {pick_up}\nDrop off service requested: {drop_off}.\n\n')
                f=(f'TOTAL BALANCE for accommodation only is £{total_price:.2f}\n')
                g=(f'To confirm payment and generate a RECIEPT, click CONFIRM PAYMENT and select yes.')
                h=a+a1+a2+b+c+d+e+f+g
                bookinginfo.set(h)
            
            
        
        #,,, Select Record from bookings treeviw(bind),,,,,,,,,,,,,,,,,,
        def select_bookings_record_for_invoice():
            try:
                selected_item=calendar_tree.focus()#selects the row.
                details=calendar_tree.item(selected_item)#gets row ref number.
                tree_cell_content=details.get("values")[14]#gets cell 14 content in row.
                tree_cell_content2=details.get("values")[1]#gets cell 14 content in row.
                #messagebox.showinfo(f'Booking ref',(tree_cell_content2 ) + '\n' + ( tree_cell_content))
                #print('Booking ref selected: ',tree_cell_content)#returns a string.
            except IndexError:
                pass
            response=messagebox.askyesno("CONFIRM PAYMENT", f"Select YES to confirm payment\nfor Booking Ref: {tree_cell_content}\nbooked by {tree_cell_content2}.")
            #,,,add logic for message box.
            if response==1:
                curr_datetime=datetime.now()#retuns (2022-07-13 19:29:37.029066) string format.
                req_format=datetime.strftime(curr_datetime,'%d/%m/%Y %H-%M-%S')
                messagebox.showinfo("RECIEPT GENERATED", f"Your Reciept can be found\nwithin Search Reciepts (left).\n\nunder ref:\n{tree_cell_content} {tree_cell_content2}\n\n(It may take 2 or 3 seconds.)")
                button_1["state"]="normal"
                #,,,Grabe entry on click
                selected=calendar_tree.focus()
                #,,,Grab record values.
                values=calendar_tree.item(selected,"values")
                #,,,variable names for output to boxes/labels.
                id=values[0]
                client_name=values[1]
                pet_name=values[2]
                staff=values[3]
                enclosure=values[4]
                pet_type=values[5]
                pick_up=values[6]
                drop_off=values[7]
                care_option=values[8]
                in_date=values[9]
                out_date=values[10]
                units=values[11]
                units=int(units)
                unit_price=values[12]
                unit_price=float(unit_price)
                total_price=values[13]
                total_price=float(total_price)
                booking_ref=values[14]
                #make_client_invoice doc
                document = Document()
                document.add_picture('C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\network.png',width=Inches(1))
                document.add_heading('Receipt',0)
                pa=document.add_paragraph('ref: ')
                pa.add_run(str(booking_ref))
                #this adds paragraph 1
                p1=document.add_paragraph('Dear ')
                p1.add_run(client_name).bold=True
                p1.add_run(',')
                #this adds paragraph 2
                p2=document.add_paragraph('Thank you for your recent booking of, ')
                p2.add_run(str(units)).bold=True
                p2.add_run(' days ')
                p2.add_run(care_option).bold=True
                p2.add_run(' Accommodation').bold=True
                p2.add_run(f',\nfrom {in_date} until {out_date}, {pet_name} will be placed in enclosure {enclosure} on arrival.\nFor future corrospondance your booking was taken by our staff member {staff}. ')
                p2.add_run(f'{pet_name} has also been given pet ID: {id} for referancing on our system.')
                [document.add_paragraph('')for _ in range(2)]#adds 2 empty lines.
                table=document.add_table(rows=1, cols=5)
                hdr_cells=table.rows[0].cells
                hdr_cells[0].text='Product Name'
                hdr_cells[1].text='Units'
                hdr_cells[2].text='Unit Price'
                hdr_cells[3].text='Vat'
                hdr_cells[4].text='Total Price'
                
                for i in range(5):
                    hdr_cells[i].paragraphs[0].runs[0].font.bold=True
                    
                row_cells=table.add_row().cells
                row_cells[0].text=care_option
                row_cells[1].text=f'{units}'
                row_cells[2].text=f'£{unit_price:,.2f}'
                row_cells[3].text=f'£{total_price - (unit_price * units):,.2f}'
                row_cells[4].text=f'£{total_price:,.2f}'
                
                [document.add_paragraph('')for _ in range(4)]#,adds empty spaces.
                
                p3=document.add_paragraph("We look forward to seeing ")
                p3.add_run(f"{pet_name} on the day. please don't forget to bring any toys or treats that will make {pet_name} comfetable.")
                document.add_paragraph('We appreciate your business and look forward to other bookings in the future.\nFor bank transfers please contact a member of staff.')
                document.add_paragraph(f'Receipt date: {req_format}\nPayment received with thanks.')
                document.add_paragraph(f'Staff member.\n{staff}')
                document.save(r'C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\booking_invoice_files\\'f'{booking_ref} {client_name}.docx')
                
                #converts receipt to pdf
                '''def docx_to_pdf(src,dst):#src is the source file ie docx file,,, dst is destination of pdf you make.
                    word=win32com.client.Dispatch('Word.Application')
                    wdfFormatPDF=17 #,look into other numbers in documentation for other formats.
                    doc=word.Documents.Open(src)
                    doc.SaveAs(dst,FileFormat=wdfFormatPDF)
                    doc.Close()
                    word.Quit()
                #this finds the file path of the docx and places the new pdf in a path.
                docx_to_pdf(r'C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\booking_invoice_files\\'f'{booking_ref} {client_name}.docx',
                            r'C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\booking_invoice_pdf_files\\'f'{booking_ref} {client_name}.pdf')
                #messagebox.showinfo("PDF Generated", "Your Reciept can be found\nin the booking invoice Folder")   
                '''
        
        
            
        def help_bookings():
            a=(f'Click SHOW ALL or REFRESH to display all records\n')
            b=(f'then select a record to display its information.\n\n')
            c=(f'To search click SEARCH and enter the pet name in the popup window.\n\n')
            #d=(f"Click SHOW CURRENT for all bookings as of today's date.\n\n")
            e=(f'WARNING... Deleteing will permanently remove from database.\n')
            f=(f'To remove a record, select a record in view, click DELETE RECORD.\n')
            g=(f'(you will have the option to continue or cancel the deletion.)\n\n')
            h=(f'To confirm payment (generate receipt) select a record and, CONFIRM PAYMENT.')
            
            answer=a+b+c+e+f+g+h
            bookinginfo.set(answer)
        
        
        def talk_search():
            engine=pyttsx3.init()
            engine.say(f'''Sorry, NO results found for that search, check for spelling errors or try another name.''')
            engine.runAndWait()
        
        
        def talk_bookings():
            engine=pyttsx3.init()
            engine.say(f''' WARNING,, you have chosen to delete this record, is this correct,.''')
            engine.runAndWait()
        
        def delete_booking():
            talk_bookings()
            app.bell()
            response = messagebox.askyesno("WARNING!!", "This will permanently DELETE the selected record\n\nAre You Sure?!.")
            #,,,add logic for message box.
            if response==1:
                #,,,desegnate selections
                x=calendar_tree.selection()
                
                #,,,create list of ID's.
                ids_to_delete=[]
                
                #,,,add selections to ids_to_delete list.
                for record in x:
                    ids_to_delete.append(calendar_tree.item(record,"values")[0])
                #,,,Delete from treeview.
                for record in x:
                    calendar_tree.delete(record)
                #,,,create database.    
                conn=sqlite3.connect("tree_crm.db")
                #,,,create a cursor.
                c=conn.cursor()
                #,,,delete everything from table
                c.executemany("DELETE FROM kennel_bookings WHERE oid = ?",[(a,)for a in ids_to_delete])
                #,,,commit the changes
                conn.commit()
                    #,,,close the connection
                conn.close()
                bookinginfo.set("")
                messagebox.showinfo("SELECTION DELETED", "Records have been deleted")
        
        #select a cell within the treeview.
        def get_selected_id():#https://youtu.be/yICGS9Lv86s
            try:
                selected_item=calendar_tree.focus()#selects the row.
                details=calendar_tree.item(selected_item)#gets row ref number.
                tree_cell_content=details.get("values")[14]#gets cell 14 content in row.
                tree_cell_content2=details.get("values")[1]#gets cell 14 content in row.
                messagebox.showinfo(f'Booking ref',(tree_cell_content2 ) + '\n' + ( tree_cell_content))
                #print('Booking ref selected: ',tree_cell_content)#returns a string.
            except IndexError:
                pass
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create frames,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        top_label_frame=Frame(can_3,bd=1)  
        top_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        heading=Label(top_label_frame,text="BOOKING RECORDS",font="bold 30",fg="#347083")
        heading.grid(row=0,column=0,columnspan=3,padx=30,pady=25)
        button_label_frame=Frame(can_3,bg="light blue",bd=1)  
        button_label_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        
        button_1=HoverButton(button_label_frame,text="DELETE RECORD",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,state="disabled",command=delete_booking)#lambda:popupmsg("Not yet completed"))
        button_1.grid(row=0,column=0,padx=0,pady=0,ipady=10,ipadx=30)
        #button_3=HoverButton(button_label_frame,text="LAST MONTH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lambda:popupmsg("Not yet completed"))
        #button_3.grid(row=0,column=2,padx=0,pady=0,ipady=10,ipadx=20)
        button_4=HoverButton(button_label_frame,text="SEARCH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=lookup_bookings)
        button_4.grid(row=0,column=3,padx=1,pady=1,ipady=10,ipadx=20)
        button_2=HoverButton(button_label_frame,text="SHOW ALL",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=fetch_bookings)
        button_2.grid(row=0,column=1,padx=1,pady=1,ipady=10,ipadx=20)
        button_5=HoverButton(button_label_frame,text="REFRESH",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=fetch_bookings)#lambda:popupmsg("Not yet completed"))
        button_5.grid(row=0,column=4,padx=1,pady=1,ipady=10,ipadx=20)
        button_6=HoverButton(button_label_frame,text="CONFIRM PAYMENT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,state="disabled",command=select_bookings_record_for_invoice)#lambda:popupmsg("Not yet completed"))
        button_6.grid(row=0,column=5,padx=1,pady=1,ipady=10,ipadx=20)
        button_7=HoverButton(button_label_frame,text="INDEX",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=get_selected_id)#lambda:popupmsg("Not yet completed"))
        button_7.grid(row=0,column=6,padx=1,pady=1,ipady=10,ipadx=20)
        button_8=HoverButton(button_label_frame,text="HELP!",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,command=help_bookings)#lambda:popupmsg("Not yet completed"))
        button_8.grid(row=0,column=7,padx=1,pady=1,ipady=10,ipadx=20)
        
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,create treeview frames,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        label_frame_1=LabelFrame(can_3,text="",fg="#e2f723",font=20,bd=0)#,width=500,height=250) # Pads (space) inside of frame.
        label_frame_1.pack(padx=30,pady=20,anchor=W) # Pads (space) outside of frame.
        tree_frame=Frame(label_frame_1)  
        tree_frame.pack(pady=5,padx=5)
        #,,,Create Treeview Scroll,,,
        tree_scroll=Scrollbar(tree_frame)
        tree_scroll.pack(side=RIGHT,fill=Y)
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #,,,Create the Treeview,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        global calendar_tree
        calendar_tree=ttk.Treeview(tree_frame,height=5,yscrollcommand=tree_scroll.set,selectmode="extended")
        calendar_tree.pack()
        #,,,Configure the scrollbar,,,
        tree_scroll.config(command=calendar_tree.yview)
        #,,,,,,,,,,,,,,,,,,,,,,use this to hide columns in tree view,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        #column_names=("ID","CLIENT","PET NAME","STAFF","ENCLOSURE","PET TYPE","PICK UP","DROP OFF","CARE PLAN","IN DATE","OUT DATE","UNITS","UNIT PRICE","TOTAL PRICE","BOOKING REF",)
        #calendar_tree.configure(columns=column_names,
                                #displaycolumns=("ID","CLIENT","PET NAME","STAFF","ENCLOSURE","PET TYPE","CARE PLAN","UNITS","UNIT PRICE","TOTAL PRICE","BOOKING REF",))
        #,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        
        #,,,Define the Columns,,,
        calendar_tree["columns"]=("ID","CLIENT","PET NAME","STAFF","ENCLOSURE","PET TYPE","PICK UP","DROP OFF","CARE PLAN","IN DATE","OUT DATE","UNITS","UNIT PRICE","TOTAL PRICE","BOOKING REF",)
        #,,,Format the Columns,,,
        calendar_tree.column("#0",width=0,stretch=NO)
        calendar_tree.column("ID",anchor=W,width=15)
        calendar_tree.column("CLIENT",anchor=W,width=100)
        calendar_tree.column("PET NAME",anchor=CENTER,width=80)
        calendar_tree.column("STAFF",anchor=CENTER,width=90)
        calendar_tree.column("ENCLOSURE",anchor=CENTER,width=80)
        calendar_tree.column("PET TYPE",anchor=CENTER,width=70)
        calendar_tree.column("PICK UP",anchor=CENTER,width=70)
        calendar_tree.column("DROP OFF",anchor=CENTER,width=70)
        calendar_tree.column("CARE PLAN",anchor=CENTER,width=80)
        calendar_tree.column("IN DATE",anchor=CENTER,width=75)
        calendar_tree.column("OUT DATE",anchor=CENTER,width=75)
        calendar_tree.column("UNITS",anchor=CENTER,width=40)
        calendar_tree.column("UNIT PRICE",anchor=CENTER,width=80)
        calendar_tree.column("TOTAL PRICE",anchor=CENTER,width=80)
        calendar_tree.column("BOOKING REF",anchor=CENTER,width=120)
        #,,,Create Headings,,,,,,,,,,,,,,,,,,,,,,,,
        calendar_tree.heading("#0",text="",anchor=W)
        calendar_tree.heading("ID",text="ID",anchor=CENTER)
        calendar_tree.heading("CLIENT",text="CLIENT",anchor=CENTER)
        calendar_tree.heading("PET NAME",text="PET NAME",anchor=CENTER)
        calendar_tree.heading("STAFF",text="STAFF",anchor=CENTER)
        calendar_tree.heading("ENCLOSURE",text="ENCLOSURE",anchor=CENTER)
        calendar_tree.heading("PET TYPE",text="PET TYPE",anchor=CENTER)
        calendar_tree.heading("PICK UP",text="PICK UP",anchor=CENTER)
        calendar_tree.heading("DROP OFF",text="DROP OFF",anchor=CENTER)
        calendar_tree.heading("CARE PLAN",text="CARE PLAN",anchor=CENTER)
        calendar_tree.heading("IN DATE",text="IN DATE",anchor=CENTER)
        calendar_tree.heading("OUT DATE",text="OUT DATE",anchor=CENTER)
        calendar_tree.heading("UNITS",text="UNITS",anchor=CENTER)
        calendar_tree.heading("UNIT PRICE",text="UNIT PRICE",anchor=CENTER)
        calendar_tree.heading("TOTAL PRICE",text="TOTAL PRICE",anchor=CENTER)
        calendar_tree.heading("BOOKING REF",text="BOOKING REF",anchor=CENTER)
        #,,, create striped rows in treeview,,,,,,,,,,,,,,,,,,,,,
        calendar_tree.tag_configure("oddrow",background="white")
        calendar_tree.tag_configure("evenrow",background="lightblue")
        
        #,,, Bind the tree,,,,when clicked,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,,
        calendar_tree.bind("<ButtonRelease-1>",select_bookings_record)
        fetch_bookings()
        
        
        
        info_frame=Frame(can_3)  
        info_frame.pack(pady=0,padx=0,anchor=W,fill=X)
        #bookinginfo.set("There is a hidden Label here!.\nClick on a record in the treeview.")
        lb1=Label(info_frame,textvariable=bookinginfo,font="Verdana, 14",fg="#347083",justify=LEFT)
        lb1.grid(row=0,column=0,padx=30,pady=20)
        lb2=Label(info_frame,text="",font="Verdana, 12",fg="#347083",justify=LEFT)
        lb2.grid(row=0,column=1,padx=0,pady=20)
        #button_6=HoverButton(can_3,text="CONFIRM\nPAYMENT",font=LARGE_FONT,activebackground = 'white',cursor="hand2",bg="light blue",fg="black",relief=FLAT,width=15,state="disabled",command=select_bookings_record_for_invoice)#lambda:popupmsg("Not yet completed"))
        #button_6.place(x=870,y=400)
        

app = LangtonAndRoberts()

#,,,sets opening frame position on start up,,,,,,,,,
#app_width=1450
#app_height=600
#screen_width=app.winfo_screenwidth()
#screen_height=app.winfo_screenheight()
#x=(screen_width / 2)-(app_width / 2)
#y=(screen_height /2)-(app_height / 2)
#app.geometry(f'{app_width}x{app_height}+{int(x)}+{int(y)}')

# ,,, make start up screen fit any monitor ,,,,,,,
width_value=app.winfo_screenwidth()
height_value=app.winfo_screenheight()
app.geometry(f"{width_value}x{height_value}+0+0")




#query_database()
#app.bell()
app.mainloop()