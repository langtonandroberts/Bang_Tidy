#https://youtu.be/WmlQC2Rv1HA
#,makes an invoice template in docx/word and converts into a pdf.
#,make a Word Docx File in Python.
#https://youtu.be/7ljbGXSxMw4

#,also look at https://youtu.be/q70xzDG6nls to creat PDF with Python
#https://youtu.be/hslXVn6r0ck

from pydoc import doc
from turtle import width
from docx import Document
from docx.shared import Inches
import win32com.client



invoice_ref=123456
client="Mr Joseph Roberts"#,this will show as file name.
unit=7
product="Accomodation"
price=23.0

def make_client_invoice():
    document = Document()
    document.add_picture('C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\network.png',width=Inches(1))
    document.add_heading('Invoice',0)
    pa=document.add_paragraph('Invoice ref: ')
    pa.add_run(str(invoice_ref))
    #this adds paragraph 1
    p1=document.add_paragraph('Dear ')
    p1.add_run(client).bold=True
    p1.add_run(',')
    #this adds paragraph 2
    p2=document.add_paragraph('Thank you for your recent booking of, ')
    p2.add_run(str(unit)).bold=True
    p2.add_run(' days ')
    p2.add_run(product).bold=True
    p2.add_run(',')
    [document.add_paragraph('')for _ in range(2)]#adds 2 empty lines.
    table=document.add_table(rows=1, cols=4)
    hdr_cells=table.rows[0].cells
    hdr_cells[0].text='Product Name'
    hdr_cells[1].text='Units'
    hdr_cells[2].text='Unit Price'
    hdr_cells[3].text='Total Price'
    
    for i in range(4):
        hdr_cells[i].paragraphs[0].runs[0].font.bold=True
        
    row_cells=table.add_row().cells
    row_cells[0].text=product
    row_cells[1].text=f'{unit}'
    row_cells[2].text=f'{price:,.2f}'
    row_cells[3].text=f'{unit*price:,.2f}'
    
    [document.add_paragraph('')for _ in range(2)]#,adds empty spaces.
    
    document.add_paragraph('We appreciate your business and look forward to other bookings in the future.\nfor bank transfers please use.\nsort: 01-23-45 acc: 123456789.')
    document.add_paragraph('Sincerely')
    document.add_paragraph('Mr Test Invoice')
    document.save(f'{client}.docx')

make_client_invoice()#pass the values needed in the def.



#,change docx file into pdf file.    
def docx_to_pdf(src,dst):#src is the source file ie docx file,,, dst is destination of pdf you make.
    word=win32com.client.Dispatch('Word.Application')
    wdfFormatPDF=17 #,look into other numbers in documentation for other formats.
    doc=word.Documents.Open(src)
    doc.SaveAs(dst,FileFormat=wdfFormatPDF)
    doc.Close()
    word.Quit()
   
#this finds the file path of the docx and places the new pdf in a path.
docx_to_pdf(r'C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\Mr Joseph Roberts.docx',
            r'C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\\Bang_Tidy\\Mr Joseph Roberts.pdf')



'''            
#,,,sends an email with the attached invoice.(uses Outlook)            
def send_email(name,to_addr,attachment):
    outlook = win32com.client.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)
    mail.To = to_addr
    mail.Subject = 'Invoice from Pythoninoffice'
    mail.Body = f'Dear {name}, Please find attached invoice'
    mail.Attachments.Add(attachment)
    mail.Send()

send_email('Mr Tom Smith','joerobertsmail@aol.co.uk',r'C:\\Users\\Invate\\OneDrive\\Desktop\\Python\\Python cheat sheets\\Python Codes Projects\Bang_Tidy\\Mr Tom Smith.pdf')
'''